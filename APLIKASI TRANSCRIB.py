import streamlit as st
import streamlit.components.v1 as components
import requests
import json
import pandas as pd
import firebase_admin
from firebase_admin import credentials, firestore, auth
from datetime import datetime, timedelta
import time
import io
import math
import os
import tempfile
import urllib.parse
import random
import zlib
import base64
import re

os.environ['TZ'] = 'Asia/Jakarta'
try:
    time.tzset()
except AttributeError:
    pass

# === IMPORT UNTUK EXPORT DOCX ===
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Konfigurasi Halaman
st.set_page_config(page_title="TranscribX & SOP Gen", layout="wide", initial_sidebar_state="expanded")

# =====================================================================
# CSS CUSTOM UNTUK ANIMASI, CARD, DAN UI ENTERPRISE
# =====================================================================
custom_css = """
<style>
/* Menyembunyikan elemen bawaan Streamlit */
#MainMenu {visibility: hidden;}
header {background-color: transparent;}
footer {visibility: hidden;}

/* Menyembunyikan Toolbar Streamlit */
[data-testid="stToolbar"] { visibility: hidden; }

/* Metric Card Styling */
.metric-card {
    background: linear-gradient(135deg, #ffffff 0%, #f8fafc 100%);
    padding: 20px;
    border-radius: 16px;
    box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05), 0 2px 4px -1px rgba(0,0,0,0.03);
    border: 1px solid #e2e8f0;
    transition: transform 0.3s ease, box-shadow 0.3s ease;
    position: relative;
    overflow: hidden;
}
.metric-card:hover {
    transform: translateY(-5px);
    box-shadow: 0 10px 15px -3px rgba(0,0,0,0.1), 0 4px 6px -2px rgba(0,0,0,0.05);
}
.metric-card::before { content: ""; position: absolute; top: 0; left: 0; width: 6px; height: 100%; }
.metric-total::before { background-color: #8b5cf6; }
.metric-aktif::before { background-color: #10b981; }
.metric-nonaktif::before { background-color: #ef4444; }
.metric-admin::before { background-color: #f59e0b; }

/* Pulse Animation for Sidebar Profile */
@keyframes pulse-border {
    0% { box-shadow: 0 0 0 0 rgba(59, 130, 246, 0.4); }
    70% { box-shadow: 0 0 0 8px rgba(59, 130, 246, 0); }
    100% { box-shadow: 0 0 0 0 rgba(59, 130, 246, 0.4); }
}
@keyframes pulse-border-admin {
    0% { box-shadow: 0 0 0 0 rgba(245, 158, 11, 0.4); }
    70% { box-shadow: 0 0 0 8px rgba(245, 158, 11, 0); }
    100% { box-shadow: 0 0 0 0 rgba(245, 158, 11, 0.4); }
}
.profile-card {
    border-radius: 16px;
    padding: 20px;
    color: white;
    margin-bottom: 20px;
    position: relative;
    overflow: hidden;
    transition: all 0.3s ease;
}
.profile-card.admin { background: linear-gradient(135deg, #f59e0b 0%, #d97706 100%); animation: pulse-border-admin 2.5s infinite; }
.profile-card.user-active { background: linear-gradient(135deg, #3b82f6 0%, #1d4ed8 100%); animation: pulse-border 2.5s infinite; }
.profile-card.user-warning { background: linear-gradient(135deg, #ef4444 0%, #b91c1c 100%); }
.profile-card.user-inactive { background: linear-gradient(135deg, #64748b 0%, #334155 100%); }

/* Mempercantik Tombol Utama Streamlit */
.stButton>button {
    border-radius: 8px !important;
    transition: all 0.3s ease !important;
    font-weight: bold !important;
}
.stButton>button:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 12px rgba(59, 130, 246, 0.3) !important;
}

/* Mempercantik Tab Navigasi Dalam Aplikasi */
.stTabs [data-baseweb="tab-list"] { gap: 8px; }
.stTabs [data-baseweb="tab"] { border-radius: 8px 8px 0px 0px; padding: 10px 20px; background-color: #f1f5f9; }
.stTabs [aria-selected="true"] { background-color: #3b82f6 !important; color: white !important; }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# =====================================================================
# INISIALISASI FIREBASE ADMIN
# =====================================================================
if not firebase_admin._apps:
    cred = credentials.Certificate(dict(st.secrets["firebase_admin"]))
    firebase_admin.initialize_app(cred)

db = firestore.client()

# =====================================================================
# DATA PAKET LANGGANAN & HELPER FUNGSI FIREBASE
# =====================================================================
PAKET_LANGGANAN = {
    "BASIC": {"ai_limit": 8, "upload_limit": 1, "durasi_hari": 30},
    "EXECUTIVE": {"ai_limit": 25, "upload_limit": 5, "durasi_hari": 30},
    "MASTER": {"ai_limit": 60, "upload_limit": 15, "durasi_hari": 30},
    "NON-AKTIF": {"ai_limit": 0, "upload_limit": 0, "durasi_hari": 0}
}

ADMIN_EMAILS_CONFIG = st.secrets.get("ADMIN_EMAILS", [])
if not ADMIN_EMAILS_CONFIG:
    single_admin = st.secrets.get("ADMIN_EMAIL", "")
    ADMIN_EMAILS_CONFIG = [single_admin] if single_admin else ["sutsuga.gery@gmail.com"]

def is_admin():
    return st.session_state.get("user_email") in ADMIN_EMAILS_CONFIG

def hitung_sisa_hari(tanggal_berakhir_str):
    if not tanggal_berakhir_str: return 0
    try:
        tanggal_berakhir = datetime.fromisoformat(tanggal_berakhir_str) if isinstance(tanggal_berakhir_str, str) else tanggal_berakhir_str.replace(tzinfo=None)
        selisih = tanggal_berakhir - datetime.now()
        sisa = math.ceil(selisih.total_seconds() / 86400)
        return sisa if sisa > 0 else 0
    except: return 0

def cek_dan_update_status_kadaluarsa(uid, user_data):
    if not user_data or user_data.get("status_subscription", "non-aktif") == "non-aktif" or user_data.get("email") in ADMIN_EMAILS_CONFIG:
        return False
    if hitung_sisa_hari(user_data.get("tanggal_berakhir")) <= 0:
        db.collection("users").document(uid).update({
            "status_subscription": "non-aktif", "paket": "NON-AKTIF",
            "kuota_ai": 0, "kuota_upload": 0, "tanggal_kadaluarsa": datetime.now().isoformat()
        })
        return True
    return False

def get_user_login_history(uid):
    try:
        return [h.to_dict() for h in db.collection("users").document(uid).collection("login_history").order_by("timestamp", direction=firestore.Query.DESCENDING).limit(10).stream()]
    except: return []

def record_login(uid, email):
    try:
        waktu_wib = (datetime.utcnow() + timedelta(hours=7)).isoformat()
        db.collection("users").document(uid).collection("login_history").add({"timestamp": waktu_wib, "email": email, "platform": "Streamlit Cloud"})
        db.collection("users").document(uid).update({"last_login": waktu_wib, "login_count": firestore.Increment(1), "is_online": True})
    except: pass

def delete_user(uid, email):
    try:
        docs = db.collection("users").document(uid).collection("login_history").stream()
        for doc in docs: doc.reference.delete()
        db.collection("users").document(uid).delete()
        auth.delete_user(uid) # Hapus dari Firebase Auth
        return True, f"User {email} berhasil dihapus permanen dari semua sistem."
    except Exception as e: return False, f"Gagal menghapus user: {str(e)}"

def get_active_users_count():
    try: 
        # Cek siapa saja yang aktif dalam 5 menit terakhir
        batas_waktu = (datetime.utcnow() + timedelta(hours=7) - timedelta(minutes=5)).isoformat()
        users = db.collection("users").where("last_active", ">=", batas_waktu).stream()
        return sum(1 for u in users)
    except: return 0

def reset_user_kuota(uid, paket):
    if paket in PAKET_LANGGANAN and paket != "NON-AKTIF":
        db.collection("users").document(uid).update({
            "kuota_ai": PAKET_LANGGANAN[paket]["ai_limit"],
            "kuota_upload": PAKET_LANGGANAN[paket]["upload_limit"],
            "reset_kuota_terakhir": datetime.now().isoformat()
        })
        return True
    return False

def login_firebase(email, password):
    return requests.post(f"https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword?key={st.secrets['FIREBASE_WEB_API_KEY']}", json={"email": email, "password": password, "returnSecureToken": True}).json()

def register_firebase(email, password):
    return requests.post(f"https://identitytoolkit.googleapis.com/v1/accounts:signUp?key={st.secrets['FIREBASE_WEB_API_KEY']}", json={"email": email, "password": password, "returnSecureToken": True}).json()

def check_subscription(uid):
    doc_ref = db.collection("users").document(uid)
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict()
        cek_dan_update_status_kadaluarsa(uid, data)
        data = doc_ref.get().to_dict()
        if "paket" not in data and data.get("status_subscription") == "aktif":
            data.update({"paket": "BASIC", "kuota_ai": PAKET_LANGGANAN["BASIC"]["ai_limit"], "kuota_upload": PAKET_LANGGANAN["BASIC"]["upload_limit"], "tanggal_berakhir": (datetime.now() + timedelta(days=30)).isoformat()})
            doc_ref.update({"paket": data["paket"], "kuota_ai": data["kuota_ai"], "kuota_upload": data["kuota_upload"], "tanggal_berakhir": data["tanggal_berakhir"]})
        return data
    return {"status_subscription": "non-aktif", "paket": "NON-AKTIF"}

def cek_reset_kuota_bulanan(uid, user_data):
    if not user_data or user_data.get("status_subscription") != "aktif" or user_data.get("email") in ADMIN_EMAILS_CONFIG or user_data.get("paket", "BASIC") not in PAKET_LANGGANAN: return False
    sekarang, perlu_reset = datetime.now(), False
    if rkt := user_data.get("reset_kuota_terakhir"):
        try: perlu_reset = (sekarang - (datetime.fromisoformat(rkt) if isinstance(rkt, str) else rkt.replace(tzinfo=None))).days >= 30
        except: perlu_reset = True
    elif tm := user_data.get("tanggal_mulai"):
        try: perlu_reset = (sekarang - (datetime.fromisoformat(tm) if isinstance(tm, str) else tm.replace(tzinfo=None))).days >= 30
        except: perlu_reset = True
    else: perlu_reset = True
    
    if perlu_reset:
        reset_user_kuota(uid, user_data["paket"])
        st.session_state["user_kuota_ai"], st.session_state["user_kuota_upload"] = PAKET_LANGGANAN[user_data["paket"]]["ai_limit"], PAKET_LANGGANAN[user_data["paket"]]["upload_limit"]
        return True
    return False

# =====================================================================
# HELPER FUNCTIONS UNTUK SOP GENERATOR
# =====================================================================
def clean_json_response(raw_text):
    text = re.sub(r'```json', '', raw_text, flags=re.IGNORECASE)
    text = re.sub(r'```', '', text)
    return text.strip()

def encode_plantuml(text):
    compressed = zlib.compress(text.encode('utf-8'))[2:-4]
    b64 = base64.b64encode(compressed).decode('utf-8')
    plantuml_alphabet = '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz-_'
    standard_alphabet = 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/'
    b64 = b64.translate(str.maketrans(standard_alphabet, plantuml_alphabet))
    return b64

def get_png_data_uri(shape_type):
    if shape_type == 'start':
        b64 = "iVBORw0KGgoAAAANSUhEUgAAAFAAAAA4CAYAAABqtn+aAAABZ0lEQVR4nO2bUZLCIBAFYWvvf2X2Q9GEDDDJi2sB3QfA2PVGEZ4xpRTgOj/ffoDRQaAIAkUQKIJAEQSKIFDk966FYozDbShTSlFdIyob6RGl1bgq85JAS9yIv2hiPDo7K/KUwFLciNJqlDK9It0Ct/JmEleyFemR6PoWXkVeCPv35/mM7yYwLzK7OIucxlYSmwlcWV4I7/fdSmJV4OryMj2JpsCZ9nd3YnlpjvDq6cu0PBwEMro2tVHmMEFkJ5D0tbFSSAJFECjyEsj4+ijHmASKIFAEgSIIFEGgCAJFECiCQJGXwHxsbV31wZvymJ8EiiBQZCeQMW5j3dKRQJGDQFJoU7sj7t0Lf/CRxqHlwRR4R29uRiwv1QQyyg969Y7evfDSEj3dGFe9baV2VgjnKm6ubcx2kdnTeLYfSEP1yccbqsWL0ZHOa9DSf/CvLX1zoQFlfv1/IsBhggwCRRAogkARBIogUOQPuTLHQ7a2XJEAAAAASUVORK5CYII="
    elif shape_type == 'process':
        b64 = "iVBORw0KGgoAAAANSUhEUgAAAFAAAAA4CAYAAABqtn+aAAAAsklEQVR4nO3bMQoCMRBAUUf2/lceK0W0WPQXIrxXhVTLZ4oEsrO7F753/fUH/DsBIwEjASMBIwEjAaPjvpgZB8IP7e6YwOh43XAzOTczj7UJjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIwEjASMBIzeHpk/P6DmnAmMxm8NjQmMBIwEjASMBIwEjG61eQ1yRP1zEAAAAABJRU5ErkJggg=="
    elif shape_type == 'decision':
        b64 = "iVBORw0KGgoAAAANSUhEUgAAAFAAAAA4CAYAAABqtn+aAAAB3klEQVR4nO2b4a6DIAxGy33/d+b+0kFjK1oKH9gvWSKZo+Ts6IZiyjlT5H3+Zg9g9cADTClBHyLQAA94yBAT4jlQA5ZzTiPHchc4A0t4OefzdfU+QmAM5GCuxpVSLR+CjRAAuXUN+5/bsyFOBdhinfLZqj0L5DSAT61T+jm3Z0AcDtBindJn1R4JcijAXtYp/Z/boyAOAehhnVKranuDdAfobZ1S99z2hOgGcKR1yhiqtgdIF4CzrJPiaWP3qRwaPCJynQp2MxDhkL2LxyHdBSCidVp6HtImgCtYJ6WXja8BrmadFKuNjwGubJ0Ui42PAO5inZQ3NjYB3NE6KU9tvAW4u3VSWm0UAX7JOiktNl4C/Kp1UjQbkzbNCXi/SDaeAMO6tnAbw8DG3BrIdg4bizSfA9mHPm/j619h1sknbTT/D2SdfcbG7jMR1vnWNrrNhVmR7WwcdjWGFd3CxuHXA1nxZW2cfkWaDWYpG2HuiVQdLWAj7F25qkNQG71urne/sV4Ojn/js+K5MiHWxlhrxOosY51YH2isFytUjTVjjbSxdqzStyWeEzEmnlQyBgIgUZuNKNaVgQF4RLIRyboycACJ9HXMSPCIQAEeYTZCgTsC98B1mQMaKjwicANXyD8CavS5t3haHwAAAABJRU5ErkJggg=="
    else:
        b64 = ""
    return f"data:image/png;base64,{b64}"

def render_numbered_list(items):
    if not items: return "-"
    cleaned_items = [re.sub(r'^[\d]+[\.\)]?\s*', '', str(i)) for i in items]
    li_elements = "".join([f"<li>{i}</li>" for i in cleaned_items])
    return f'<ol style="margin:0; padding-left:20px;">{li_elements}</ol>'

def sanitize_plantuml(puml_raw):
    puml_code = re.sub(r'```plantuml', '', puml_raw, flags=re.IGNORECASE)
    puml_code = re.sub(r'```', '', puml_code).strip()
    puml_code = re.sub(r'@startuml', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'@enduml', '', puml_code, flags=re.IGNORECASE).strip()
    puml_code = re.sub(r'^\s*(actor|participant|component|swimlane).*$', '', puml_code, flags=re.IGNORECASE | re.MULTILINE)
    puml_code = re.sub(r'(then\s*\([^\)]+\)\s*\n)\s*\|[^|]+\|\s*\n', r'\1', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'(else\s*\([^\)]+\)\s*\n)\s*\|[^|]+\|\s*\n', r'\1', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'(else\s*\n)\s*\|[^|]+\|\s*\n', r'\1', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'\brepeat\b', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'\bend repeat\b.*', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'\bwhile\b\s*\(.*?\)', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'\bendwhile\b.*', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'\blabel\b\s+\w+\s*;', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'\bgoto\b\s+\w+\s*;', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'\bbreak\b;?', '', puml_code, flags=re.IGNORECASE)
    puml_code = re.sub(r'^(\s*:[^;\n]+)(?<!;)$', r'\1;', puml_code, flags=re.MULTILINE)

    if_count = len(re.findall(r'\bif\s*\(', puml_code, flags=re.IGNORECASE))
    endif_count = len(re.findall(r'\bendif\b', puml_code, flags=re.IGNORECASE))
    while if_count > endif_count:
        puml_code += '\nendif'
        endif_count += 1

    modern_theme = """
skinparam handwritten false
skinparam BackgroundColor transparent
skinparam defaultFontSize 14
skinparam margin 30
skinparam padding 15
skinparam SwimlaneBorderColor #94a3b8
skinparam SwimlaneTitleBackgroundColor #f1f5f9
skinparam SwimlaneTitleFontColor #0f172a
skinparam SwimlaneTitleFontSize 18
skinparam ActivityBackgroundColor #eff6ff
skinparam ActivityBorderColor #2563eb
skinparam ActivityFontSize 14
skinparam ActivityDiamondBackgroundColor #fef3c7
skinparam ActivityDiamondBorderColor #d97706
skinparam ActivityDiamondFontSize 14
skinparam ArrowColor #475569
skinparam ArrowFontColor #b91c1c
skinparam ArrowFontName Arial
skinparam ArrowFontSize 13
skinparam ArrowFontStyle bold
"""
    return f"@startuml\n{modern_theme}\n{puml_code}\n@enduml"

def render_akreditasi_output(full_data, rs_name, title, no_dok, dir_name, dir_nip, logo_url):
    data_cover = full_data.get('cover', {})
    data_flow = full_data.get('flow', {})
    pelaksana_list = data_flow.get('pelaksana_list', ["Petugas"])
    steps = data_flow.get('langkah', [])
    # --- TAMBAHAN: CONVERT LOGO KE BASE64 AGAR TIDAK RUSAK DI WORD ---
    logo_b64_html = logo_url
    if logo_url.startswith("http"):
        try:
            res_logo = requests.get(logo_url, timeout=5)
            if res_logo.status_code == 200:
                ctype = res_logo.headers.get('content-type', 'image/png')
                b64_logo = base64.b64encode(res_logo.content).decode('utf-8')
                logo_b64_html = f"data:{ctype};base64,{b64_logo}"
        except:
            pass # Jika gagal download, biarkan pakai URL asli

    embedded_img_html = ""
    plantuml_html = ""
    puml_raw = data_flow.get('plantuml_code', "")

    if puml_raw:
        clean_puml = sanitize_plantuml(puml_raw)
        encoded_puml = encode_plantuml(clean_puml)
        img_url_svg = f"https://www.plantuml.com/plantuml/svg/{encoded_puml}"
        img_url_png = f"https://www.plantuml.com/plantuml/png/{encoded_puml}"

        try:
            response = requests.get(img_url_png, timeout=10)
            if response.status_code == 200:
                b64_image = base64.b64encode(response.content).decode('utf-8')
                embedded_img_html = f"data:image/png;base64,{b64_image}"
        except:
            embedded_img_html = img_url_png

        plantuml_html = f"""
        <div class="plantuml-scroll-container" style="page-break-before: always; margin-top: 40px; padding-top: 20px;">
            <h3 style="text-align:center; font-size:14pt; font-weight:bold; margin-bottom:15px; color:#1e3a8a;">DIAGRAM ALUR PROSES (SWIMLANE)</h3>
            <div style="background:#f8fafc; padding:20px; border-radius:10px; border:1px solid #e2e8f0; text-align:center; overflow-x:auto;">
                <img src="{img_url_svg}" alt="SOP Swimlane Flowchart" style="max-width:none; min-width:800px; display:block; margin:0 auto;" />
            </div>
        </div>
        """

    def get_primary_col_idx(pelaksana_str):
        if not pelaksana_str: return 0
        actors = re.split(r'[,/&]', pelaksana_str)
        for a in actors:
            a_clean = a.strip().lower()
            for idx, role in enumerate(pelaksana_list):
                if role.lower() in a_clean or a_clean in role.lower():
                    return idx
        return 0

    img_start = get_png_data_uri('start')
    img_process = get_png_data_uri('process')
    img_decision = get_png_data_uri('decision')

    steps_html = ""
    for idx, step in enumerate(steps):
        is_last = idx == (len(steps) - 1)
        col_idx = get_primary_col_idx(step.get('pelaksana', ''))
        next_col_idx = get_primary_col_idx(steps[idx+1].get('pelaksana', '')) if not is_last else -1
        kegiatan_txt = str(step.get('kegiatan', '')).lower()

        pelaksana_cols = ""
        for k in range(len(pelaksana_list)):
            if k == col_idx:
                img_src = img_process
                width = "55"
                if idx == 0 or is_last: img_src = img_start
                elif any(kw in kegiatan_txt for kw in ["jika", "?", "apakah", "memutuskan", "verifikasi"]):
                    img_src = img_decision
                    width = "45"

                arrow = ""
                if not is_last:
                    if next_col_idx > col_idx: arrow = "<div style='text-align:center;font-size:18px;font-weight:bold;margin-top:2px;'>↳</div>"
                    elif next_col_idx < col_idx: arrow = "<div style='text-align:center;font-size:18px;font-weight:bold;margin-top:2px;'>↲</div>"
                    else: arrow = "<div style='text-align:center;font-size:16px;font-weight:bold;margin-top:2px;'>↓</div>"

                pelaksana_cols += f"<td align='center' style='padding-top:10px;vertical-align:top;'><img src='{img_src}' width='{width}' style='display:block;margin:0 auto;'/>{arrow}</td>"
            else:
                pelaksana_cols += "<td></td>"

        steps_html += f"""
        <tr>
            <td align="center">{idx + 1}</td>
            <td style="padding:10px; text-align:justify;">{step.get('kegiatan', '-')}</td>
            {pelaksana_cols}
            <td style="padding:10px;">{step.get('kelengkapan', '-')}</td>
            <td align="center">{step.get('waktu', '-')}</td>
            <td style="padding:10px;">{step.get('output', '-')}</td>
            <td style="padding:10px;">{step.get('keterangan', '-')}</td>
        </tr>
        """

    headers_pelaksana = "".join([f'<th style="width:70px; padding:0;"><div style="writing-mode: vertical-rl; transform: rotate(180deg); height: 100px; padding: 5px; margin: 0 auto;">{p}</div></th>' for p in pelaksana_list])

    doc_html = f"""
    <div style="background:white; padding:1cm; font-family:Arial, sans-serif; width:100%; border:1px solid #d1d5db; box-shadow: 0 10px 15px -3px rgba(0,0,0,0.1);">
        <table style="border-collapse:collapse; width:100%; border:1px solid black; margin-bottom:20px;">
            <tr>
                <td rowspan="6" style="text-align:center; width:45%; padding:15px; border:1px solid black;">
                    <img src="{logo_b64_html}" width="80" height="80" style="width:80px; height:80px; object-fit:contain; display:block; margin:0 auto 10px auto;">
                    <div style="font-weight:bold; font-size:11pt;">RS {rs_name.upper()}</div>
                </td>
                <td style="width:20%; border:1px solid black; padding:5px;">Nomor SOP</td>
                <td style="width:35%; font-weight:bold; border:1px solid black; padding:5px;">{no_dok}</td>
            </tr>
            <tr><td style="border:1px solid black; padding:5px;">Tgl Pembuatan</td><td style="border:1px solid black; padding:5px;">-</td></tr>
            <tr><td style="border:1px solid black; padding:5px;">Tgl Revisi</td><td style="border:1px solid black; padding:5px;">-</td></tr>
            <tr><td style="border:1px solid black; padding:5px;">Tgl Pengesahan</td><td style="border:1px solid black; padding:5px;">-</td></tr>
            <tr>
                <td style="border:1px solid black; padding:5px;">Disahkan Oleh</td>
                <td style="text-align:center; padding:15px 5px; border:1px solid black;">
                    <b style="font-size: 11pt;">DIREKTUR UTAMA</b><br><br><br><br>
                    <u><b style="font-size: 11pt;">{dir_name}</b></u><br>
                    <span style="font-size: 10pt;">NIP. {dir_nip}</span>
                </td>
            </tr>
            <tr><td style="border:1px solid black; padding:5px;">Judul SOP</td><td style="font-weight:bold; border:1px solid black; padding:5px;">{title.upper()}</td></tr>
        </table>

        <table style="border-collapse:collapse; width:100%; border:1px solid black; margin-bottom:20px;">
            <tr>
                <td width="50%" style="background:#eeeeee; text-align:center; font-weight:bold; border:1px solid black; padding:8px;">DASAR HUKUM</td>
                <td style="background:#eeeeee; text-align:center; font-weight:bold; border:1px solid black; padding:8px;">KUALIFIKASI PELAKSANA</td>
            </tr>
            <tr>
                <td style="border:1px solid black; padding:8px;">{render_numbered_list(data_cover.get('dasar_hukum', []))}</td>
                <td style="border:1px solid black; padding:8px;">{render_numbered_list(data_cover.get('kualifikasi', []))}</td>
            </tr>
            <tr>
                <td style="background:#eeeeee; text-align:center; font-weight:bold; border:1px solid black; padding:8px;">KETERKAITAN</td>
                <td style="background:#eeeeee; text-align:center; font-weight:bold; border:1px solid black; padding:8px;">PERALATAN</td>
            </tr>
            <tr>
                <td style="border:1px solid black; padding:8px;">{render_numbered_list(data_cover.get('keterkaitan', []))}</td>
                <td style="border:1px solid black; padding:8px;">{render_numbered_list(data_cover.get('peralatan', []))}</td>
            </tr>
            <tr>
                <td style="background:#eeeeee; text-align:center; font-weight:bold; border:1px solid black; padding:8px;">PERINGATAN</td>
                <td style="background:#eeeeee; text-align:center; font-weight:bold; border:1px solid black; padding:8px;">PENCATATAN</td>
            </tr>
            <tr>
                <td style="border:1px solid black; padding:8px;">{data_cover.get('peringatan', '-')}</td>
                <td style="border:1px solid black; padding:8px;">{render_numbered_list(data_cover.get('pencatatan', []))}</td>
            </tr>
        </table>

        <div style="text-align:center; font-weight:bold; margin-top:20px; margin-bottom:10px;">RINCIAN PROSEDUR</div>
        <table style="border-collapse:collapse; width:100%; border:1px solid black;">
            <thead>
                <tr>
                    <th rowspan="2" width="5%" style="border:1px solid black; background:#fff;">No</th>
                    <th rowspan="2" width="20%" style="border:1px solid black; background:#fff;">Kegiatan</th>
                    <th colspan="{len(pelaksana_list)}" style="border:1px solid black; background:#fff;">Pelaksana</th>
                    <th colspan="3" style="border:1px solid black; background:#fff;">Mutu Baku</th>
                    <th rowspan="2" width="10%" style="border:1px solid black; background:#fff;">Ket</th>
                </tr>
                <tr>
                    {headers_pelaksana}
                    <th width="15%" style="border:1px solid black; background:#fff;">Kelengkapan</th>
                    <th width="10%" style="border:1px solid black; background:#fff;">Waktu</th>
                    <th width="10%" style="border:1px solid black; background:#fff;">Output</th>
                </tr>
            </thead>
            <tbody>
                {steps_html}
            </tbody>
        </table>

        {plantuml_html}
    </div>
    """

    word_html_string = doc_html.replace(
        f'src="{img_url_svg}"',
        f'src="{embedded_img_html}" height="500"'
    )

    word_wrapper = f"""
    <html xmlns:o='urn:schemas-microsoft-com:office:office' xmlns:w='urn:schemas-microsoft-com:office:word' xmlns='http://www.w3.org/TR/REC-html40'>
    <head><meta charset='utf-8'><title>{title}</title>
    <style>
        @page {{ size: 29.7cm 21cm; margin: 1.5cm; mso-page-orientation: landscape; }}
        body {{ font-family: Arial, sans-serif; font-size: 10pt; }}
        table {{ border-collapse: collapse; width: 100%; table-layout: auto; margin-bottom: 20px; }}
        td, th {{ border: 1px solid #000000; padding: 6px; vertical-align: middle; word-wrap: break-word; }}
        th {{ background-color: #f0f0f0; font-weight: bold; text-align: center; }}
        img {{ max-width: 100%; }}
    </style></head><body>{word_html_string}</body></html>
    """
    return doc_html, word_wrapper

# =====================================================================
# FUNGSI GENERATE DOCX UNTUK NOTULENSI
# =====================================================================
def generate_notulensi_docx(data):
    doc = Document()
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True
    
    cell_left = header_table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    run_logo = p_left.add_run("SMARTDOSE\n")
    run_logo.bold = True
    run_logo.font.size = Pt(28)
    run_logo.font.color.rgb = RGBColor(30, 58, 138) 
    
    run_sub = p_left.add_run("Enterprise AI Transcription\nHealthcare & Productivity")
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    cell_right = header_table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_org = p_right.add_run("SMARTDOSE ENTERPRISE\nDIVISI INOVASI DIGITAL\n")
    run_org.bold = True
    run_org.font.size = Pt(12)
    run_org.font.color.rgb = RGBColor(30, 58, 138)
    run_motto = p_right.add_run("Sistem Notulensi Cerdas & Presisi")
    run_motto.italic = True
    run_motto.font.size = Pt(10)

    doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() 

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NOTULEN RAPAT")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(30, 58, 138)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notulensi = data.get('notulensi_rapat', {})
    run_subtitle = p_subtitle.add_run(notulensi.get('agenda', 'Koordinasi dan Pembahasan Internal'))
    run_subtitle.bold = True
    run_subtitle.font.size = Pt(12)
    doc.add_paragraph() 

    table_meta = doc.add_table(rows=6, cols=2)
    table_meta.style = 'Table Grid'
    for row in table_meta.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    meta_data = [
        ("Hari / Tanggal", datetime.now().strftime("%A, %d %B %Y")),
        ("Waktu", datetime.now().strftime("%H:%M WIB")),
        ("Media", "SmartDose TranscribX (Offline/Live)"),
        ("Notulis", "AI Transcription System"),
        ("Peserta", ", ".join(notulensi.get('peserta', [])) if isinstance(notulensi.get('peserta'), list) else notulensi.get('peserta', '-')),
        ("Agenda", notulensi.get('agenda', '-'))
    ]

    for i, (key, val) in enumerate(meta_data):
        cell_k = table_meta.cell(i, 0)
        cell_k.text = key
        cell_k.paragraphs[0].runs[0].bold = True
        table_meta.cell(i, 1).text = str(val)
        
    doc.add_paragraph()

    def add_section_header(num, title):
        p = doc.add_paragraph()
        run = p.add_run(f"{num}. {title}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 58, 138)

    def add_bullet_points(items, is_numbered=False):
        style = 'List Number' if is_numbered else 'List Bullet'
        for item in items: doc.add_paragraph(str(item), style=style)

    add_section_header("1", "RINGKASAN EKSEKUTIF")
    for r in data.get('ringkasan_eksekutif', []): doc.add_paragraph(r)

    add_section_header("2", "POKOK PEMBAHASAN / JALANNYA DISKUSI")
    add_bullet_points(notulensi.get('jalannya_diskusi', []))

    add_section_header("3", "KEPUTUSAN RAPAT")
    add_bullet_points(notulensi.get('keputusan', []), is_numbered=True)

    add_section_header("4", "TINDAK LANJUT")
    tindak_lanjut = notulensi.get('rencana_tindak_lanjut', [])
    if tindak_lanjut:
        table_tl = doc.add_table(rows=1, cols=4)
        table_tl.style = 'Table Grid'
        hdr_cells = table_tl.rows[0].cells
        headers = ['Tugas', 'PIC', 'Deadline', 'Prioritas']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for tl in tindak_lanjut:
            row_cells = table_tl.add_row().cells
            row_cells[0].text = str(tl.get('tugas', '-'))
            row_cells[1].text = str(tl.get('pic', '-'))
            row_cells[2].text = str(tl.get('deadline', '-'))
            row_cells[3].text = str(tl.get('prioritas', '-'))
    else: doc.add_paragraph("- Tidak ada tindak lanjut khusus.", style='List Bullet')

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_footer = p_footer.add_run(f"\n\nJakarta, {datetime.now().strftime('%d %B %Y')}\n\n\n\n( Tim Notulen SmartDose )")
    
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# =====================================================================
# INISIALISASI SESSION STATE
# =====================================================================
if "logged_in" not in st.session_state: st.session_state["logged_in"] = False
if "offline_transcript" not in st.session_state: st.session_state["offline_transcript"] = ""
if "offline_summary" not in st.session_state: st.session_state["offline_summary"] = None
if "offline_segments" not in st.session_state: st.session_state["offline_segments"] = []
if "offline_audio_b64" not in st.session_state: st.session_state["offline_audio_b64"] = ""
if "offline_audio_mime" not in st.session_state: st.session_state["offline_audio_mime"] = ""
if "confirm_delete" not in st.session_state: st.session_state["confirm_delete"] = None
if "sop_preview_html" not in st.session_state: st.session_state["sop_preview_html"] = None
if "sop_word_html" not in st.session_state: st.session_state["sop_word_html"] = None
if "sop_title" not in st.session_state: st.session_state["sop_title"] = ""

if "captcha_n1" not in st.session_state:
    st.session_state.captcha_n1 = random.randint(1, 10)
    st.session_state.captcha_n2 = random.randint(1, 10)
if "last_reg_attempt" not in st.session_state:
    st.session_state.last_reg_attempt = datetime.min

# =====================================================================
# HALAMAN DEPAN (BERANDA, LOGIN, REGISTRASI)
# =====================================================================
if not st.session_state["logged_in"]:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Balsamiq+Sans:wght@700&display=swap');
    
    /* === PEMBASMI GHOST TOOLTIP CHART YANG NYANGKUT SAAT LOGOUT === */
    #vg-tooltip-element, .vg-tooltip { display: none !important; opacity: 0 !important; visibility: hidden !important; }
    .stApp, [data-testid="stAppViewContainer"], [data-testid="stHeader"] { background: transparent !important; }
    [data-testid="stSidebar"] { display: none; }
    
    /* === STYLING TAB HALAMAN DEPAN SUPER ELEGAN === */
    div[data-testid="stTabs"] { margin-bottom: 20px; }
    div[data-testid="stTabs"] button {
        color: #f1f5f9 !important;
        background-color: rgba(30, 41, 59, 0.6) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 12px !important;
        margin-right: 10px !important;
        padding: 8px 24px !important;
        font-weight: 600 !important;
        font-size: 15px !important;
        transition: all 0.3s ease !important;
    }
    div[data-testid="stTabs"] button:hover {
        background-color: rgba(56, 189, 248, 0.2) !important;
        border-color: rgba(56, 189, 248, 0.5) !important;
        transform: translateY(-2px);
    }
    div[data-testid="stTabs"] button[aria-selected="true"] {
        color: #ffffff !important;
        background-color: #3b82f6 !important;
        border: 1px solid #60a5fa !important;
        box-shadow: 0 0 15px rgba(59, 130, 246, 0.5) !important;
    }
    div[data-baseweb="tab-list"] { border-bottom: none !important; background-color: transparent !important; }

    div[data-testid="stForm"] {
        background: rgba(15, 23, 42, 0.4) !important;
        backdrop-filter: blur(16px) saturate(180%);
        -webkit-backdrop-filter: blur(16px) saturate(180%);
        padding: 40px 35px !important;
        border-radius: 24px !important;
        box-shadow: 0 0 30px rgba(56, 189, 248, 0.15), inset 0 0 20px rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(56, 189, 248, 0.3) !important;
        transition: all 0.3s ease-in-out;
    }
    div[data-testid="stForm"]:hover {
        box-shadow: 0 0 50px rgba(56, 189, 248, 0.25), inset 0 0 20px rgba(255, 255, 255, 0.05) !important;
        border: 1px solid rgba(56, 189, 248, 0.6) !important;
    }
    
    div[data-testid="stForm"] p, div[data-testid="stForm"] label { color: #e2e8f0 !important; font-weight: 600 !important; letter-spacing: 0.5px; }
    
    div[data-baseweb="input"] { background-color: rgba(30, 41, 59, 0.7) !important; border-radius: 12px !important; border: 1px solid rgba(100, 116, 139, 0.5) !important; }
    div[data-baseweb="input"]:focus-within { border-color: #38bdf8 !important; box-shadow: 0 0 15px rgba(56, 189, 248, 0.4) !important; }
    div[data-baseweb="input"] input { color: #ffffff !important; }
    div[data-baseweb="input"] input::placeholder { color: #64748b !important; }
    
    div[data-testid="stForm"] button[type="submit"] {
        background: linear-gradient(135deg, #0ea5e9 0%, #3b82f6 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        font-weight: 800 !important;
        letter-spacing: 1px;
        padding: 0.75rem 1rem !important;
        box-shadow: 0 10px 25px -5px rgba(14, 165, 233, 0.5) !important;
        transition: all 0.3s ease !important;
        margin-top: 15px;
    }
    div[data-testid="stForm"] button[type="submit"]:hover {
        transform: translateY(-3px);
        box-shadow: 0 15px 35px -5px rgba(14, 165, 233, 0.7) !important;
        background: linear-gradient(135deg, #38bdf8 0%, #2563eb 100%) !important;
    }
    
    .login-title {
        font-family: 'Balsamiq Sans', cursive; 
        color: #FFD166; font-weight: 700; font-size: 4.5rem;
        text-shadow: 3px 3px 0px #FF9F1C, 0px 5px 15px rgba(0, 0, 0, 0.6); margin: 0; line-height: 1;
    }
    
    @keyframes float-login { 0%, 100% { transform: translateY(0px) rotate(0deg); } 50% { transform: translateY(-12px) rotate(3deg); } }
    @keyframes signal-login { 0% { transform: scale(0.5); opacity: 0; } 50% { opacity: 1; } 100% { transform: scale(1.5); opacity: 0; } }
    @keyframes pulse-login { 0%, 100% { opacity: 1; } 50% { opacity: 0.3; } }
    
    .germic-login-wrapper { width: 140px; height: 140px; animation: float-login 4s ease-in-out infinite; margin-right: 20px; filter: drop-shadow(0px 6px 15px rgba(0,0,0,0.5)); }
    .signal-wave-login { transform-origin: 50px 12px; animation: signal-login 2s infinite; }
    .signal-wave-2-login { transform-origin: 50px 12px; animation-delay: 0.6s; animation: signal-login 2s infinite; }
    .animate-pulse-login { animation: pulse-login 2s cubic-bezier(0.4, 0, 0.6, 1) infinite; }
    </style>
    """, unsafe_allow_html=True)
    
    components.html("""
    <script>
        const parentWindow = window.parent; const parentDoc = parentWindow.document;
        if (parentWindow.nodeAnimFrame) cancelAnimationFrame(parentWindow.nodeAnimFrame);
        
        let canvas = parentDoc.getElementById('node-bg-canvas');
        if (!canvas) {
            canvas = parentDoc.createElement('canvas'); canvas.id = 'node-bg-canvas';
            Object.assign(canvas.style, { position: 'fixed', top: '0', left: '0', width: '100vw', height: '100vh', zIndex: '-1', background: 'radial-gradient(circle at center, #1e1b4b 0%, #020617 100%)' });
            parentDoc.body.prepend(canvas);
        }

        const ctx = canvas.getContext('2d');
        let w, h, points = [], edges = [], sparks = [], time = 0, maxLayer = 6; let coreX, coreY;

        function resize() { w = canvas.width = parentWindow.innerWidth; h = canvas.height = parentWindow.innerHeight; coreX = w / 2; coreY = h * 0.25; initNetwork(); }

        function initNetwork() {
            points = []; edges = []; sparks = []; points.push({ x: coreX, y: coreY, layer: 0 });
            for(let l = 1; l <= maxLayer; l++) {
                let radius = l * (Math.max(w, h) / 1.8) / maxLayer; let count = l * 12; 
                for(let i = 0; i < count; i++) {
                    let angle = (i / count) * Math.PI * 2 + (Math.random() * 0.4);
                    points.push({ x: coreX + Math.cos(angle) * (radius + (Math.random() * 60 - 30)), y: coreY + Math.sin(angle) * (radius + (Math.random() * 60 - 30)), layer: l });
                }
            }
            for(let i = 1; i < points.length; i++) {
                let p = points[i]; let targets = points.filter(t => t.layer === p.layer - 1);
                if (targets.length > 0) {
                    let closest = targets.reduce((prev, curr) => (Math.hypot(p.x - curr.x, p.y - curr.y) < Math.hypot(p.x - prev.x, p.y - prev.y)) ? curr : prev);
                    edges.push({ from: p, to: closest }); 
                }
            }
        }

        class Spark {
            constructor(startNode) { this.currentNode = startNode; this.findNextEdge(); this.progress = 0; this.speed = 0.01 + Math.random() * 0.015; }
            findNextEdge() { let nextEdges = edges.filter(e => e.from === this.currentNode); this.edge = nextEdges.length > 0 ? nextEdges[Math.floor(Math.random() * nextEdges.length)] : null; }
            update() {
                if (!this.edge) return true; 
                this.progress += this.speed;
                if (this.progress >= 1) { this.currentNode = this.edge.to; this.progress = 0; this.findNextEdge(); if (!this.edge) return true; }
                return false;
            }
            draw() {
                if (!this.edge) return;
                let x = this.edge.from.x + (this.edge.to.x - this.edge.from.x) * this.progress; let y = this.edge.from.y + (this.edge.to.y - this.edge.from.y) * this.progress;
                ctx.beginPath(); ctx.arc(x, y, 2.5, 0, Math.PI * 2); ctx.fillStyle = '#38bdf8'; ctx.shadowBlur = 15; ctx.shadowColor = '#0ea5e9'; ctx.fill(); ctx.shadowBlur = 0;
            }
        }

        parentWindow.addEventListener('resize', resize); resize();

        function animate() {
            ctx.clearRect(0, 0, w, h); time += 0.05; ctx.lineWidth = 1;
            edges.forEach(e => { ctx.beginPath(); ctx.moveTo(e.from.x, e.from.y); ctx.lineTo(e.to.x, e.to.y); ctx.strokeStyle = 'rgba(99, 102, 241, 0.15)'; ctx.stroke(); });

            if (Math.random() < 0.3) {
                let outerNodes = points.filter(p => p.layer === maxLayer || p.layer === maxLayer - 1);
                if(outerNodes.length > 0) sparks.push(new Spark(outerNodes[Math.floor(Math.random() * outerNodes.length)]));
            }

            for (let i = sparks.length - 1; i >= 0; i--) { let s = sparks[i]; if (s.update()) sparks.splice(i, 1); else s.draw(); }

            let pulse = Math.sin(time) * 8; 
            ctx.beginPath(); ctx.arc(coreX, coreY, 35 + pulse, 0, Math.PI * 2);
            let grad = ctx.createRadialGradient(coreX, coreY, 5, coreX, coreY, 45 + pulse);
            grad.addColorStop(0, '#fcd34d'); grad.addColorStop(0.4, '#f59e0b'); grad.addColorStop(1, 'rgba(239, 68, 68, 0)');
            ctx.fillStyle = grad; ctx.fill();
            
            ctx.beginPath(); ctx.arc(coreX, coreY, 12 + (pulse/3), 0, Math.PI * 2);
            ctx.fillStyle = '#fffbeb'; ctx.shadowBlur = 20; ctx.shadowColor = '#f59e0b'; ctx.fill(); ctx.shadowBlur = 0;
            parentWindow.nodeAnimFrame = requestAnimationFrame(animate);
        }
        animate();

        parentWindow.addEventListener('mousemove', (e) => {
            const face = parentDoc.getElementById('germic-login-face');
            if (face) {
                const limit = 8; const rect = face.getBoundingClientRect();
                const mouseX = e.clientX - (rect.left + rect.width / 2); const mouseY = e.clientY - (rect.top + rect.height / 2);
                face.style.transform = `translate(${Math.max(Math.min(mouseX / 30, limit), -limit)}px, ${Math.max(Math.min(mouseY / 30, limit), -limit)}px)`;
                face.style.transition = "transform 0.1s ease-out";
            }
        });
    </script>
    """, height=0, width=0)

    tab_home, tab_login, tab_register = st.tabs(["🏠 Beranda Platform", "🔐 Login Portal", "📝 Daftar Akun Baru"])

    with tab_home:
        st.markdown("""
        <div style="text-align: center; padding: 40px 20px; color: white; animation: fadeIn 1s ease-in-out;">
            <h1 style="font-size: 3.5rem; font-weight: 800; margin-bottom: 10px; background: -webkit-linear-gradient(#38bdf8, #3b82f6); -webkit-background-clip: text; -webkit-text-fill-color: transparent; text-shadow: 0px 4px 15px rgba(56,189,248,0.3);">
                TranscribX Enterprise
            </h1>
            <p style="font-size: 1.2rem; color: #cbd5e1; max-width: 700px; margin: 0 auto 30px auto; line-height: 1.6;">
                Pusat Notulensi & Edukasi Cerdas AI. Mengubah suara rapat Anda menjadi teks, ringkasan eksekutif, daftar penugasan, dan SOP Akreditasi secara instan.
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("""
        <style>
        .glass-hover-card { background: rgba(30, 41, 59, 0.6); backdrop-filter: blur(12px); border: 1px solid rgba(56, 189, 248, 0.3); border-radius: 16px; padding: 25px; color: white; height: 100%; transition: transform 0.3s ease; }
        .glass-hover-card:hover { transform: translateY(-5px); }
        </style>
        """, unsafe_allow_html=True)
        
        st.markdown("<h3 style='text-align: center; color: #38bdf8; margin-top: 20px; font-weight: bold;'>✨ Kenapa Memilih TranscribX?</h3>", unsafe_allow_html=True)
        st.write("")
        col_f1, col_f2, col_f3 = st.columns(3)
        glass_card = "<div class='glass-hover-card'><div style='font-size: 35px; margin-bottom: 15px;'>{icon}</div><h4 style='color: #e0f2fe; margin-top: 0; font-size: 18px;'>{title}</h4><p style='color: #94a3b8; font-size: 14px; line-height: 1.6;'>{desc}</p></div>"
        with col_f1: st.markdown(glass_card.format(icon="🎙️", title="Transkripsi Live & Cepat", desc="Tangkap kata dari rapat Zoom/YouTube dengan teknologi Speech-to-Text mutakhir secara real-time."), unsafe_allow_html=True)
        with col_f2: st.markdown(glass_card.format(icon="🤖", title="AI Smart Summary", desc="AI otomatis membuat ringkasan eksekutif, daftar tugas (Action Items), dan Peta Konsep (Mindmap)."), unsafe_allow_html=True)
        with col_f3: st.markdown(glass_card.format(icon="🏥", title="SOP Generator (KARS/JCI)", desc="Ubah keputusan rapat menjadi dokumen SOP resmi Akreditasi RS lengkap dengan Flowchart Swimlane otomatis."), unsafe_allow_html=True)

        st.markdown("""
        <style>
        .pricing-card { background-color: rgba(30, 41, 59, 0.4); padding: 30px 20px; border-radius: 15px; border: 1px solid rgba(255,255,255,0.1); height: 100%; min-height: 420px; text-align: center; position: relative; display: flex; flex-direction: column; transition: transform 0.2s ease; }
        .pricing-card:hover { transform: scale(1.02); }
        </style>
        """, unsafe_allow_html=True)
        
        st.markdown("<h3 style='text-align: center; color: #38bdf8; margin-top: 20px; font-weight: bold;'>✨ Paket Langganan Tersedia</h3>", unsafe_allow_html=True)
        st.write("")
        col_p1, col_p2, col_p3 = st.columns(3)
        
        with col_p1:
            st.markdown("""<div class="pricing-card">
                <h3 style="color:#94a3b8; margin-top:0; font-size:20px; margin-bottom:10px; text-transform:uppercase; letter-spacing:1px;">Paket BASIC</h3>
                <h2 style="color:#e0f2fe; font-size:36px; margin:10px 0;">Rp 35.000</h2><p style="color:#94a3b8; font-size:13px; margin:0 0 20px 0;">/ 30 hari</p><hr style="border-color:rgba(255,255,255,0.1); margin:20px 0;">
                <ul style="list-style:none; padding:0; margin:0; font-size:14px; color:#cbd5e1; text-align:left; line-height:2.2; flex-grow:1;">
                    <li>✅ <b>Unlimited</b> Live Transcribe</li><li>✅ <b>8x</b> Premium AI Summary</li><li>❌ <del>AI SOP Generator</del></li><li>✅ <b>1x</b> Upload Audio (Max 45mnt)</li><li>⏳ <b>30 Hari</b> Masa Aktif</li>
                </ul></div>""", unsafe_allow_html=True)
        
        with col_p2:
            st.markdown("""<div class="pricing-card">
                <div style="position:absolute; top:-15px; left:50%; transform:translateX(-50%); background:linear-gradient(135deg, #ef4444 0%, #dc2626 100%); color:white; padding:6px 18px; border-radius:20px; font-size:12px; font-weight:bold; white-space:nowrap; box-shadow:0 4px 15px rgba(239,68,68,0.4); letter-spacing:0.5px;">🔥 BEST SELLER</div>
                <h3 style="color:#38bdf8; margin-top:0; font-size:20px; margin-bottom:10px; text-transform:uppercase; letter-spacing:1px;">EXECUTIVE</h3>
                <h2 style="color:white; font-size:32px; margin:10px 0;">Rp 69.000</h2><p style="color:#94a3b8; font-size:13px; margin:0 0 20px 0;">/ 30 hari</p><hr style="border-color:rgba(255,255,255,0.1); margin:20px 0;">
                <ul style="list-style:none; padding:0; margin:0; font-size:14px; color:#cbd5e1; text-align:left; line-height:2.2; flex-grow:1;">
                    <li>✅ <b>Unlimited</b> Live Transcribe</li><li>✅ <b>25x</b> Premium AI Summary</li><li>✅ <b>AI SOP Generator</b></li><li>✅ <b>5x</b> Upload Audio (Max 45mnt)</li><li>🌟 Bisa <b>Top-Up Kuota</b> Kapan Saja</li>
                </ul></div>""", unsafe_allow_html=True)
        
        with col_p3:
            st.markdown("""<div class="pricing-card">
                <h3 style="color:#fb7185; margin-top:0; font-size:20px; margin-bottom:10px; text-transform:uppercase; letter-spacing:1px;">Paket MASTER</h3>
                <p style="color:#fca5a5; font-size:12px; margin:0 0 8px 0; font-weight:bold;">VIP / ENTERPRISE</p>
                <h2 style="color:#e0f2fe; font-size:36px; margin:10px 0;">Rp 149.000</h2><p style="color:#94a3b8; font-size:13px; margin:0 0 20px 0;">/ 30 hari</p><hr style="border-color:rgba(255,255,255,0.1); margin:20px 0;">
                <ul style="list-style:none; padding:0; margin:0; font-size:14px; color:#cbd5e1; text-align:left; line-height:2.2; flex-grow:1;">
                    <li>✅ <b>Unlimited</b> Live Transcribe</li><li>✅ <b>60x</b> Premium AI Summary</li><li>✅ <b>AI SOP Generator</b></li><li>✅ <b>15x</b> Upload Audio (Max 45mnt)</li><li>🌟 <b>Prioritas Support</b> via WA 24/7</li>
                </ul></div>""", unsafe_allow_html=True)
         
        st.markdown("""<div style="text-align:center; margin-top:30px; padding:20px; background:rgba(56, 189, 248, 0.1); border-radius:12px; border:1px solid rgba(56, 189, 248, 0.3);">
            <h4 style="color:#38bdf8; margin:0;">Ingin aktivasi atau perpanjang paket?</h4>
            <p style="color:#64748b; margin:10px 0 0 0; font-size:14px;">Silakan <b>Daftar / Login</b> terlebih dahulu melalui tab di atas untuk melakukan pembelian paket secara otomatis. 🚀</p>
        </div>""", unsafe_allow_html=True)

    with tab_login:
        col1, col2, col3 = st.columns([1, 1.5, 1])
        with col2:
            st.write("")
            st.write("")
            st.write("")
            title_html = """
            <div style='display: flex; justify-content: center; align-items: center; margin-bottom: 10px;'>
                <div class="germic-login-wrapper">
                    <svg width="100%" height="100%" viewBox="0 0 100 100" fill="none">
                        <circle class="signal-wave-login" cx="50" cy="12" r="8" stroke="#fb7185" stroke-width="1" />
                        <circle class="signal-wave-login signal-wave-2-login" cx="50" cy="12" r="8" stroke="#fb7185" stroke-width="1" />
                        <circle cx="50" cy="12" r="8" stroke="#3b82f6" stroke-opacity="0.2" />
                        <path d="M50 25V15M50 15L45 10M50 15L55 10" stroke="#3b82f6" stroke-width="2" stroke-linecap="round"/>
                        <circle cx="50" cy="12" r="3" fill="#fb7185" class="animate-pulse-login"/>
                        <rect x="5" y="45" width="8" height="20" rx="4" fill="#1e293b"/>
                        <rect x="87" y="45" width="8" height="20" rx="4" fill="#1e293b"/>
                        <rect x="15" y="25" width="70" height="65" rx="18" fill="#f1f5f9" stroke="#cbd5e1" stroke-width="2"/>
                        <rect x="22" y="35" width="56" height="40" rx="12" fill="#1e1b4b"/>
                        <g id="germic-login-face"><rect x="33" y="45" width="12" height="15" rx="3" fill="#38bdf8"/><rect x="55" y="45" width="12" height="15" rx="3" fill="#38bdf8"/><rect x="42" y="65" width="16" height="3" rx="1.5" fill="#818cf8"/></g>
                    </svg>
                </div>
                <h1 class='login-title'>TranscribX</h1>
            </div>
            """
            st.markdown(title_html, unsafe_allow_html=True)
            
            with st.form("login_form"):
                st.markdown("""<h3 style='text-align: center; color: #e0f2fe; text-shadow: 0 0 10px rgba(56,189,248,0.5); margin-bottom: 5px; letter-spacing: 1px;'>Secure Login</h3><p style='text-align: center; color: #94a3b8; font-size: 0.9rem; margin-bottom: 25px;'>Portal Notulensi AI Enterprise. Masuk untuk melanjutkan.</p>""", unsafe_allow_html=True)
                email_login = st.text_input("Email Address", placeholder="Ketik email Anda di sini...")
                pass_login = st.text_input("Password", type="password", placeholder="••••••••")
                st.write("")
                btn_login = st.form_submit_button("🚀 Masuk ke Sistem", use_container_width=True, type="primary")
                
                if btn_login:
                    if email_login and pass_login:
                        with st.spinner("Memverifikasi kredensial..."):
                            user_data = login_firebase(email_login, pass_login)
                            if "idToken" in user_data:
                                uid = user_data["localId"]
                                record_login(uid, email_login)
                                if email_login in ADMIN_EMAILS_CONFIG:
                                    st.session_state.update({"logged_in": True, "user_email": email_login, "user_uid": uid, "user_paket": "ADMIN", "user_kuota_ai": 999999, "user_kuota_upload": 999999, "sisa_hari": 999999})
                                    st.success("✅ Selamat datang, Admin! Akses Unlimited diaktifkan.")
                                    st.rerun()
                                user_db_info = check_subscription(uid)
                                if user_db_info.get("status_subscription", "non-aktif") == "aktif":
                                    cek_reset_kuota_bulanan(uid, user_db_info)
                                    st.session_state.update({"logged_in": True, "user_email": email_login, "user_uid": uid, "user_paket": user_db_info.get("paket", "BASIC"), "user_kuota_ai": user_db_info.get("kuota_ai", 0), "user_kuota_upload": user_db_info.get("kuota_upload", 0), "sisa_hari": hitung_sisa_hari(user_db_info.get("tanggal_berakhir")), "tanggal_berakhir": user_db_info.get("tanggal_berakhir", "")})
                                    st.success("Login berhasil! Memuat sistem...")
                                    st.rerun()
                                else:
                                    st.session_state.update({"logged_in": True, "user_email": email_login, "user_uid": uid, "user_paket": "NON-AKTIF", "user_kuota_ai": 0, "user_kuota_upload": 0, "sisa_hari": 0})
                                    st.warning("⚠️ Akun Anda belum memiliki paket aktif. Silakan beli paket di menu Info Paket.")
                                    st.rerun()
                            else:
                                st.error(f"⚠️ {user_data.get('error', {}).get('message', 'Login gagal')}")
                    else: st.warning("Silakan masukkan email dan password.")

    with tab_register:
        col_r1, col_r2, col_r3 = st.columns([1, 1.5, 1])
        with col_r2:
            st.write("")
            st.write("")
            st.write("")
            st.markdown("""<div style='text-align: center; margin-bottom: 20px;'><h1 class='login-title' style='font-size: 3rem;'>Daftar Akun</h1><p style='color: #94a3b8; font-size: 0.9rem;'>Buat akun baru untuk mulai menggunakan TranscribX.</p></div>""", unsafe_allow_html=True)
            with st.form("register_form_user"):
                st.markdown("<h3 style='text-align: center; color: #e0f2fe; margin-bottom: 15px; letter-spacing: 1px;'>Form Pendaftaran</h3>", unsafe_allow_html=True)
                email_reg_user = st.text_input("Email Address", placeholder="Masukkan email aktif Anda...")
                pass_reg_user = st.text_input("Password", type="password", placeholder="Minimal 6 karakter")
                pass_confirm = st.text_input("Konfirmasi Password", type="password", placeholder="Ulangi password")
                st.markdown(f"<p style='color:#e2e8f0; font-size:14px; margin-top:10px; margin-bottom:5px;'>🛡️ Verifikasi Keamanan: <b>Berapa hasil dari {st.session_state.captcha_n1} + {st.session_state.captcha_n2} ?</b></p>", unsafe_allow_html=True)
                captcha_answer = st.text_input("Jawaban Captcha", placeholder="Ketik angka jawaban...", label_visibility="collapsed")
                st.write("")
                btn_reg_user = st.form_submit_button("📝 Buat Akun Sekarang", use_container_width=True, type="primary")
                
                if btn_reg_user:
                    now = datetime.now()
                    expected_ans = str(st.session_state.captcha_n1 + st.session_state.captcha_n2)
                    if (now - st.session_state.last_reg_attempt).total_seconds() < 10:
                        st.error("⏳ Terlalu banyak permintaan! Sistem mendeteksi aktivitas mencurigakan. Silakan tunggu 10 detik.")
                    elif captcha_answer.strip() != expected_ans:
                        st.session_state.last_reg_attempt = now 
                        st.error("❌ Jawaban Captcha salah! Silakan hitung dengan benar.")
                        st.session_state.captcha_n1 = random.randint(1, 10)
                        st.session_state.captcha_n2 = random.randint(1, 10)
                    elif not email_reg_user or len(pass_reg_user) < 6:
                        st.session_state.last_reg_attempt = now
                        st.warning("⚠️ Masukkan email yang valid dan password minimal 6 karakter.")
                    elif pass_reg_user != pass_confirm:
                        st.session_state.last_reg_attempt = now
                        st.warning("⚠️ Password dan Konfirmasi Password tidak cocok!")
                    else:
                        st.session_state.last_reg_attempt = now
                        with st.spinner("Membuat akun Anda..."):
                            new_user = register_firebase(email_reg_user, pass_reg_user)
                            if "idToken" in new_user:
                                uid = new_user["localId"]
                                sekarang = datetime.now()
                                db.collection("users").document(uid).set({
                                    "email": email_reg_user, "status_subscription": "non-aktif", "paket": "NON-AKTIF",
                                    "kuota_ai": 0, "kuota_upload": 0, "tanggal_mulai": sekarang.isoformat(),
                                    "tanggal_berakhir": sekarang.isoformat(), "reset_kuota_terakhir": sekarang.isoformat(),
                                    "last_login": "Belum pernah login", "login_count": 0
                                })
                                st.success("✅ Pendaftaran berhasil! Silakan klik tab '🔐 Login Portal' di atas untuk masuk dan memilih paket.")
                                st.session_state.captcha_n1 = random.randint(1, 10)
                                st.session_state.captcha_n2 = random.randint(1, 10)
                            else: st.error(f"⚠️ Gagal memproses: Email mungkin sudah terdaftar.")

# =====================================================================
# APLIKASI UTAMA (SETELAH LOGIN)
# =====================================================================
else:
# --- TAMBAHKAN BLOK KODE INI ---
    try:
        if st.session_state.get("user_uid"):
            waktu_wib = (datetime.utcnow() + timedelta(hours=7)).isoformat()
            db.collection("users").document(st.session_state["user_uid"]).update({
                "last_active": waktu_wib,
                "is_online": True
            })
    except: pass
    st.markdown("""
    <style>
    #node-bg-canvas { display: none !important; opacity: 0 !important; visibility: hidden !important; }
    .stApp { background: #ffffff !important; animation: none !important; }
    [data-testid="stSidebar"] { display: flex !important; }
    </style>
    """, unsafe_allow_html=True)
    components.html("<script>try { const canvas = window.parent.document.getElementById('node-bg-canvas'); if (canvas) { canvas.remove(); } } catch(e) {}</script>", height=0)
    
    with st.sidebar:
        st.markdown("<h3 style='text-align: center; color: #475569;'>🤖 AI Assistant</h3>", unsafe_allow_html=True)
        germic_html = """
        <style>
            @keyframes float { 0%, 100% { transform: translateY(0px) rotate(0deg); } 50% { transform: translateY(-15px) rotate(2deg); } }
            @keyframes signal { 0% { transform: scale(0.5); opacity: 0; } 50% { opacity: 1; } 100% { transform: scale(1.5); opacity: 0; } }
            body { margin: 0; padding: 0; display: flex; justify-content: center; align-items: center; height: 250px; background-color: transparent; overflow: hidden; }
            .germic-container { width: 180px; height: 180px; animation: float 4s ease-in-out infinite; position: relative; cursor: pointer; }
            .signal-wave { transform-origin: 50px 12px; animation: signal 2s infinite; }
            .signal-wave-2 { transform-origin: 50px 12px; animation-delay: 0.6s; animation: signal 2s infinite; }
            .animate-pulse { animation: pulse 2s cubic-bezier(0.4, 0, 0.6, 1) infinite; }
            @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.3; } }
        </style>
        <div id="germic-wrapper" class="germic-container">
            <svg width="100%" height="100%" viewBox="0 0 100 100" fill="none">
                <circle class="signal-wave" cx="50" cy="12" r="8" stroke="#fb7185" stroke-width="1" />
                <circle class="signal-wave signal-wave-2" cx="50" cy="12" r="8" stroke="#fb7185" stroke-width="1" />
                <circle cx="50" cy="12" r="8" stroke="#3b82f6" stroke-opacity="0.2" />
                <path d="M50 25V15M50 15L45 10M50 15L55 10" stroke="#3b82f6" stroke-width="2" stroke-linecap="round"/>
                <circle cx="50" cy="12" r="3" fill="#fb7185" class="animate-pulse"/>
                <rect x="5" y="45" width="8" height="20" rx="4" fill="#1e293b"/>
                <rect x="87" y="45" width="8" height="20" rx="4" fill="#1e293b"/>
                <rect x="15" y="25" width="70" height="65" rx="18" fill="#f1f5f9" stroke="#cbd5e1" stroke-width="2"/>
                <rect x="22" y="35" width="56" height="40" rx="12" fill="#1e1b4b"/>
                <g id="germic-face"><rect id="eye-l" x="33" y="45" width="12" height="15" rx="3" fill="#38bdf8"/><rect id="eye-r" x="55" y="45" width="12" height="15" rx="3" fill="#38bdf8"/><rect id="mouth" x="42" y="65" width="16" height="3" rx="1.5" fill="#818cf8"/></g>
            </svg>
        </div>
        <script>
            const face = document.getElementById('germic-face');
            function trackMouse(clientX, clientY, screenWidth) {
                if (!face) return;
                const robotX = screenWidth * 0.2; const robotY = 125;
                const mouseX = clientX - robotX; const mouseY = clientY - robotY;
                const limit = 8;
                face.style.transform = `translate(${Math.max(Math.min(mouseX/50, limit), -limit)}px, ${Math.max(Math.min(mouseY/50, limit), -limit)}px)`;
                face.style.transition = "transform 0.1s ease-out";
            }
            document.addEventListener('mousemove', (e) => trackMouse(e.clientX, e.clientY, window.innerWidth));
            try { window.parent.document.addEventListener('mousemove', (e) => trackMouse(e.clientX, e.clientY, window.parent.innerWidth)); } catch (err) { }
        </script>
        """
        components.html(germic_html, height=250)
        st.markdown("<p style='text-align: center; color: #94a3b8; font-size: 13px; font-weight: bold; margin-top:-20px;'>Sistem Online</p>", unsafe_allow_html=True)
        
        profile_placeholder = st.empty()
        def render_sidebar_profile():
            if st.session_state.get("logged_in"):
                user_email = st.session_state.get('user_email')
                if is_admin():
                    card_class = "profile-card admin"
                    paket_str = "👑 ADMIN (Unlimited)"
                    ai_str = "♾️ Unlimited"
                    up_str = "♾️ Unlimited"
                    hari_str = "Status: Permanen"
                else:
                    try:
                        uid = st.session_state.get("user_uid")
                        if uid:
                            fresh_data = db.collection("users").document(uid).get().to_dict()
                            if fresh_data:
                                st.session_state["user_kuota_ai"] = fresh_data.get("kuota_ai", 0)
                                st.session_state["user_kuota_upload"] = fresh_data.get("kuota_upload", 0)
                                st.session_state["user_paket"] = fresh_data.get("paket", "NON-AKTIF")
                                st.session_state["sisa_hari"] = hitung_sisa_hari(fresh_data.get("tanggal_berakhir"))
                    except: pass
                    
                    paket = st.session_state.get('user_paket', 'NON-AKTIF')
                    sisa_hari = st.session_state.get('sisa_hari', 0)
                    if paket == 'NON-AKTIF':
                        card_class = "profile-card user-inactive"
                        paket_str = "NON-AKTIF"
                        ai_str, up_str = "0x", "0x"
                        hari_str = "Masa Aktif: Habis"
                    else:
                        paket_str = paket 
                        ai_str = f"{st.session_state.get('user_kuota_ai', 0)}x"
                        up_str = f"{st.session_state.get('user_kuota_upload', 0)}x"
                        hari_str = f"⏳ Sisa: {sisa_hari} hari"
                        card_class = "profile-card user-warning" if sisa_hari <= 3 else "profile-card user-active"

                profile_html = f"""<div class="{card_class}"><div style="font-size: 11px; text-transform: uppercase; letter-spacing: 1px; opacity: 0.8; margin-bottom: 5px;">Akses Profil</div><div style="font-weight: 800; font-size: 14px; margin-bottom: 15px; word-break: break-all;">{user_email}</div><div style="background: rgba(255,255,255,0.15); padding: 10px; border-radius: 8px; margin-bottom: 10px;"><div style="font-size: 12px; margin-bottom: 5px;">🏷️ <b>Paket:</b> {paket_str}</div><div style="font-size: 12px; margin-bottom: 5px;">🎙️ <b>Transcribe:</b> ∞ Unlimited</div><div style="font-size: 12px; margin-bottom: 5px;">✨ <b>Sisa AI:</b> {ai_str}</div><div style="font-size: 12px;">📁 <b>Sisa Audio:</b> {up_str}</div></div><div style="font-size: 12px; font-weight: bold; text-align: center; margin-top: 10px;">{hari_str}</div></div>"""
                with profile_placeholder.container():
                    st.markdown(profile_html, unsafe_allow_html=True)
                    if not is_admin() and st.session_state.get('sisa_hari', 0) <= 3 and st.session_state.get('sisa_hari', 0) > 0: st.warning("⚠️ Masa aktif hampir habis! Segera perpanjang.")
        
        render_sidebar_profile()

        clock_html = """
        <style>@import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@700&display=swap'); .digital-clock { background: linear-gradient(145deg, #0f172a, #1e293b); border: 1px solid rgba(56, 189, 248, 0.4); box-shadow: 0 4px 10px rgba(0,0,0,0.3), inset 0 0 15px rgba(56, 189, 248, 0.1); border-radius: 12px; padding: 15px 10px; text-align: center; margin-top: 5px; } .time-display { font-family: 'Orbitron', sans-serif; color: #38bdf8; font-size: 24px; font-weight: bold; letter-spacing: 2px; text-shadow: 0 0 10px rgba(56, 189, 248, 0.6); display: flex; justify-content: center; align-items: center; } .date-display { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; color: #94a3b8; font-size: 11px; text-transform: uppercase; letter-spacing: 1px; margin-top: 6px; font-weight: 700; } .blink { animation: blinker 1s linear infinite; margin: 0 2px; color: #7dd3fc; } @keyframes blinker { 50% { opacity: 0.2; } }</style>
        <div class="digital-clock"><div class="time-display"><span id="hour">00</span><span class="blink">:</span><span id="minute">00</span><span class="blink">:</span><span id="second">00</span></div><div class="date-display" id="date">SENIN, 01 JAN 2026 • WIB</div></div>
        <script>
            function updateClock() {
                const now = new Date(); const options = { timeZone: 'Asia/Jakarta' }; const wibDate = new Date(now.toLocaleString('en-US', options));
                document.getElementById('hour').innerText = String(wibDate.getHours()).padStart(2, '0'); document.getElementById('minute').innerText = String(wibDate.getMinutes()).padStart(2, '0'); document.getElementById('second').innerText = String(wibDate.getSeconds()).padStart(2, '0');
                const hari = ['Minggu', 'Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu']; const bulan = ['Jan', 'Feb', 'Mar', 'Apr', 'Mei', 'Jun', 'Jul', 'Ags', 'Sep', 'Okt', 'Nov', 'Des'];
                document.getElementById('date').innerText = hari[wibDate.getDay()] + ', ' + String(wibDate.getDate()).padStart(2, '0') + ' ' + bulan[wibDate.getMonth()] + ' ' + wibDate.getFullYear() + ' • WIB';
            }
            setInterval(updateClock, 1000); updateClock();
        </script>
        """
        components.html(clock_html, height=100)
        
        if is_admin():
            if st.button("🔄 Refresh Klien Online", use_container_width=True): pass 
            try:
                # GANTI MENJADI KODE BARU INI:
                batas_waktu_online = (datetime.utcnow() + timedelta(hours=7) - timedelta(minutes=5)).isoformat()
                active_users_ref = db.collection("users").where("last_active", ">=", batas_waktu_online).stream()
                online_list_html, count_online = "", 0
                for doc in active_users_ref:
                    u_data = doc.to_dict()
                    if u_data.get("is_online", True) == False: continue 
                    email_klien = u_data.get("email", "-")
                    display_email = email_klien if len(email_klien) <= 22 else email_klien[:19] + "..."
                    online_list_html += f"<li style='margin-bottom: 8px; border-bottom: 1px dashed #e2e8f0; padding-bottom: 4px; display: flex; align-items: center; gap: 6px;'><span style='color: #10b981; font-size: 14px;'>●</span> {display_email}</li>"
                    count_online += 1
                if count_online == 0: online_list_html = "<li style='color: #94a3b8; font-style: italic; text-align: center; padding-top: 10px;'>Belum ada klien online</li>"
                admin_panel_html = f"""<div style="background: #ffffff; border: 1px solid #e2e8f0; border-radius: 12px; padding: 15px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05); margin-bottom: 20px;"><div style="font-size: 11px; font-weight: 800; color: #1e293b; margin-bottom: 15px; display: flex; align-items: center; gap: 8px; border-bottom: 2px solid #f1f5f9; padding-bottom: 8px;"><div style="position: relative; width: 10px; height: 10px;"><div style="position: absolute; top: -1px; left: -1px; width: 12px; height: 12px; background: #10b981; border-radius: 50%; animation: ping-dot 1.5s cubic-bezier(0, 0, 0.2, 1) infinite;"></div><div style="position: absolute; top: 1px; left: 1px; width: 8px; height: 8px; background: #059669; border-radius: 50%;"></div></div>KLIEN ONLINE ({count_online})</div><ul style="list-style: none; padding: 0; margin: 0; font-size: 12px; color: #475569; max-height: 180px; overflow-y: auto;">{online_list_html}</ul></div><style>@keyframes ping-dot {{ 75%, 100% {{ transform: scale(2.5); opacity: 0; }} }} ul::-webkit-scrollbar {{ width: 4px; }} ul::-webkit-scrollbar-track {{ background: #f1f5f9; border-radius: 4px; }} ul::-webkit-scrollbar-thumb {{ background: #cbd5e1; border-radius: 4px; }}</style>"""
                st.markdown(admin_panel_html, unsafe_allow_html=True)
            except: pass

    colA, colB = st.columns([8, 1])
    with colB:
        if st.button("🚪 Logout", use_container_width=True):
            try:
                uid = st.session_state.get("user_uid")
                if uid: db.collection("users").document(uid).update({"is_online": False})
            except: pass
            st.session_state.clear()
            st.session_state["logged_in"] = False
            st.rerun()

    st.title("🎙️ TranscribX: SmartDose Enterprise Transcription")

    if is_admin():
        tabs = st.tabs(["👑 Admin Panel", "🔴 Live Capture", "📁 Upload Rekaman", "📝 SOP Generator", "💳 Info Paket"])
        tab_admin, tab1, tab2, tab_sop, tab_paket = tabs[0], tabs[1], tabs[2], tabs[3], tabs[4]
        
        with tab_admin:
            st.markdown("### 👑 Dashboard Admin: Enterprise Control Center")
            with st.spinner("Memuat metrik & menyegarkan data klien..."):
                users_ref = db.collection("users").stream()
                users_list = []
                for doc in users_ref:
                    user_info = doc.to_dict()
                    email_user = user_info.get("email", "-")
                    cek_dan_update_status_kadaluarsa(doc.id, user_info)
                    user_info = db.collection("users").document(doc.id).get().to_dict()
                    if email_user in ADMIN_EMAILS_CONFIG:
                        status_user, paket_user, sisa_ai, sisa_up, sisa_hari_display = "admin", "ADMIN", "∞", "∞", "∞"
                        last_login, login_count = user_info.get("last_login", "Tidak diketahui"), user_info.get("login_count", 0)
                    else:
                        status_user = user_info.get("status_subscription", "non-aktif")
                        paket_user = user_info.get("paket", "BASIC" if status_user == "aktif" else "-")
                        sisa_ai, sisa_up = user_info.get("kuota_ai", 0), user_info.get("kuota_upload", 0)
                        last_login, login_count = user_info.get("last_login", "Belum pernah login"), user_info.get("login_count", 0)
                        tanggal_berakhir = user_info.get("tanggal_berakhir")
                        if status_user == "aktif" and tanggal_berakhir:
                            sisa_hari = hitung_sisa_hari(tanggal_berakhir)
                            sisa_hari_display = "Kadaluarsa" if sisa_hari <= 0 else (f"⚠️ {sisa_hari} hari" if sisa_hari <= 3 else f"{sisa_hari} hari")
                            if sisa_hari <= 0: status_user = "non-aktif ⚠️"
                        else: sisa_hari_display = "Kadaluarsa"
                    
                    last_login_display = last_login[:19].replace("T", " ") if isinstance(last_login, str) and len(last_login) > 19 and last_login not in ["Belum pernah login", "Tidak diketahui"] else last_login
                    users_list.append({"Email": email_user, "Status": status_user.split()[0], "Paket": paket_user, "Sisa AI": sisa_ai, "Sisa Upload": sisa_up, "Sisa Hari": sisa_hari_display, "Last Login": last_login_display, "Login Count": login_count, "UID": doc.id, "UID_Short": doc.id[:8] + "..."})

            total_users, total_aktif, total_nonaktif, total_admin = len(users_list), sum(1 for u in users_list if u['Status'] == 'aktif'), sum(1 for u in users_list if 'non-aktif' in u['Status']), sum(1 for u in users_list if u['Status'] == 'admin')
            active_now = get_active_users_count()

            col_m1, col_m2, col_m3, col_m4 = st.columns(4)
            def build_metric_card(title, value, subtext, icon, type_class): return f"""<div class="metric-card metric-{type_class}"><div style="display: flex; justify-content: space-between; align-items: flex-start;"><div><p style="margin: 0; color: #64748b; font-size: 13px; font-weight: 600; text-transform: uppercase;">{title}</p><h2 style="margin: 5px 0 0 0; color: #0f172a; font-size: 28px; line-height: 1;">{value}</h2><p style="margin: 8px 0 0 0; color: #10b981; font-size: 11px; font-weight: 600; background: #d1fae5; display: inline-block; padding: 2px 8px; border-radius: 10px;">{subtext}</p></div><div style="font-size: 28px; opacity: 0.9;">{icon}</div></div></div>"""
            with col_m1: st.markdown(build_metric_card("Total Users", total_users, f"↑ {active_now} online", "👥", "total"), unsafe_allow_html=True)
            with col_m2: st.markdown(build_metric_card("Klien Aktif", total_aktif, "Paket Berjalan", "✅", "aktif"), unsafe_allow_html=True)
            with col_m3: st.markdown(build_metric_card("Non-Aktif", total_nonaktif, "Perlu Follow-up", "❌", "nonaktif"), unsafe_allow_html=True)
            with col_m4: st.markdown(build_metric_card("Admin", total_admin, "Superuser", "👑", "admin"), unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            with st.expander("➕ Registrasi Akun Klien Baru (Manual)", expanded=False):
                with st.form("admin_register_form", clear_on_submit=True):
                    col_reg1, col_reg2, col_reg3 = st.columns(3)
                    with col_reg1: email_reg = st.text_input("Email Klien Baru", placeholder="email@klien.com")
                    with col_reg2: pass_reg = st.text_input("Password Klien", type="password", help="Minimal 6 karakter")
                    with col_reg3: paket_reg = st.selectbox("Pilih Paket Awal", ["BASIC", "EXECUTIVE", "MASTER", "NON-AKTIF"])
                    if st.form_submit_button("Daftarkan Klien", type="primary", use_container_width=True):
                        if email_reg and len(pass_reg) >= 6:
                            if email_reg in ADMIN_EMAILS_CONFIG: st.error("⚠️ Email ini terdaftar sebagai Admin.")
                            else:
                                with st.spinner("Mendaftarkan..."):
                                    new_user = register_firebase(email_reg, pass_reg)
                                    if "idToken" in new_user:
                                        uid, sekarang = new_user["localId"], datetime.now()
                                        db.collection("users").document(uid).set({"email": email_reg, "status_subscription": "aktif" if paket_reg != "NON-AKTIF" else "non-aktif", "paket": paket_reg, "kuota_ai": PAKET_LANGGANAN[paket_reg]["ai_limit"] if paket_reg != "NON-AKTIF" else 0, "kuota_upload": PAKET_LANGGANAN[paket_reg]["upload_limit"] if paket_reg != "NON-AKTIF" else 0, "tanggal_mulai": sekarang.isoformat(), "tanggal_berakhir": (sekarang + timedelta(days=30)).isoformat(), "reset_kuota_terakhir": sekarang.isoformat(), "last_login": "Belum pernah login", "login_count": 0})
                                        st.success(f"✅ Akun {email_reg} berhasil dibuat!")
                                        st.rerun()
                                    else: st.error(f"⚠️ Gagal mendaftar: {new_user.get('error', {}).get('message', 'Gagal')}")
                        else: st.warning("Pastikan email terisi dan password minimal 6 karakter.")

            st.markdown("#### 📊 Analitik Distribusi")
            col_chart1, col_chart2 = st.columns(2)
            with col_chart1:
                st.markdown("<p style='text-align: center; color: #64748b; font-size: 14px;'><b>Distribusi Status Pengguna</b></p>", unsafe_allow_html=True)
                st.bar_chart(pd.DataFrame([u['Status'] for u in users_list], columns=['Status']).value_counts().reset_index(name='Jumlah'), x="Status", y="Jumlah", color="#3b82f6")
            with col_chart2:
                st.markdown("<p style='text-align: center; color: #64748b; font-size: 14px;'><b>Distribusi Paket Langganan</b></p>", unsafe_allow_html=True)
                paket_list = [u['Paket'] for u in users_list if u['Paket'] != 'ADMIN']
                if paket_list: st.bar_chart(pd.DataFrame(paket_list, columns=['Paket']).value_counts().reset_index(name='Jumlah'), x="Paket", y="Jumlah", color="#10b981")
                else: st.info("Belum ada data paket klien.")

            st.markdown("---")
            st.markdown("### 📋 Tabel Manajemen Klien")
            col_search1, col_search2, col_search3 = st.columns([3, 2, 2])
            with col_search1: search_query = st.text_input("🔍 Cari email atau paket...", placeholder="Ketik untuk filter...")
            with col_search2: filter_status = st.selectbox("📊 Filter Status", ["Semua", "Aktif", "Non-Aktif", "Admin"])
            with col_search3: filter_paket = st.selectbox("📦 Filter Paket", ["Semua", "BASIC", "EXECUTIVE", "MASTER", "NON-AKTIF"])
            
            filtered_users = users_list.copy()
            if search_query: filtered_users = [u for u in filtered_users if search_query.lower() in u['Email'].lower() or search_query.lower() in u['Paket'].lower()]
            if filter_status != "Semua": filtered_users = [u for u in filtered_users if filter_status.lower() in str(u['Status']).lower()]
            if filter_paket != "Semua": filtered_users = [u for u in filtered_users if u['Paket'] == filter_paket]

            if filtered_users:
                df_display = pd.DataFrame([{ "📧 Email": u['Email'], "📊 Status": str(u['Status']).upper(), "📦 Paket": u['Paket'], "💎 AI": str(u['Sisa AI']), "📤 Upload": str(u['Sisa Upload']), "⏳ Sisa Hari": str(u['Sisa Hari']), "🕒 Last Login": u['Last Login'], "🔑 ID": u['UID_Short'] } for u in filtered_users])
                def color_status(val): return 'background-color: #dbeafe; color: #1e40af; font-weight: bold' if val == "ADMIN" else 'background-color: #d1fae5; color: #065f46; font-weight: bold' if val == "AKTIF" else 'background-color: #fee2e2; color: #991b1b; font-weight: bold' if "NON-AKTIF" in str(val) else ''
                def color_sisa_hari(val): return 'background-color: #dbeafe; color: #1e40af; font-weight: bold' if val == "∞" else 'background-color: #fee2e2; color: #991b1b; font-weight: bold' if "Kadaluarsa" in str(val) else 'background-color: #fef3c7; color: #92400e; font-weight: bold' if "⚠️" in str(val) else 'background-color: #d1fae5; color: #065f46' if "hari" in str(val) and int(''.join(filter(str.isdigit, str(val)))) > 7 else ''
                try: styled_df = df_display.astype(str).style.map(color_status, subset=['📊 Status']).map(color_sisa_hari, subset=['⏳ Sisa Hari'])
                except AttributeError: styled_df = df_display.astype(str).style.applymap(color_status, subset=['📊 Status']).applymap(color_sisa_hari, subset=['⏳ Sisa Hari'])
                st.dataframe(styled_df, use_container_width=True, hide_index=True, height=350)
            else: st.info("🔍 Tidak ada user yang sesuai dengan filter.")
                
            st.markdown("---")
            st.markdown("### 🛠️ Action Center (Kelola & Edit Klien)")
            user_options = [u['Email'] for u in filtered_users if u['Status'] != 'admin']
            if user_options:
                with st.container(border=True):
                    selected_email = st.selectbox("🎯 Pilih Akun Klien yang Ingin Dikelola:", user_options, key="action_center_select")
                    selected_user = next((u for u in filtered_users if u['Email'] == selected_email), None)
                    if selected_user:
                        st.markdown(f"""<div style='background: #f8fafc; padding: 15px 20px; border-radius: 10px; border-left: 5px solid #3b82f6; margin-bottom: 20px;'><h4 style='margin: 0; color: #1e293b;'>Pengaturan: <b>{selected_user['Email']}</b></h4><div style='display: flex; gap: 20px; margin-top: 10px; font-size: 14px; color: #475569;'><span>📦 Paket: <b>{selected_user['Paket']}</b></span><span>💎 AI: <b>{selected_user['Sisa AI']}</b></span><span>📤 Upload: <b>{selected_user['Sisa Upload']}</b></span><span>⏳ Masa Aktif: <b>{selected_user['Sisa Hari']}</b></span></div></div>""", unsafe_allow_html=True)
                        tab_edit, tab_extend, tab_topup, tab_history, tab_danger = st.tabs(["✏️ Ubah Paket", "📅 Perpanjang Masa Aktif", "🔋 Top-Up Kuota", "🔍 Riwayat & Reset", "❌ Hapus Akun"])
                        with tab_edit:
                            st.caption("Mengubah paket akan mengatur ulang kuota sesuai paket baru dan mereset masa aktif menjadi 30 hari dari sekarang.")
                            with st.form("form_edit_paket"):
                                col_e1, col_e2 = st.columns([3, 1])
                                with col_e1: new_paket = st.selectbox("Pilih Paket Baru", ["BASIC", "EXECUTIVE", "MASTER", "NON-AKTIF"], index=["BASIC", "EXECUTIVE", "MASTER", "NON-AKTIF"].index(selected_user['Paket']) if selected_user['Paket'] in ["BASIC", "EXECUTIVE", "MASTER", "NON-AKTIF"] else 0, label_visibility="collapsed")
                                if st.form_submit_button("💾 Simpan", type="primary", use_container_width=True):
                                    uid, sekarang = selected_user['UID'], datetime.now()
                                    if new_paket != "NON-AKTIF": db.collection("users").document(uid).update({"status_subscription": "aktif", "paket": new_paket, "kuota_ai": PAKET_LANGGANAN[new_paket]["ai_limit"], "kuota_upload": PAKET_LANGGANAN[new_paket]["upload_limit"], "tanggal_mulai": sekarang.isoformat(), "tanggal_berakhir": (sekarang + timedelta(days=30)).isoformat(), "reset_kuota_terakhir": sekarang.isoformat()})
                                    else: db.collection("users").document(uid).update({"status_subscription": "non-aktif", "paket": "NON-AKTIF", "kuota_ai": 0, "kuota_upload": 0, "tanggal_berakhir": sekarang.isoformat()})
                                    st.success(f"✅ Paket diubah ke {new_paket}.")
                                    time.sleep(1); st.rerun()
                        with tab_extend:
                            st.caption("Menambah masa aktif hari tanpa mengubah/mereset sisa kuota saat ini.")
                            with st.form("form_extend_hari"):
                                col_x1, col_x2 = st.columns([3, 1])
                                with col_x1: hari_tambahan = st.number_input("Tambah Masa Aktif (Hari)", min_value=1, max_value=365, value=30, label_visibility="collapsed")
                                if st.form_submit_button("⏳ Tambah Hari", type="primary", use_container_width=True):
                                    uid = selected_user['UID']
                                    tgl_lama = db.collection("users").document(uid).get().to_dict().get("tanggal_berakhir")
                                    try: t_baru = max(datetime.fromisoformat(tgl_lama) if isinstance(tgl_lama, str) else tgl_lama.replace(tzinfo=None), datetime.now()) + timedelta(days=hari_tambahan)
                                    except: t_baru = datetime.now() + timedelta(days=hari_tambahan)
                                    db.collection("users").document(uid).update({"status_subscription": "aktif", "tanggal_berakhir": t_baru.isoformat()})
                                    st.success(f"✅ Masa aktif ditambah {hari_tambahan} hari!"); time.sleep(1); st.rerun()
                        with tab_topup:
                            st.caption("Top-up atau tambah kuota AI & Upload secara eceran tanpa mengubah masa aktif paket bulanan.")
                            with st.form("form_topup_kuota"):
                                col_tp1, col_tp2 = st.columns(2)
                                with col_tp1: tambah_ai = st.number_input("➕ Tambah Kuota AI", min_value=0, max_value=500, value=10)
                                with col_tp2: tambah_upload = st.number_input("➕ Tambah Kuota Upload", min_value=0, max_value=50, value=3)
                                if st.form_submit_button("🔋 Eksekusi Top-Up", type="primary", use_container_width=True):
                                    if tambah_ai > 0 or tambah_upload > 0:
                                        db.collection("users").document(selected_user['UID']).update({"kuota_ai": firestore.Increment(tambah_ai), "kuota_upload": firestore.Increment(tambah_upload)})
                                        st.success("✅ Top-Up Berhasil!"); time.sleep(1); st.rerun()
                                    else: st.warning("⚠️ Masukkan nominal top-up lebih dari 0.")
                        with tab_history:
                            col_h1, col_h2 = st.columns([1, 1.5], gap="large")
                            with col_h1:
                                st.markdown("<h5 style='color:#334155; margin-bottom:10px;'>🔄 Reset Kuota Manual</h5>", unsafe_allow_html=True)
                                if st.button("🔄 Reset Kuota Sekarang", use_container_width=True):
                                    if selected_user['Paket'] in PAKET_LANGGANAN: reset_user_kuota(selected_user['UID'], selected_user['Paket']); st.success("✅ Kuota berhasil direset!"); time.sleep(1); st.rerun()
                                    else: st.error("Paket tidak valid untuk direset.")
                                st.markdown("<br><h5 style='color:#334155; margin-bottom:10px;'>▶️ Status Akses</h5>", unsafe_allow_html=True)
                                if selected_user['Status'] == 'non-aktif' and selected_user['Paket'] != 'NON-AKTIF':
                                    if st.button("Aktifkan Kembali Akses", type="primary", use_container_width=True):
                                        sekarang = datetime.now()
                                        db.collection("users").document(selected_user['UID']).update({"status_subscription": "aktif", "tanggal_berakhir": (sekarang + timedelta(days=30)).isoformat(), "tanggal_mulai": sekarang.isoformat()})
                                        st.success("✅ Akun diaktifkan kembali!"); time.sleep(1); st.rerun()
                                else: st.info("Akun sedang berjalan normal (Aktif).")
                            with col_h2:
                                st.markdown("<h5 style='color:#334155; margin-bottom:10px;'>📜 Riwayat Login</h5>", unsafe_allow_html=True)
                                if st.button("Tampilkan Riwayat", use_container_width=True):
                                    history = get_user_login_history(selected_user['UID'])
                                    if history: st.dataframe(pd.DataFrame([{"Waktu": h.get('timestamp', '')[:19].replace("T", " "), "Platform": h.get('platform', '-')} for h in history[:10]]), use_container_width=True, hide_index=True)
                                    else: st.info("Belum ada riwayat login.")
                        with tab_danger:
                            st.markdown("<h5 style='color:#ef4444; margin-bottom:10px;'>❌ Hapus Akun Klien Permanen</h5>", unsafe_allow_html=True)
                            if st.button("Hapus Permanen User Ini", type="secondary", use_container_width=True):
                                if st.session_state.get('confirm_delete') != selected_user['UID']:
                                    st.session_state['confirm_delete'] = selected_user['UID']
                                    st.error(f"⚠️ KLIK SEKALI LAGI UNTUK KONFIRMASI MENGHAPUS **{selected_user['Email']}**!")
                                else:
                                    success, msg = delete_user(selected_user['UID'], selected_user['Email'])
                                    if success: st.success(msg); st.session_state.pop('confirm_delete', None); time.sleep(1); st.rerun()
                                    else: st.error(msg)
            else: st.info("Tidak ada klien yang terdaftar atau sesuai filter untuk dikelola.")
    else:
        tabs = st.tabs(["🔴 Live Capture", "📁 Upload Rekaman", "📝 SOP Generator", "💳 Info Paket"])
        tab1, tab2, tab_sop, tab_paket = tabs[0], tabs[1], tabs[2], tabs[3]

    # =====================================================================
    # TAB INFO PAKET LANGGANAN (USER BIASA)
    # =====================================================================
    with tab_paket:
        st.markdown("### 📊 Pilihan Paket Langganan TranscribX")
        st.write("Tingkatkan produktivitas rapat Anda dengan memilih paket yang sesuai. Klik tombol beli di bawah untuk upgrade secara otomatis.")
        st.write("")
        col_p1, col_p2, col_p3 = st.columns(3)
        with col_p1:
            st.markdown("""<div style='background-color:#ffffff; padding:20px; border-radius:15px; border:1px solid #e2e8f0; height:100%; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05); transition: transform 0.2s;' onmouseover="this.style.transform='scale(1.02)'" onmouseout="this.style.transform='scale(1)'"><h3 style='color:#334155; margin-top:0;'>1. Paket BASIC</h3><h4 style='color:#3b82f6;'>Rp 35.000 <span style='font-size:14px; color:#94a3b8;'>/ 30 hari</span></h4><p style='font-size:14px; color:#64748b; margin-bottom:20px;'>Cocok untuk mahasiswa, asisten peneliti, atau staf admin.</p><hr style='border-color:#f1f5f9; margin-bottom:20px;'><ul style='font-size:14px; color:#334155; padding-left:20px; line-height:1.8;'><li>✅ <b>Unlimited</b> Live Transcribe</li><li>✅ <b>8x</b> Premium AI Summary</li><li>❌ <del>AI SOP Generator</del></li><li>✅ <b>1x</b> Upload File Audio</li><li>⏳ <b>30 Hari</b> Masa Aktif</li></ul></div>""", unsafe_allow_html=True)
        with col_p2:
            st.markdown("""<div style='background-color:#eff6ff; padding:20px; border-radius:15px; border:2px solid #3b82f6; height:100%; box-shadow: 0 10px 15px -3px rgba(0,0,0,0.1); position:relative; transition: transform 0.2s;' onmouseover="this.style.transform='scale(1.02)'" onmouseout="this.style.transform='scale(1)'"><div style='position:absolute; top:-12px; right:20px; background:#ef4444; color:white; padding:4px 12px; border-radius:20px; font-size:12px; font-weight:bold;'>🔥 Best Seller</div><h3 style='color:#1e3a8a; margin-top:0;'>2. Paket EXECUTIVE</h3><h4 style='color:#2563eb;'>Rp 69.000 <span style='font-size:14px; color:#94a3b8;'>/ 30 hari</span></h4><p style='font-size:14px; color:#475569; margin-bottom:20px;'>Cocok untuk ketua komite, manajer, atau profesional.</p><hr style='border-color:#bfdbfe; margin-bottom:20px;'><ul style='font-size:14px; color:#1e3a8a; padding-left:20px; line-height:1.8;'><li>✅ <b>Unlimited</b> Live Transcribe</li><li>✅ <b>25x</b> Premium AI Summary</li><li>✅ <b>AI SOP Generator</b></li><li>✅ <b>5x</b> Upload File Audio</li><li>🌟 Bisa <b>Top-Up Kuota</b> Kapan Saja</li><li>⏳ <b>30 Hari</b> Masa Aktif</li></ul></div>""", unsafe_allow_html=True)
        with col_p3:
            st.markdown("""<div style='background-color:#fff1f2; padding:20px; border-radius:15px; border:1px solid #fecdd3; height:100%; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05); transition: transform 0.2s;' onmouseover="this.style.transform='scale(1.02)'" onmouseout="this.style.transform='scale(1)'"><h3 style='color:#881337; margin-top:0;'>3. Paket MASTER / VIP</h3><h4 style='color:#e11d48;'>Rp 149.000 <span style='font-size:14px; color:#94a3b8;'>/ 30 hari</span></h4><p style='font-size:14px; color:#64748b; margin-bottom:20px;'>Cocok untuk panitia masterclass atau institusi.</p><hr style='border-color:#ffe4e6; margin-bottom:20px;'><ul style='font-size:14px; color:#881337; padding-left:20px; line-height:1.8;'><li>✅ <b>Unlimited</b> Live Transcribe</li><li>✅ <b>60x</b> Premium AI Summary</li><li>✅ <b>AI SOP Generator</b></li><li>✅ <b>15x</b> Upload File Audio</li><li>🌟 <b>Prioritas Support</b> via WhatsApp</li><li>⏳ <b>30 Hari</b> Masa Aktif</li></ul></div>""", unsafe_allow_html=True)

        user_email_aktif = st.session_state.get('user_email', '')
        email_encoded = urllib.parse.quote(user_email_aktif)

        st.markdown("<h3 style='text-align: center; color: #38bdf8; margin-top: 40px; font-weight: bold;'>🔋 Paket Top-Up Kuota (Eceran)</h3>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: center; color: #64748b; margin-bottom: 20px;'>Kehabisan kuota sebelum masa aktif habis? Beli kuota tambahan tanpa harus memperpanjang masa aktif bulanan.</p>", unsafe_allow_html=True)
        
        link_tp_ai_10 = f"https://lynk.id/gerrysutsuga/LINK_TOPUP_AI_10/checkout?email={email_encoded}"
        link_tp_ai_25 = f"https://lynk.id/gerrysutsuga/LINK_TOPUP_AI_25/checkout?email={email_encoded}"
        link_tp_mp3_3 = f"https://lynk.id/gerrysutsuga/LINK_TOPUP_MP3_3/checkout?email={email_encoded}"
        link_tp_mp3_10 = f"https://lynk.id/gerrysutsuga/LINK_TOPUP_MP3_10/checkout?email={email_encoded}"

        col_t1, col_t2 = st.columns(2)
        with col_t1:
            st.markdown(f"""<div style='background-color:#f8fafc; padding:20px; border-radius:15px; border:1px solid #e2e8f0; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);'>
<h4 style='color:#8b5cf6; margin-top:0;'>✨ Top-Up AI Summary</h4>
<hr style='border-color:#e2e8f0; margin:10px 0;'>
<div style='display:flex; align-items:center; justify-content:space-between; margin-bottom:12px; padding: 12px; background: #ffffff; border-radius: 8px; border: 1px solid #f1f5f9;'>
<div><div style='font-size:15px; font-weight:bold; color:#334155;'>+ 10x Kuota AI Summary</div><div style='font-size:14px; font-weight:bold; color:#3b82f6;'>Rp 10.000</div></div>
<a href="{link_tp_ai_10}" target="_blank" style="text-decoration:none;"><button style="background:#8b5cf6; color:white; border:none; padding:8px 16px; border-radius:6px; cursor:pointer; font-weight:bold; font-size:13px; transition:0.2s;" onmouseover="this.style.opacity='0.8'" onmouseout="this.style.opacity='1'">Beli</button></a>
</div>
<div style='display:flex; align-items:center; justify-content:space-between; padding: 12px; background: #ffffff; border-radius: 8px; border: 1px solid #f1f5f9;'>
<div><div style='font-size:15px; font-weight:bold; color:#334155;'>+ 25x Kuota AI Summary</div><div style='font-size:14px; font-weight:bold; color:#3b82f6;'>Rp 20.000</div></div>
<a href="{link_tp_ai_25}" target="_blank" style="text-decoration:none;"><button style="background:#8b5cf6; color:white; border:none; padding:8px 16px; border-radius:6px; cursor:pointer; font-weight:bold; font-size:13px; transition:0.2s;" onmouseover="this.style.opacity='0.8'" onmouseout="this.style.opacity='1'">Beli</button></a>
</div>
</div>""", unsafe_allow_html=True)
            
        with col_t2:
            st.markdown(f"""<div style='background-color:#f8fafc; padding:20px; border-radius:15px; border:1px solid #e2e8f0; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);'>
<h4 style='color:#10b981; margin-top:0;'>📁 Top-Up Upload MP3</h4>
<hr style='border-color:#e2e8f0; margin:10px 0;'>
<div style='display:flex; align-items:center; justify-content:space-between; margin-bottom:12px; padding: 12px; background: #ffffff; border-radius: 8px; border: 1px solid #f1f5f9;'>
<div><div style='font-size:15px; font-weight:bold; color:#334155;'>+ 3x Upload MP3</div><div style='font-size:14px; font-weight:bold; color:#3b82f6;'>Rp 10.000</div></div>
<a href="{link_tp_mp3_3}" target="_blank" style="text-decoration:none;"><button style="background:#10b981; color:white; border:none; padding:8px 16px; border-radius:6px; cursor:pointer; font-weight:bold; font-size:13px; transition:0.2s;" onmouseover="this.style.opacity='0.8'" onmouseout="this.style.opacity='1'">Beli</button></a>
</div>
<div style='display:flex; align-items:center; justify-content:space-between; padding: 12px; background: #ffffff; border-radius: 8px; border: 1px solid #f1f5f9;'>
<div><div style='font-size:15px; font-weight:bold; color:#334155;'>+ 10x Upload MP3</div><div style='font-size:14px; font-weight:bold; color:#3b82f6;'>Rp 25.000</div></div>
<a href="{link_tp_mp3_10}" target="_blank" style="text-decoration:none;"><button style="background:#10b981; color:white; border:none; padding:8px 16px; border-radius:6px; cursor:pointer; font-weight:bold; font-size:13px; transition:0.2s;" onmouseover="this.style.opacity='0.8'" onmouseout="this.style.opacity='1'">Beli</button></a>
</div>
</div>""", unsafe_allow_html=True)
            
        st.write("")
        st.info("💡 **INFO UPGRADE PAKET:** Membeli paket baru saat paket lama masih aktif akan **mereset ulang** masa aktif Anda menjadi 30 hari ke depan dan mengatur ulang kuota sesuai paket terbaru (sisa kuota lama tidak diakumulasi). Jika hanya ingin menambah kuota, gunakan fitur **Top-Up Eceran** di atas.")

        st.markdown("<hr style='margin-top: 30px; border-color:#e2e8f0;'>", unsafe_allow_html=True)
        st.markdown("<h3 style='text-align: center; color: #38bdf8;'>🚀 Pilih & Beli Paket Sekarang</h3>", unsafe_allow_html=True)
        link_basic = f"http://lynk.id/gerrysutsuga/yw8d3d5r1m5l/checkout?email={email_encoded}"
        link_executive = f"https://lynk.id/gerrysutsuga/yw8d3d5r1m5l/checkout?email={email_encoded}"
        link_master = f"https://lynk.id/gerrysutsuga/LINK_PRODUK_MASTER_KAMU/checkout?email={email_encoded}"

        col_btn1, col_btn2, col_btn3 = st.columns(3)
        with col_btn1: st.markdown(f"""<a href="{link_basic}" target="_blank" style="text-decoration: none;"><button style="width: 100%; padding: 12px; border-radius: 8px; background-color: #f1f5f9; color: #334155; border: 1px solid #cbd5e1; font-weight: bold; cursor: pointer; transition: 0.3s;">🛒 Beli Paket BASIC</button></a>""", unsafe_allow_html=True)
        with col_btn2: st.markdown(f"""<a href="{link_executive}" target="_blank" style="text-decoration: none;"><button style="width: 100%; padding: 12px; border-radius: 8px; background-color: #3b82f6; color: white; border: none; font-weight: bold; cursor: pointer; box-shadow: 0 4px 6px rgba(59,130,246,0.3); transition: 0.3s;">💳 Beli EXECUTIVE</button></a>""", unsafe_allow_html=True)
        with col_btn3: st.markdown(f"""<a href="{link_master}" target="_blank" style="text-decoration: none;"><button style="width: 100%; padding: 12px; border-radius: 8px; background-color: #e11d48; color: white; border: none; font-weight: bold; cursor: pointer; box-shadow: 0 4px 6px rgba(225,29,72,0.3); transition: 0.3s;">👑 Beli MASTER</button></a>""", unsafe_allow_html=True)

        st.write("")
        if st.button("🔄 Cek Status Pembayaran Saya", use_container_width=True, type="secondary"):
            with st.spinner("Menyinkronkan data..."):
                time.sleep(2); render_sidebar_profile(); st.success("✅ Sinkronisasi selesai!"); time.sleep(2); st.rerun()

    # =====================================================================
    # TAB 1: LIVE CAPTURE DENGAN ANIMASI OTAK AI
    # =====================================================================
    with tab1:
        st.markdown("### 🎙️ Live Transcribe - Screen Capture (Zoom / YouTube)")
        st.info("💡 **TIPS:** Klik Start Capture → Pilih tab/window yang menjalankan Zoom atau YouTube → Centang **'Share tab audio'** → Klik Share.")
        st.warning("⚠️ **PENTING:** Saat dialog share muncul, pastikan kamu memilih tab/window Zoom/YouTube dan **CENTANG 'Share tab audio'**!")
        
        html_code = """
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <script src="https://cdn.tailwindcss.com"></script>
            <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
            <script src="https://cdn.jsdelivr.net/npm/d3@7/dist/d3.min.js"></script>
            <script src="https://cdn.jsdelivr.net/npm/markmap-lib@0.15.4/dist/browser/index.js"></script>
            <script src="https://cdn.jsdelivr.net/npm/markmap-view@0.15.4/dist/browser/index.js"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/cytoscape/3.26.0/cytoscape.min.js"></script>
            <script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js"></script>
            <script src="https://cdn.jsdelivr.net/npm/echarts@5.5.0/dist/echarts.min.js"></script>
            <style>
                * { box-sizing: border-box; }
                body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; background: transparent; margin: 0; padding: 10px; color: #1e293b; }
                
                .controls-wrapper { 
                    background: #ffffff; border-radius: 20px; padding: 24px; 
                    box-shadow: 0 10px 25px -5px rgba(0,0,0,0.05), 0 8px 10px -6px rgba(0,0,0,0.01); 
                    border: 1px solid #e2e8f0; margin-bottom: 24px;
                    display: flex; flex-direction: column; gap: 16px;
                }
                
                .controls-row { display: flex; flex-wrap: wrap; align-items: center; gap: 12px; }
                
                .visualizer-container { 
                    width: 100%; height: 80px; border-radius: 16px; overflow: hidden; 
                    background: linear-gradient(180deg, #0f172a 0%, #1e293b 100%); 
                    position: relative; box-shadow: inset 0 4px 6px rgba(0,0,0,0.3);
                }
                #visualizer { width: 100%; height: 100%; display: block; }
                
                .transcript-box { 
                    background: #f8fafc; border: 1px solid #cbd5e1; padding: 16px 20px; 
                    border-radius: 20px; height: 300px; overflow-y: auto; 
                    box-shadow: inset 0 2px 4px rgba(0,0,0,0.02); margin-bottom: 16px; 
                }
                
                .line-final { 
                    margin-bottom: 12px; padding: 12px 16px; background: #ffffff; 
                    border-radius: 12px; border-left: 5px solid #3b82f6; 
                    font-size: 14px; line-height: 1.6; box-shadow: 0 2px 4px rgba(0,0,0,0.02);
                }
                .line-interim { 
                    margin-bottom: 12px; padding: 12px 16px; background: rgba(255,255,255,0.5); 
                    border-radius: 12px; border-left: 5px solid #94a3b8; 
                    font-size: 14px; opacity: 0.6; font-style: italic; 
                }
                .timestamp { font-weight: 700; color: #3b82f6; margin-right: 10px; font-size: 12px; background: #eff6ff; padding: 2px 8px; border-radius: 6px; display: inline-block; }
                
                .btn-custom { 
                    font-family: inherit; color: white; padding: 10px 20px; border: none; 
                    border-radius: 10px; cursor: pointer; font-weight: 600; 
                    transition: all 0.2s ease; display: inline-flex; align-items: center; gap: 8px; font-size: 14px; 
                    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                }
                .btn-custom:hover { transform: translateY(-2px); box-shadow: 0 4px 8px rgba(0,0,0,0.15); }
                .btn-custom:active { transform: translateY(0); }
                .btn-custom:disabled { background: #cbd5e1 !important; cursor: not-allowed; transform: none; box-shadow: none; color: #64748b; }
                
                .btn-start { background: #3b82f6; } .btn-start:hover { background: #2563eb; }
                .btn-stop { background: #ef4444; } .btn-stop:hover { background: #dc2626; }
                .btn-ai { background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%); } .btn-ai:hover { background: linear-gradient(135deg, #7c3aed 0%, #6d28d9 100%); }
                .btn-secondary { background: #f1f5f9; color: #475569; border: 1px solid #cbd5e1; box-shadow: none; } 
                .btn-secondary:hover { background: #e2e8f0; color: #1e293b; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
                .btn-green { background: #10b981; } .btn-green:hover { background: #059669; }
                
                select.btn-secondary, input.api-input { 
                    outline: none; border: 1px solid #cbd5e1; padding: 10px 14px; 
                    border-radius: 10px; font-family: inherit; font-size: 14px; background: white;
                }
                select.btn-secondary:focus, input.api-input:focus { border-color: #3b82f6; box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1); }
                input.api-input { flex: 1; min-width: 200px; }
                
                .status-badge { display: flex; align-items: center; gap: 10px; background: #f8fafc; padding: 6px 16px; border-radius: 20px; border: 1px solid #e2e8f0; font-weight: 600; font-size: 13px; }
                .indicator-dot { width: 12px; height: 12px; border-radius: 50%; background: #cbd5e1; transition: all 0.3s; flex-shrink: 0; }
                .indicator-dot.recording { background: #ef4444; box-shadow: 0 0 10px #ef4444; animation: blink 1s infinite; }
                @keyframes blink { 0%, 100% { opacity: 1; } 50% { opacity: 0.3; } }
                
                .ai-section { background: #f8fafc; border: 1px solid #e2e8f0; padding: 16px 20px; border-radius: 16px; display: flex; flex-direction: column; gap: 12px; margin-top: 8px; }
                .ai-section .ai-row { display: flex; flex-wrap: wrap; align-items: center; gap: 12px; }
                
                #audioContainer { display: flex; flex-direction: column; gap: 10px; margin-top: 12px; }
                .audio-item { display: flex; align-items: center; gap: 12px; padding: 12px 16px; background: #ffffff; border: 1px solid #e2e8f0; border-radius: 12px; box-shadow: 0 2px 4px rgba(0,0,0,0.02); flex-wrap: wrap; }
                .audio-item audio { height: 36px; flex: 1; min-width: 200px; }
                
                .fade-in { animation: fadeIn 0.5s ease; }
                @keyframes fadeIn { from { opacity: 0; transform: translateY(10px); } to { opacity: 1; transform: translateY(0); } }
                
                .instruction-box { background: #fffbeb; border: 2px solid #f59e0b; padding: 16px; border-radius: 12px; margin-bottom: 16px; font-size: 13px; color: #92400e; }

                @media (max-width: 640px) {
                    .controls-row { flex-direction: column; align-items: stretch; }
                    .controls-row > * { width: 100%; }
                    .status-badge { justify-content: center; }
                    .ai-section .ai-row { flex-direction: column; }
                    .ai-section .ai-row > * { width: 100%; }
                }
            </style>
        </head>
        <body>
            <div class="instruction-box">
                <strong>📺 CARA SCREEN CAPTURE:</strong><br>
                1. Buka <b>Zoom/YouTube</b> di tab browser <b>TERPISAH</b>.<br>
                2. Pastikan <b>VOLUME SPEAKER LAPTOP ANDA NYALA</b>.<br>
                3. Klik <b>"Start Capture"</b>, izinkan akses Microphone jika diminta.<br>
                4. Pilih tab/window <b>Zoom</b> atau <b>YouTube</b>.<br>
                5. <b>CENTANG "Share tab audio"</b> lalu klik Share.
            </div>
            
            <div class="controls-wrapper">
                <div class="controls-row">
                    <select id="langSelect" class="btn-custom btn-secondary" style="min-width:120px;">
                        <option value="id-ID">🇮🇩 Indonesia</option>
                        <option value="en-US">🇬🇧 English</option>
                        <option value="ja-JP">🇯🇵 Japanese</option>
                    </select>
                    <button id="startBtn" class="btn-custom btn-start">▶️ Start Capture</button>
                    <button id="stopBtn" class="btn-custom btn-stop" disabled>⏹️ Stop & Simpan Rekaman</button>
                    <div class="status-badge" style="margin-left:auto;">
                        <div id="indicator" class="indicator-dot"></div>
                        <span id="status" style="color: #64748b;">Standby</span>
                    </div>
                </div>
                
                <div class="visualizer-container">
                    <canvas id="visualizer"></canvas>
                </div>
                
                <div class="ai-section">
                    <div id="aiBrainContainer" style="width: 100%; height: 180px; background: #1e1e2f; border-radius: 12px; margin-bottom: 15px; position: relative; overflow: hidden; box-shadow: inset 0 0 30px rgba(0,0,0,0.6); border: 1px solid #334155;">
                        <canvas id="aiBrainCanvas"></canvas>
                        <div id="aiBrainText" style="position: absolute; top: 50%; left: 50%; transform: translate(-50%, -50%); color: #64748b; font-size: 13px; font-weight: 800; pointer-events: none; letter-spacing: 5px; z-index: 10; text-align: center; text-transform: uppercase;">NEURAL NETWORK IDLE</div>
                    </div>
                    
                    <div class="ai-row">
                        <span style="font-size: 13px; font-weight: 700; color: #475569; white-space:nowrap;">🔑 API Key:</span>
                        <input type="password" id="apiKeyInput" class="api-input" placeholder="Masukkan API Key LiteLLM / Gemini..." style="flex:1; min-width:150px;">
                        <button id="aiBtn" class="btn-custom btn-ai" style="white-space:nowrap;">✨ Generate AI Summary</button>
                    </div>
                </div>
            </div>

            <div style="display: flex; flex-wrap: wrap; justify-content: flex-end; gap: 8px; margin-bottom: 12px;">
                <button id="copyBtn" class="btn-custom btn-secondary" style="padding: 6px 14px; font-size: 13px;">📋 Copy</button>
                <button id="clearBtn" class="btn-custom btn-secondary" style="padding: 6px 14px; font-size: 13px;">🗑️ Clear</button>
                <button id="downloadTxtBtn" class="btn-custom btn-secondary" style="padding: 6px 14px; font-size: 13px;">📝 Save Raw TXT</button>
                <button id="downloadNotulensiBtn" class="btn-custom btn-green" style="padding: 6px 14px; font-size: 13px; display: none;">📑 Download Full Notulensi</button>
                <button id="downloadDocxBtn" class="btn-custom btn-ai" style="padding: 6px 14px; font-size: 13px; display: none; background: #ef4444;">📄 Download Resmi (DOCX)</button>
            </div>

            <div id="transcriptBox" class="transcript-box">
                <div id="placeholder" style="text-align: center; color: #94a3b8; margin-top: 100px; font-weight: 600;">
                    🎤 Klik "Start Capture" → Pilih tab Zoom/YouTube → Centang "Share audio" → Share
                </div>
            </div>

            <div id="aiContent" class="w-full"></div>
            
            <div style="margin-top: 24px; background: #ffffff; padding: 16px 20px; border-radius: 16px; border: 1px solid #e2e8f0;">
                <h3 style="margin: 0 0 12px 0; font-size: 15px; color: #1e293b; font-weight: 700;">🎧 Arsip Rekaman Screen Capture</h3>
                <div id="audioContainer">
                    <p style="color:#94a3b8; font-size:13px; text-align:center;" id="audioPlaceholder">Rekaman akan muncul di sini setelah Stop</p>
                </div>
            </div>

            <script>
                (function() {
                    'use strict';

                    // ======== DOM REFS ========
                    const startBtn = document.getElementById('startBtn');
                    const stopBtn = document.getElementById('stopBtn');
                    const copyBtn = document.getElementById('copyBtn');
                    const clearBtn = document.getElementById('clearBtn');
                    const downloadTxtBtn = document.getElementById('downloadTxtBtn');
                    const downloadNotulensiBtn = document.getElementById('downloadNotulensiBtn');
                    const downloadDocxBtn = document.getElementById('downloadDocxBtn');
                    const aiBtn = document.getElementById('aiBtn');
                    const apiKeyInput = document.getElementById('apiKeyInput');
                    const aiContent = document.getElementById('aiContent');
                    const langSelect = document.getElementById('langSelect');
                    const status = document.getElementById('status');
                    const indicator = document.getElementById('indicator');
                    const transcriptBox = document.getElementById('transcriptBox');
                    const audioContainer = document.getElementById('audioContainer');
                    const visualizer = document.getElementById('visualizer');
                    const canvasCtx = visualizer.getContext('2d');

                    let lastAiData = null;

                    if (window.mermaid) {
                        mermaid.initialize({ startOnLoad: false, theme: 'default' });
                    }

                    // ======== AI BRAIN VISUALIZER ========
                    const brainCanvas = document.getElementById('aiBrainCanvas');
                    const brainCtx = brainCanvas ? brainCanvas.getContext('2d') : null;
                    const brainText = document.getElementById('aiBrainText');
                    let brainParticles = [];
                    let isThinking = false;
                    let timeOscillator = 0;

                    function initBrain() {
                        if (!brainCanvas) return;
                        const rect = brainCanvas.parentElement.getBoundingClientRect();
                        brainCanvas.width = rect.width; brainCanvas.height = 180;
                        brainParticles = [];
                        const numParticles = 160;
                        for(let i=0; i<numParticles; i++) {
                            let isCore = Math.random() > 0.95;
                            brainParticles.push({
                                angle: Math.random() * Math.PI * 2,
                                orbitRadius: Math.random() * 65 + 5, 
                                speed: Math.random() * 0.015 + 0.002,
                                baseRadius: isCore ? (Math.random() * 2 + 2) : (Math.random() * 1.5 + 0.5),
                                isCore: isCore, x: 0, y: 0
                            });
                        }
                    }

                    function drawBrain() {
                        if (!brainCanvas) return;
                        brainCtx.clearRect(0, 0, brainCanvas.width, brainCanvas.height);
                        timeOscillator += 0.02;
                        
                        const cx = (brainCanvas.width / 2) + Math.cos(timeOscillator) * 15;
                        const cy = (brainCanvas.height / 2) + Math.sin(timeOscillator * 0.8) * 8;
                        const maxDistance = isThinking ? 40 : 25;
                        const nodeColor = isThinking ? "rgba(216, 180, 254, 0.9)" : "rgba(148, 163, 184, 0.8)"; 
                        const coreColor = isThinking ? "rgba(45, 212, 191, 1)" : "rgba(255, 255, 255, 1)";
                        
                        for(let i=0; i<brainParticles.length; i++) {
                            let p = brainParticles[i];
                            p.angle += isThinking ? p.speed * 4 : p.speed;
                            let breath = Math.sin(timeOscillator * 2 + p.angle) * (isThinking ? 15 : 5);
                            let currentRadius = p.orbitRadius + breath;
                            
                            p.x = cx + Math.cos(p.angle) * currentRadius * 1.5; 
                            p.y = cy + Math.sin(p.angle) * currentRadius;
                            
                            brainCtx.beginPath();
                            brainCtx.arc(p.x, p.y, isThinking ? p.baseRadius * 1.2 : p.baseRadius, 0, Math.PI * 2);
                            brainCtx.fillStyle = p.isCore ? coreColor : nodeColor;
                            
                            if (isThinking) {
                                brainCtx.shadowBlur = p.isCore ? 15 : 5;
                                brainCtx.shadowColor = "#2dd4bf";
                            } else { brainCtx.shadowBlur = 0; }
                            brainCtx.fill();
                        }
                        
                        for(let i=0; i<brainParticles.length; i++) {
                            let p = brainParticles[i];
                            for(let j=i+1; j<brainParticles.length; j++) {
                                let p2 = brainParticles[j];
                                let dx = p.x - p2.x; let dy = p.y - p2.y;
                                let dist = Math.sqrt(dx*dx + dy*dy);
                                if(dist < maxDistance) {
                                    brainCtx.beginPath(); brainCtx.moveTo(p.x, p.y); brainCtx.lineTo(p2.x, p2.y);
                                    let opacity = 1 - (dist / maxDistance);
                                    let alpha = isThinking ? opacity * 0.7 : opacity * 0.2;
                                    brainCtx.strokeStyle = `rgba(148, 163, 184, ${alpha})`;
                                    brainCtx.lineWidth = isThinking ? 1.0 : 0.4;
                                    brainCtx.stroke();
                                }
                            }
                        }
                        requestAnimationFrame(drawBrain);
                    }

                    if (brainCanvas) {
                        setTimeout(() => { initBrain(); drawBrain(); }, 500);
                        window.addEventListener('resize', initBrain);
                    }

                    // ======== STATE MANAGEMENT ========
                    let isRecording = false;
                    let isVisualizerActive = false;
                    let recognition = null;
                    let mediaRecorder = null;
                    let audioChunks = [];
                    let displayStream = null;
                    let globalAudioCtx = null;
                    let analyser = null;
                    let dataArray = null;
                    let drawVisual = null;
                    let currentInterimDiv = null;
                    let lastFinalText = "";

                    function initSpeechRecognition() {
                        const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
                        if (!SpeechRecognition) { status.innerText = "❌ Browser tidak mendukung Speech API"; return null; }
                        const rec = new SpeechRecognition();
                        rec.continuous = true; rec.interimResults = true; rec.lang = langSelect.value;
                        rec.onstart = function() {
                            isRecording = true;
                            status.innerText = "🎤 Menangkap audio... (Pastikan volume speaker nyala!)";
                            indicator.className = 'indicator-dot recording';
                            startBtn.disabled = true; stopBtn.disabled = false;
                            const placeholder = document.getElementById('placeholder');
                            if (placeholder) placeholder.style.display = 'none';
                        };
                        rec.onerror = function(event) {
                            if (event.error === 'no-speech' || event.error === 'aborted') return;
                            status.innerText = "⚠️ Speech Error: " + event.error;
                        };
                        rec.onend = function() {
                            if (isRecording) { setTimeout(() => { try { rec.start(); } catch(e) {} }, 200); }
                        };
                        rec.onresult = function(event) {
                            let interimTranscript = ''; let finalTranscript = '';
                            for (let i = event.resultIndex; i < event.results.length; i++) {
                                const result = event.results[i];
                                if (result.isFinal) finalTranscript += result[0].transcript + ' ';
                                else interimTranscript += result[0].transcript;
                            }
                            if (finalTranscript.trim()) {
                                const cleanFinal = finalTranscript.trim();
                                if (cleanFinal === lastFinalText.trim()) return;
                                lastFinalText = cleanFinal;
                                const now = new Date();
                                const timeStr = '[' + String(now.getHours()).padStart(2,'0') + ':' + String(now.getMinutes()).padStart(2,'0') + ':' + String(now.getSeconds()).padStart(2,'0') + ']';
                                if (currentInterimDiv) {
                                    currentInterimDiv.className = 'line-final';
                                    currentInterimDiv.innerHTML = '<span class="timestamp">' + timeStr + '</span> ' + cleanFinal;
                                    currentInterimDiv = null;
                                } else {
                                    const line = document.createElement('div');
                                    line.className = 'line-final';
                                    line.innerHTML = '<span class="timestamp">' + timeStr + '</span> ' + cleanFinal;
                                    transcriptBox.appendChild(line);
                                }
                                transcriptBox.scrollTop = transcriptBox.scrollHeight;
                            } else if (interimTranscript.trim()) {
                                if (!currentInterimDiv) {
                                    currentInterimDiv = document.createElement('div');
                                    currentInterimDiv.className = 'line-interim';
                                    transcriptBox.appendChild(currentInterimDiv);
                                }
                                currentInterimDiv.textContent = '💬 ' + interimTranscript;
                                transcriptBox.scrollTop = transcriptBox.scrollHeight;
                            }
                        };
                        return rec;
                    }

                    function setupVisualizer(stream) {
                        try {
                            const audioTracks = stream.getAudioTracks();
                            if (audioTracks.length === 0) return;
                            const audioOnlyStream = new MediaStream(audioTracks);
                            const source = globalAudioCtx.createMediaStreamSource(audioOnlyStream);
                            analyser = globalAudioCtx.createAnalyser();
                            analyser.fftSize = 256;
                            dataArray = new Uint8Array(analyser.frequencyBinCount);
                            source.connect(analyser);
                        } catch(e) { console.error(e); }
                    }

                    function drawVisualizer() {
                        if (!isVisualizerActive) return;
                        const width = visualizer.width; const height = visualizer.height;
                        canvasCtx.clearRect(0, 0, width, height);
                        if (analyser && dataArray) {
                            analyser.getByteFrequencyData(dataArray);
                            const bufferLength = analyser.frequencyBinCount;
                            const barWidth = (width / bufferLength) * 2.5; let x = 0;
                            for (let i = 0; i < bufferLength; i++) {
                                const barHeight = dataArray[i] / 2;
                                canvasCtx.fillStyle = 'rgb(' + Math.min(barHeight + 100, 255) + ',158,11)';
                                const y = (height / 2) - (barHeight / 2);
                                canvasCtx.fillRect(x, y, Math.max(barWidth, 1), Math.max(barHeight, 2));
                                x += barWidth + 1;
                            }
                        }
                        drawVisual = requestAnimationFrame(drawVisualizer);
                    }

                    function setupMediaRecorder(stream) {
                        audioChunks = [];
                        let options = { mimeType: 'video/webm;codecs=vp9,opus' };
                        if (!MediaRecorder.isTypeSupported(options.mimeType)) options = { mimeType: 'video/webm' };
                        try { mediaRecorder = new MediaRecorder(stream, options.mimeType ? options : undefined); } 
                        catch(e) { mediaRecorder = new MediaRecorder(stream); }
                        mediaRecorder.ondataavailable = function(e) { if (e.data && e.data.size > 0) audioChunks.push(e.data); };
                        mediaRecorder.onstop = function() {
                            if (audioChunks.length > 0) {
                                const mimeType = mediaRecorder.mimeType || 'video/webm';
                                const blob = new Blob(audioChunks, { type: mimeType });
                                const audioUrl = URL.createObjectURL(blob);
                                const audioItem = document.createElement('div');
                                audioItem.className = 'audio-item fade-in';
                                audioItem.innerHTML = `
                                    <audio controls src="${audioUrl}" preload="auto" style="flex:1; min-width:200px;"></audio>
                                    <a href="${audioUrl}" download="Capture_${Date.now()}.webm" class="btn-custom btn-green" style="padding:6px 14px; font-size:12px; text-decoration:none;">💾 Download</a>`;
                                document.getElementById('audioPlaceholder')?.remove();
                                audioContainer.appendChild(audioItem);
                            }
                        };
                    }

                    startBtn.onclick = async function() {
                        try {
                            lastFinalText = ""; currentInterimDiv = null; audioChunks = [];
                            transcriptBox.innerHTML = '<div style="text-align: center; color: #94a3b8; margin-top: 100px; font-weight: 600;">🎤 Menangkap audio... Rapat/Tab sedang berjalan.</div>';
                            await navigator.mediaDevices.getUserMedia({ audio: true });
                            if (!globalAudioCtx) globalAudioCtx = new (window.AudioContext || window.webkitAudioContext)();
                            if (globalAudioCtx.state === 'suspended') globalAudioCtx.resume();
                            
                            displayStream = await navigator.mediaDevices.getDisplayMedia({ 
                                video: { width: 640, height: 480, frameRate: 1 },
                                audio: { echoCancellation: false, noiseSuppression: false }
                            });
                            
                            if (displayStream.getAudioTracks().length === 0) {
                                status.innerText = "⚠️ CENTANG 'Share tab audio'!";
                                displayStream.getVideoTracks().forEach(t => t.stop());
                                return;
                            }
                            
                            setupMediaRecorder(displayStream);
                            mediaRecorder.start(1000);
                            setupVisualizer(displayStream);
                            isVisualizerActive = true; drawVisualizer();
                            
                            recognition = initSpeechRecognition();
                            if (recognition) recognition.start();
                            
                            displayStream.getVideoTracks()[0].addEventListener('ended', () => { if (isRecording) stopBtn.click(); });
                            resizeVisualizer();
                        } catch(err) { status.innerText = "❌ Gagal: " + err.message; }
                    };

                    stopBtn.onclick = function() {
                        isRecording = false; isVisualizerActive = false;
                        if (drawVisual) cancelAnimationFrame(drawVisual);
                        if (recognition) { try { recognition.stop(); } catch(e) {} recognition = null; }
                        if (mediaRecorder && mediaRecorder.state === 'recording') mediaRecorder.stop();
                        if (displayStream) { displayStream.getTracks().forEach(t => t.stop()); displayStream = null; }
                        status.innerText = "⏸️ Stopped - Cek Arsip Rekaman ↓"; indicator.className = 'indicator-dot';
                        startBtn.disabled = false; stopBtn.disabled = true; langSelect.disabled = false;
                    };

                    function resizeVisualizer() {
                        visualizer.width = visualizer.parentElement.clientWidth || 600;
                        visualizer.height = 80;
                    }
                    window.addEventListener('resize', resizeVisualizer);

                    copyBtn.onclick = function() {
                        const lines = transcriptBox.querySelectorAll('.line-final');
                        if (lines.length === 0) { alert('Belum ada teks!'); return; }
                        const text = Array.from(lines).map(line => line.innerText).join('\\n');
                        navigator.clipboard.writeText(text).then(() => {
                            copyBtn.innerText = '✅ Copied!';
                            setTimeout(() => copyBtn.innerText = '📋 Copy', 2000);
                        });
                    };

                    downloadTxtBtn.onclick = function() {
                        const lines = transcriptBox.querySelectorAll('.line-final');
                        if (lines.length === 0) { alert('Belum ada teks!'); return; }
                        const text = Array.from(lines).map(line => line.innerText).join('\\n');
                        const blob = new Blob([text], { type: 'text/plain' });
                        const a = document.createElement('a');
                        a.href = URL.createObjectURL(blob);
                        a.download = 'Transkrip_Live_' + Date.now() + '.txt';
                        a.click();
                    };
                    
                    downloadNotulensiBtn.onclick = function() {
                        if (!lastAiData) { alert('Belum ada data notulensi! Silakan Generate AI Summary terlebih dahulu.'); return; }
                        const d = lastAiData.notulensi_rapat || {};
                        let text = "NOTULENSI RAPAT SMARTDOSE ENTERPRISE\\n";
                        text += "========================================\\n\\n";
                        text += "🌟 RINGKASAN EKSEKUTIF:\\n";
                        (lastAiData.ringkasan_eksekutif || []).forEach(r => text += "- " + r + "\\n");
                        text += "\\n📌 AGENDA: " + (d.agenda || "-") + "\\n";
                        text += "👥 PESERTA: " + (d.peserta ? d.peserta.join(', ') : "-") + "\\n\\n";
                        if (d.transkrip_dialog && d.transkrip_dialog.length > 0) {
                            text += "💬 TRANSKRIP DIALOG:\\n";
                            d.transkrip_dialog.forEach(l => text += l + "\\n");
                            text += "\\n";
                        }
                        text += "🗣️ JALANNYA DISKUSI:\\n";
                        (d.jalannya_diskusi || []).forEach(j => text += "- " + j + "\\n");
                        text += "\\n✅ KEPUTUSAN:\\n";
                        (d.keputusan || []).forEach(k => text += "- " + k + "\\n");
                        text += "\\n📅 ACTION ITEMS (Tugas | PIC | Deadline | Prioritas):\\n";
                        (d.rencana_tindak_lanjut || []).forEach(t => { text += `- ${t.tugas} | ${t.pic} | ${t.deadline} | ${t.prioritas}\\n`; });

                        const blob = new Blob([text], { type: 'text/plain' });
                        const a = document.createElement('a');
                        a.href = URL.createObjectURL(blob);
                        a.download = 'Notulensi_Lengkap_' + Date.now() + '.txt';
                        a.click();
                    };

                    clearBtn.onclick = function() {
                        transcriptBox.innerHTML = '<div id="placeholder" style="text-align: center; color: #94a3b8; margin-top: 100px; font-weight: 600;">🎤 Klik "Start Capture"</div>';
                        lastFinalText = ""; currentInterimDiv = null; aiContent.innerHTML = "";
                        lastAiData = null;
                        downloadNotulensiBtn.style.display = 'none';
                        downloadDocxBtn.style.display = 'none';
                    };

                    function getTranscriptText() {
                        return Array.from(transcriptBox.querySelectorAll('.line-final')).map(line => line.innerText).join('\\n');
                    }

                    window.dlCyLive = function() {
                        if (window.cyInstance) {
                            const a = document.createElement('a'); 
                            a.href = window.cyInstance.png({full: true, scale: 4, bg: 'white'}); 
                            a.download = 'Cytoscape_Live.png'; a.click();
                        }
                    };

                    window.dlSunburstLive = function() {
                        if (window.sunburstChartLive) {
                            const a = document.createElement('a');
                            a.href = window.sunburstChartLive.getDataURL({ type: 'png', pixelRatio: 3, backgroundColor: '#ffffff' });
                            a.download = 'Sunburst_Live.png';
                            a.click();
                        }
                    };

                    // =========================================================
                    // FUNGSI SAKTI: EXPORT DOCX VIA JAVASCRIPT DENGAN 4 GAMBAR
                    // =========================================================
                    downloadDocxBtn.onclick = async function() {
                        if (!lastAiData) { alert('Belum ada data notulensi!'); return; }
                        
                        const originalBtnText = downloadDocxBtn.innerHTML;
                        downloadDocxBtn.innerHTML = "⏳ Memproses Gambar & Dokumen...";
                        downloadDocxBtn.disabled = true;

                        try {
                            const d = lastAiData.notulensi_rapat || {};
                            const now = new Date();
                            const tanggalStr = now.toLocaleDateString('en-GB', { weekday: 'long', day: '2-digit', month: 'long', year: 'numeric' });
                            const waktuStr = now.toLocaleTimeString('en-GB', { hour: '2-digit', minute: '2-digit' }) + ' WIB';
                            const tanggalFooterStr = now.toLocaleDateString('en-GB', { day: '2-digit', month: 'long', year: 'numeric' });
                            
                            // --- 1. CAPTURE CYTOSCAPE (NETWORK) ---
                            let cyImageHtml = "";
                            if (window.cyInstance) {
                                try {
                                    const cyBase64 = window.cyInstance.png({full: true, scale: 2, bg: 'white'});
                                    cyImageHtml = `
                                    <p style="font-size: 10pt; font-weight: bold; margin-bottom: 5px;">A. Cytoscape Network (Hubungan Topik)</p>
                                    <img src="${cyBase64}" style="width: 100%; max-width: 600px; height: auto; border: 1px solid #ccc; margin-bottom: 15px;"><br>`;
                                } catch (err) { console.error("Gagal screenshot Cytoscape:", err); }
                            }

                            // --- 2. CAPTURE MERMAID (MINDMAP 1) ---
                            let mermaidImageHtml = "";
                            const mermaidContainer = document.getElementById('merContainerLive'); 
                            if (mermaidContainer && mermaidContainer.querySelector('svg')) {
                                try {
                                    const svgMermaid = mermaidContainer.querySelector('svg');
                                    const origW = mermaidContainer.style.width;
                                    const origH = mermaidContainer.style.height;
                                    const origOverflow = mermaidContainer.style.overflow;
                                    
                                    const bboxMermaid = svgMermaid.getBBox();
                                    const padding = 40;
                                    const mWidth = Math.max(bboxMermaid.width + padding*2, mermaidContainer.clientWidth);
                                    const mHeight = Math.max(bboxMermaid.height + padding*2, mermaidContainer.clientHeight);

                                    mermaidContainer.style.width = mWidth + 'px';
                                    mermaidContainer.style.height = mHeight + 'px';
                                    mermaidContainer.style.overflow = 'visible';
                                    mermaidContainer.style.backgroundColor = '#ffffff';
                                    
                                    svgMermaid.style.width = mWidth + 'px';
                                    svgMermaid.style.height = mHeight + 'px';
                                    svgMermaid.style.maxWidth = 'none';

                                    const canvasMermaid = await html2canvas(mermaidContainer, { scale: 2, useCORS: true, backgroundColor: '#ffffff', width: mWidth, height: mHeight });
                                    const merBase64 = canvasMermaid.toDataURL('image/png');
                                    mermaidImageHtml = `
                                    <p style="font-size: 10pt; font-weight: bold; margin-bottom: 5px;">B. Mermaid (Mindmap Struktur)</p>
                                    <img src="${merBase64}" style="width: 100%; max-width: 600px; height: auto; border: 1px solid #ccc; margin-bottom: 15px;"><br>`;
                                    
                                    mermaidContainer.style.width = origW;
                                    mermaidContainer.style.height = origH;
                                    mermaidContainer.style.overflow = origOverflow;
                                    svgMermaid.style.width = '100%';
                                    svgMermaid.style.height = '100%';
                                } catch (err) { console.error("Gagal screenshot Mermaid:", err); }
                            }

                            // --- 3. CAPTURE MARKMAP (MINDMAP 2) ---
                            let markmapImageHtml = "";
                            const markmapContainer = document.getElementById('markmapLiveWrapper');
                            const markmapSvg = markmapContainer ? markmapContainer.querySelector('svg') : null;
                            if (markmapContainer && markmapSvg) {
                                try {
                                    const origW = markmapContainer.style.width;
                                    const origH = markmapContainer.style.height;
                                    const origOverflow = markmapContainer.style.overflow;
                                    const origViewBox = markmapSvg.getAttribute('viewBox');
                                    
                                    const g = markmapSvg.querySelector('g');
                                    const originalTransform = g ? g.getAttribute('transform') : null;
                                    if (g) g.setAttribute('transform', 'translate(50,50) scale(1)');
                                    
                                    await new Promise(resolve => setTimeout(resolve, 150));
                                    
                                    const bbox = g ? g.getBBox() : markmapSvg.getBBox();
                                    const padding = 60;
                                    const trueWidth = Math.max(bbox.width, 800) + (padding * 2);
                                    const trueHeight = Math.max(bbox.height, 600) + (padding * 2);

                                    markmapContainer.style.width = trueWidth + 'px';
                                    markmapContainer.style.height = trueHeight + 'px';
                                    markmapContainer.style.overflow = 'visible';
                                    markmapContainer.style.backgroundColor = '#ffffff';
                                    markmapSvg.setAttribute('viewBox', `${(bbox.x || 0) - padding} ${(bbox.y || 0) - padding} ${trueWidth} ${trueHeight}`);

                                    const canvasMarkmap = await html2canvas(markmapContainer, { 
                                        scale: 2, useCORS: true, backgroundColor: '#ffffff', width: trueWidth, height: trueHeight 
                                    });
                                    const mmBase64 = canvasMarkmap.toDataURL('image/png');
                                    markmapImageHtml = `
                                    <p style="font-size: 10pt; font-weight: bold; margin-bottom: 5px;">C. Markmap (Peta Konsep Detail)</p>
                                    <img src="${mmBase64}" style="width: 100%; max-width: 600px; height: auto; border: 1px solid #ccc; margin-bottom: 15px;"><br>`;

                                    markmapContainer.style.width = origW;
                                    markmapContainer.style.height = origH;
                                    markmapContainer.style.overflow = origOverflow;
                                    if (origViewBox) markmapSvg.setAttribute('viewBox', origViewBox);
                                    else markmapSvg.removeAttribute('viewBox');
                                    if (g && originalTransform) g.setAttribute('transform', originalTransform);
                                } catch (err) { console.error("Gagal screenshot Markmap:", err); }
                            }

                            // --- 4. CAPTURE SUNBURST (ECHARTS) ---
                            let sunburstImageHtml = "";
                            if (window.sunburstChartLive) {
                                try {
                                    const sbBase64 = window.sunburstChartLive.getDataURL({ type: 'png', pixelRatio: 3, backgroundColor: '#ffffff' });
                                    sunburstImageHtml = `
                                    <p style="font-size: 10pt; font-weight: bold; margin-bottom: 5px;">D. Sunburst Hierarchy (Anatomi Rapat)</p>
                                    <img src="${sbBase64}" style="width: 100%; max-width: 600px; height: auto; border: 1px solid #ccc; margin-bottom: 15px;"><br>`;
                                } catch (err) { console.error("Gagal screenshot Sunburst:", err); }
                            }

                            let actionItemsHtml = `<ul style="margin-top:0;"><li style="list-style: none;">- Tidak ada tindak lanjut khusus.</li></ul>`;
                            if (d.rencana_tindak_lanjut && d.rencana_tindak_lanjut.length > 0) {
                                actionItemsHtml = `
                                <table border="1" cellpadding="6" cellspacing="0" style="border-collapse: collapse; width: 100%; font-family: Arial, sans-serif; font-size: 11pt;">
                                    <tr style="background-color: #f1f5f9;">
                                        <th><b>Tugas</b></th><th><b>PIC</b></th><th><b>Deadline</b></th><th><b>Prioritas</b></th>
                                    </tr>
                                    ${d.rencana_tindak_lanjut.map(t => `
                                    <tr>
                                        <td>${t.tugas || '-'}</td><td>${t.pic || '-'}</td><td>${t.deadline || '-'}</td><td><b>${t.prioritas || '-'}</b></td>
                                    </tr>`).join('')}
                                </table>`;
                            }

                            let content = `
                            <html xmlns:o='urn:schemas-microsoft-com:office:office' xmlns:w='urn:schemas-microsoft-com:office:word' xmlns='http://www.w3.org/TR/REC-html40'>
                            <head><meta charset='utf-8'><title>Notulen Rapat</title></head>
                            <body style="font-family: Arial, sans-serif; font-size: 11pt; line-height: 1.5;">
                                
                                <table style="width: 100%; border: none; border-collapse: collapse;">
                                    <tr>
                                        <td style="width: 50%; text-align: left; vertical-align: top; border: none; padding: 0;">
                                            <span style="font-size: 28pt; font-weight: bold; color: #1e3a8a;">SMARTDOSE</span><br>
                                            <span style="font-size: 10pt; color: #64748b;">Enterprise AI Transcription<br>Healthcare & Productivity</span>
                                        </td>
                                        <td style="width: 50%; text-align: right; vertical-align: top; border: none; padding: 0;">
                                            <span style="font-size: 12pt; font-weight: bold; color: #1e3a8a;">SMARTDOSE ENTERPRISE<br>DIVISI INOVASI DIGITAL</span><br>
                                            <span style="font-size: 10pt; font-style: italic; color: #000000;">Sistem Notulensi Cerdas & Presisi</span>
                                        </td>
                                    </tr>
                                </table>
                                
                                <div style="text-align: center; margin-top: 5px; margin-bottom: 20px;">
                                    ________________________________________________________________________________
                                </div>

                                <div style="text-align: center; margin-bottom: 20px;">
                                    <span style="font-size: 16pt; font-weight: bold; color: #1e3a8a;">NOTULEN RAPAT</span><br><br>
                                    <span style="font-size: 12pt; font-weight: bold;">${d.agenda || 'Koordinasi dan Pembahasan Internal'}</span>
                                </div>
                                
                                <table border="1" cellpadding="5" cellspacing="0" style="border-collapse: collapse; width: 100%; margin-bottom: 20px; font-family: Arial, sans-serif; font-size: 11pt;">
                                    <tr>
                                        <td style="width: 1.5in;"><b>Hari / Tanggal</b></td>
                                        <td style="width: 4.5in;">${tanggalStr}</td>
                                    </tr>
                                    <tr>
                                        <td><b>Waktu</b></td>
                                        <td>${waktuStr}</td>
                                    </tr>
                                    <tr>
                                        <td><b>Media</b></td>
                                        <td>SmartDose TranscribX (Offline/Live)</td>
                                    </tr>
                                    <tr>
                                        <td><b>Notulis</b></td>
                                        <td>AI Transcription System</td>
                                    </tr>
                                    <tr>
                                        <td><b>Peserta</b></td>
                                        <td>${d.peserta ? (Array.isArray(d.peserta) ? d.peserta.join(', ') : d.peserta) : '-'}</td>
                                    </tr>
                                    <tr>
                                        <td><b>Agenda</b></td>
                                        <td>${d.agenda || '-'}</td>
                                    </tr>
                                </table>
                                
                                <p style="color: #1e3a8a; font-size: 11pt; margin-bottom: 5px;"><b>1. RINGKASAN EKSEKUTIF</b></p>
                                ${(lastAiData.ringkasan_eksekutif || []).map(r => `<p style="margin-top: 0; margin-bottom: 5px;">${r}</p>`).join('')}
                                <br>

                                <p style="color: #1e3a8a; font-size: 11pt; margin-bottom: 5px;"><b>2. POKOK PEMBAHASAN / JALANNYA DISKUSI</b></p>
                                <ul style="margin-top: 0;">
                                    ${(d.jalannya_diskusi || []).map(j => `<li style="margin-bottom: 5px;">${j}</li>`).join('')}
                                </ul>
                                <br>

                                <p style="color: #1e3a8a; font-size: 11pt; margin-bottom: 5px;"><b>3. KEPUTUSAN RAPAT</b></p>
                                <ol style="margin-top: 0;">
                                    ${(d.keputusan || []).map(k => `<li style="margin-bottom: 5px;">${k}</li>`).join('')}
                                </ol>
                                <br>

                                <p style="color: #1e3a8a; font-size: 11pt; margin-bottom: 5px;"><b>4. TINDAK LANJUT</b></p>
                                ${actionItemsHtml}
                                <br>
                                
                                <p style="color: #1e3a8a; font-size: 11pt; margin-bottom: 10px;"><b>5. LAMPIRAN VISUALISASI AI</b></p>
                                ${cyImageHtml}
                                ${mermaidImageHtml}
                                ${markmapImageHtml}
                                ${sunburstImageHtml}
                                
                                <br><br>
                                
                                <p style="text-align: right;">
                                    Jakarta, ${tanggalFooterStr}<br><br><br><br>
                                    ( Tim Notulen SmartDose )
                                </p>
                            </body>
                            </html>`;
                            
                            const blob = new Blob(['\\ufeff', content], { type: 'application/msword' });
                            const url = URL.createObjectURL(blob);
                            const a = document.createElement('a');
                            a.href = url;
                            a.download = 'Notulen_Live_SmartDose_' + Date.now() + '.doc';
                            document.body.appendChild(a);
                            a.click();
                            document.body.removeChild(a);

                        } catch (error) {
                            console.error("Error generating DOCX:", error);
                            alert("Terjadi kesalahan saat memproses gambar ke dalam dokumen.");
                        } finally {
                            downloadDocxBtn.innerHTML = originalBtnText;
                            downloadDocxBtn.disabled = false;
                        }
                    };

                    window.dlMermaidLive = function() {
                        const mDiv = document.getElementById('mermaidLive');
                        const container = document.getElementById('merContainerLive');
                        const svgEl = mDiv.querySelector('svg');
                        if (!svgEl) return;
                        
                        const btn = document.getElementById('dlBtnMermaidLive');
                        if (btn) { btn.innerHTML = "⏳..."; btn.disabled = true; }
                    
                        const bbox = svgEl.getBBox();
                        const padding = 40;
                        const trueWidth = Math.max(bbox.width, svgEl.clientWidth) + padding*2;
                        const trueHeight = Math.max(bbox.height, svgEl.clientHeight) + padding*2;
                    
                        const origSvgW = svgEl.style.width;
                        const origSvgH = svgEl.style.height;
                        const origSvgMaxW = svgEl.style.maxWidth;
                        const origDivCssText = mDiv.style.cssText;
                        const origClasses = mDiv.className;
                        const origContainerOverflow = container.style.overflow;
                    
                        mDiv.className = ''; 
                        mDiv.style.width = trueWidth + 'px';
                        mDiv.style.height = trueHeight + 'px';
                        mDiv.style.backgroundColor = '#ffffff';
                        mDiv.style.display = 'block'; 
                        
                        svgEl.style.width = trueWidth + 'px';
                        svgEl.style.height = trueHeight + 'px';
                        svgEl.style.maxWidth = 'none';
                        
                        container.style.overflow = 'visible';
                    
                        html2canvas(mDiv, { scale: 3, useCORS: true, backgroundColor: '#ffffff' })
                        .then(canvas => {
                            svgEl.style.width = origSvgW;
                            svgEl.style.height = origSvgH;
                            svgEl.style.maxWidth = origSvgMaxW;
                            
                            mDiv.className = origClasses;
                            mDiv.style.cssText = origDivCssText;
                            container.style.overflow = origContainerOverflow;
                            
                            const link = document.createElement('a'); 
                            link.download = 'Mermaid_Live.png'; 
                            link.href = canvas.toDataURL('image/png', 1.0); 
                            link.click();
                            if (btn) { btn.innerHTML = "📸 PNG"; btn.disabled = false; }
                        }).catch((e) => { 
                            console.error("Download Error:", e);
                            if (btn) { btn.innerHTML = "📸 PNG"; btn.disabled = false; } 
                        });
                    };
                    
                    window.dlMarkmapLive = function() {
                        const container = document.getElementById('markmapLiveWrapper');
                        const svgEl = container.querySelector('svg'); if (!svgEl) return;
                        
                        const btn = document.getElementById('dlBtnMMLive');
                        if (btn) { btn.innerHTML = "⏳..."; btn.disabled = true; }
                        
                        const originalWidth = container.style.width; 
                        const originalHeight = container.style.height; 
                        const originalOverflow = container.style.overflow;
                        
                        const g = svgEl.querySelector('g'); 
                        const originalTransform = g ? g.getAttribute('transform') : null;
                        if (g) g.setAttribute('transform', 'translate(50,50) scale(1)');
                        
                        setTimeout(() => {
                            const bbox = g ? g.getBBox() : svgEl.getBBox(); const padding = 60;
                            const trueWidth = Math.max(bbox.width, 800) + (padding * 2);
                            const trueHeight = Math.max(bbox.height, 600) + (padding * 2);
                            
                            container.style.width = trueWidth + 'px'; container.style.height = trueHeight + 'px'; container.style.overflow = 'visible';
                            svgEl.setAttribute('viewBox', `${(bbox.x || 0) - padding} ${(bbox.y || 0) - padding} ${trueWidth} ${trueHeight}`);
                            
                            html2canvas(container, { scale: 3, useCORS: true, backgroundColor: '#ffffff', width: trueWidth, height: trueHeight })
                            .then(canvas => {
                                container.style.width = originalWidth; container.style.height = originalHeight; container.style.overflow = originalOverflow;
                                if (g && originalTransform) g.setAttribute('transform', originalTransform);
                                
                                const link = document.createElement('a'); link.download = 'Markmap_Live.png';
                                link.href = canvas.toDataURL('image/png', 1.0); link.click();
                                if (btn) { btn.innerHTML = "📸 PNG HD"; btn.disabled = false; }
                            }).catch(() => { if (btn) { btn.innerHTML = "📸 PNG HD"; btn.disabled = false; } });
                        }, 500);
                    };

                    aiBtn.onclick = async function() {
                        const transcript = getTranscriptText(); const apiKey = apiKeyInput.value.trim();
                        if (!apiKey || !transcript) { alert('API Key atau Transkrip kosong!'); return; }
                        
                        aiContent.innerHTML = `
                            <div class="p-4 bg-yellow-50 text-yellow-700 rounded-xl mb-4 border border-yellow-200 animate-pulse">
                                <strong>⏳ AI sedang membaca transkrip...</strong>
                                <p style="margin:0; font-size:0.9em;">Sedang melakukan analisis mendalam dan menyusun visualisasi. Mohon tunggu sebentar.</p>
                            </div>
                        `;
                        
                        isThinking = true;
                        if (brainText) { brainText.innerText = "PROCESSING NEURAL DATA..."; brainText.style.color = "#d8b4fe"; }
                        aiBtn.innerHTML = '⏳ Memproses...'; aiBtn.disabled = true;

                        const prompt = `Anda adalah Ahli Pembuat Notulensi dan Visual Mapping. Analisis transkrip rapat berikut dan hasilkan JSON.
                        ATURAN STRUKTUR OUTPUT:
                        - ringkasan_eksekutif: Array of strings (poin-poin padat).
                        - transkrip_dialog: Array of strings (format: "Pembicara: Isi").
                        - jalannya_diskusi: Array of strings (Narasi detail & lengkap).
                        - keputusan: Array of strings.
                        - rencana_tindak_lanjut: Array of objects (tugas, pic, deadline, prioritas).
                        - hubungan_topik: Array of objects (sumber, target, relasi).
                        
                        ATURAN MERMAID (SANGAT KETAT):
                        1. Hasilkan flowchart berstruktur pohon dari kiri ke kanan dengan awalan 'graph LR'.
                        2. Konten/materinya HARUS SAMA DETAIL DAN BERCABANG seperti Markmap (Topik Utama -> Sub Topik -> Detail).
                        3. ID Node HARUS 1 HURUF/ANGKA saja tanpa spasi (misal: A, B, C1).
                        4. Teks label WAJIB DIAPIT TANDA KUTIP GANDA. Contoh: A["Teks Label Utama"] --> B["Teks Label Lain"]. DILARANG KERAS menggunakan kurung siku di dalam teks label itu sendiri, gunakan kurung biasa saja untuk singkatan.
                        
                        ATURAN MARKMAP (MUTLAK):
                        Hasilkan rancangan mindmap horizontal left-to-right tree yang sangat detail dan bercabang dalam menggunakan Markdown murni. 
                        Gunakan hierarki heading (# Topik Utama, ## Sub Topik, ### Detail Sub) dan bullet points (- Poin).
                        
                        Transkrip Rapat: "${transcript}"`;

                        const payload = {
                            model: "gemini/gemini-2.5-flash", messages: [{ role: "user", content: prompt }], temperature: 0.2,
                            response_format: {
                                type: "json_schema",
                                json_schema: {
                                    name: "meeting_summary", strict: true,
                                    schema: {
                                        type: "object",
                                        properties: {
                                            ringkasan_eksekutif: { type: "array", items: { type: "string" } },
                                            notulensi_rapat: {
                                                type: "object",
                                                properties: {
                                                    agenda: { type: "string" }, peserta: { type: "array", items: { type: "string" } },
                                                    jalannya_diskusi: { type: "array", items: { type: "string" } }, keputusan: { type: "array", items: { type: "string" } },
                                                    rencana_tindak_lanjut: { type: "array", items: { type: "object", properties: { tugas: { type: "string" }, pic: { type: "string" }, deadline: { type: "string" }, prioritas: { type: "string" } }, required: ["tugas", "pic", "deadline", "prioritas"], additionalProperties: false } },
                                                    hubungan_topik: { type: "array", items: { type: "object", properties: { sumber: { type: "string" }, target: { type: "string" }, relasi: { type: "string" } }, required: ["sumber", "target", "relasi"], additionalProperties: false } }
                                                }, required: ["agenda", "peserta", "jalannya_diskusi", "keputusan", "rencana_tindak_lanjut", "hubungan_topik"], additionalProperties: false
                                            },
                                            visual_mindmap: { type: "string" }, markmap_code: { type: "string" }
                                        }, required: ["ringkasan_eksekutif", "notulensi_rapat", "visual_mindmap", "markmap_code"], additionalProperties: false
                                    }
                                }
                            }
                        };

                        try {
                            const response = await fetch("https://litellm.koboi2026.biz.id/v1/chat/completions", {
                                method: "POST", 
                                headers: { 
                                    "Authorization": "Bearer " + apiKey, 
                                    "Content-Type": "application/json" 
                                },
                                body: JSON.stringify(payload)
                            });
                            const data = JSON.parse(JSON.parse(await response.text()).choices[0].message.content);
                            
                            lastAiData = data;
                            downloadNotulensiBtn.style.display = 'inline-flex';
                            downloadDocxBtn.style.display = 'inline-flex';
                            
                            let taskRows = (data.notulensi_rapat.rencana_tindak_lanjut || []).map(t => 
                                `<tr class="text-xs border-b"><td class="p-2 border-r">${t.tugas}</td><td class="p-2 border-r">${t.pic}</td><td class="p-2 border-r">${t.deadline}</td><td class="p-2 font-bold">${t.prioritas}</td></tr>`
                            ).join('');
                            
                            aiContent.innerHTML = `
                                <div class="fade-in mt-6 mb-10">
                                    <div class="mb-4"><strong>🌟 RINGKASAN EKSEKUTIF:</strong><div class="bg-blue-50 p-4 rounded-xl mt-2 text-sm"><ul class="list-disc ml-5">${(data.ringkasan_eksekutif || []).map(r => '<li>' + r + '</li>').join('')}</ul></div></div>
                                    <div class="mb-4"><strong>🗣️ JALANNYA DISKUSI:</strong><div class="bg-white p-4 rounded-xl border mt-2 text-sm"><ul>${(data.notulensi_rapat.jalannya_diskusi || []).map(d => '<li class="mb-2">- ' + d + '</li>').join('')}</ul></div></div>
                                    <div class="mb-4"><strong>✅ KEPUTUSAN UTAMA:</strong><ul class="list-disc ml-5 text-sm">${(data.notulensi_rapat.keputusan || []).map(k => '<li>' + k + '</li>').join('')}</ul></div>
                                    <div class="mb-8"><strong>📅 ACTION ITEMS:</strong><table class="w-full text-sm border mt-2"><thead class="bg-gray-100"><tr><th class="p-2 border-r">Tugas</th><th class="p-2 border-r">PIC</th><th class="p-2 border-r">Deadline</th><th class="p-2">Prioritas</th></tr></thead><tbody>${taskRows}</tbody></table></div>
                                    <h3 class="font-bold text-lg mb-4 border-b pb-2">🕸️ Visualisasi Map</h3>
                                    <div class="grid grid-cols-1 md:grid-cols-2 gap-4">
                                        <div><p class="font-bold text-sm mb-2">Cytoscape.js</p>
                                            <div class="relative bg-white border rounded-xl p-2">
                                                <button onclick="dlCyLive()" class="absolute top-2 right-2 z-10 bg-emerald-500 text-white px-3 py-1 rounded text-xs font-bold">📸 PNG</button>
                                                <div id="cyLiveContainer" style="width:100%; height:380px;"></div>
                                            </div>
                                        </div>
                                        <div><p class="font-bold text-sm mb-2">Mermaid (Mindmap)</p>
                                            <div class="relative bg-white border rounded-xl" style="height:396px;">
                                                <button id="dlBtnMermaidLive" onclick="dlMermaidLive()" class="absolute top-2 right-2 z-10 bg-emerald-500 text-white px-3 py-1 rounded text-xs font-bold">📸 PNG</button>
                                                <div id="merContainerLive" style="width:100%; height:100%; overflow:auto; border-radius:12px;">
                                                    <pre id="mermaidLive" class="mermaid w-full h-full m-0 flex justify-center items-center bg-white"></pre>
                                                </div>
                                            </div>
                                        </div>
                                    </div>
                                    
                                    <div class="mt-4"><p class="font-bold text-sm mb-2">🌿 Visualisasi Markmap (Peta Konsep Rapat)</p><div class="relative bg-white border rounded-xl overflow-hidden"><button id="dlBtnMMLive" onclick="dlMarkmapLive()" class="absolute top-4 right-4 z-10 bg-emerald-500 text-white px-3 py-1 rounded text-xs font-bold">📸 PNG HD</button><div id="markmapLiveWrapper" style="width:100%; height:500px;"><svg id="markmapLive" style="width:100%; height:100%;"></svg></div></div></div>
                                    
                                    <div class="mt-4">
                                        <p class="font-bold text-sm mb-2">☀️ Sunburst Hierarchy Chart (Anatomi Rapat)</p>
                                        <div class="relative bg-white border rounded-xl overflow-hidden p-2">
                                            <button onclick="dlSunburstLive()" class="absolute top-4 right-4 z-10 bg-emerald-500 text-white px-3 py-1 rounded text-xs font-bold">📸 PNG HD</button>
                                            <div id="sunburstLiveContainer" style="width:100%; height:600px;"></div>
                                        </div>
                                    </div>
                                </div>`;

                            // INIT CYTOSCAPE
                            setTimeout(() => {
                                const nodesSet = new Set(); const cyElements = [];
                                (data.notulensi_rapat.hubungan_topik || []).forEach(rel => {
                                    if (!nodesSet.has(rel.sumber)) { nodesSet.add(rel.sumber); cyElements.push({ data: { id: rel.sumber, label: rel.sumber } }); }
                                    if (!nodesSet.has(rel.target)) { nodesSet.add(rel.target); cyElements.push({ data: { id: rel.target, label: rel.target } }); }
                                    cyElements.push({ data: { source: rel.sumber, target: rel.target, label: rel.relasi } });
                                });
                                window.cyInstance = cytoscape({
                                    container: document.getElementById('cyLiveContainer'), elements: cyElements,
                                    style: [{ selector: 'node', style: { 'background-color': '#f43f5e', 'label': 'data(label)', 'color': '#1e293b', 'font-size': '11px', 'text-valign': 'top' } }, { selector: 'edge', style: { 'width': 2, 'line-color': '#cbd5e1', 'target-arrow-shape': 'triangle', 'label': 'data(label)', 'font-size': '9px' } }],
                                    layout: { name: 'cose', padding: 20 }
                                });
                            }, 100);

                            // INIT MERMAID
                            setTimeout(() => {
                                let rawMer = (data.visual_mindmap || "").replace(/```mermaid/gi, "").replace(/```/g, "").trim();
                                if (!rawMer.toLowerCase().includes('graph') && !rawMer.toLowerCase().includes('flowchart') && !rawMer.toLowerCase().includes('mindmap')) {
                                    rawMer = `graph LR\\n` + rawMer;
                                }
                                rawMer = rawMer.replace(/`/g, "").replace(/\\[([A-Z0-9]+)\\]/g, "($1)");
                                
                                const mDiv = document.getElementById('mermaidLive'); 
                                mDiv.textContent = rawMer; 
                                mDiv.removeAttribute('data-processed');
                                
                                try {
                                    mermaid.run({ querySelector: '#mermaidLive' }).then(() => {
                                        const svg = mDiv.querySelector('svg');
                                        if (svg) { 
                                            svg.style.maxWidth = 'none'; 
                                            svg.style.height = 'auto';
                                        }
                                    }).catch(e => {
                                        console.error("Mermaid error:", e);
                                        mDiv.innerHTML = "<div style='color:red; padding:20px;'>Gagal render Mermaid. Transkrip mungkin mengandung karakter ilegal.</div>";
                                    });
                                } catch(e) {}
                            }, 100);

                            // INIT MARKMAP
                            setTimeout(() => {
                                let rawMm = (data.markmap_code || "").replace(/```markdown/gi, "").replace(/```/g, "").trim();
                                const { Transformer, Markmap } = window.markmap;
                                const { root } = new Transformer().transform(rawMm);
                                Markmap.create('#markmapLive', null, root);
                            }, 100);

                            // INIT SUNBURST ECHARTS
                            setTimeout(() => {
                                let rawMm = (data.markmap_code || "").replace(/```markdown/gi, "").replace(/```/g, "").trim();
                                
                             function parseMarkdownToSunburst(md) {
                                        const lines = md.split('\\n');
                                        let root = { name: "Root", children: [] };
                                        let stack = [ {level: -1, node: root} ];

                                        for (let i = 0; i < lines.length; i++) {
                                            let line = lines[i];
                                            if (!line.trim()) continue;

                                            let level = 0;
                                            let text = "";

                                            let matchHeader = line.match(/^(\\s*)(#+)\\s+(.*)/);
                                            let matchList = line.match(/^(\\s*)[-*]\\s+(.*)/);

                                            if (matchHeader) {
                                                level = matchHeader[2].length;
                                                text = matchHeader[3];
                                            } else if (matchList) {
                                                let indent = matchList[1].length;
                                                level = 100 + indent;
                                                text = matchList[2];
                                            } else {
                                                continue;
                                            }

                                            // Bersihkan sisa format markdown
                                            text = text.replace(/\\*\\*/g, '').replace(/_/g, '').trim();

                                            let newNode = { name: text, children: [] };

                                            while (stack.length > 1) {
                                                let topLevel = stack[stack.length - 1].level;
                                                if (level < 100) {
                                                    if (topLevel >= 100 || topLevel >= level) stack.pop();
                                                    else break;
                                                } else {
                                                    if (topLevel >= 100 && topLevel >= level) stack.pop();
                                                    else break;
                                                }
                                            }

                                            let parent = stack[stack.length - 1].node;
                                            parent.children.push(newNode);
                                            stack.push({ level: level, node: newNode });
                                        }

                                        function assignValues(node) {
                                            if (node.children.length === 0) {
                                                node.value = 1;
                                            } else {
                                                node.children.forEach(assignValues);
                                            }
                                        }
                                        assignValues(root);
                                        
                                        // PINTAR MEMILIH ROOT: Agar Topik Utama selalu ngisi di tengah
                                        if (root.children.length === 1) {
                                            return root.children; // Jika cuma 1 topik utama, jadikan dia pusatnya
                                        } else if (root.children.length > 1) {
                                            return [{ name: "Anatomi Rapat", children: root.children }]; // Bungkus jika banyak
                                        }
                                        return [{name: "Data tidak tersedia", value: 1}];
                                    }

                                    let sunburstData = parseMarkdownToSunburst(rawMm);
                                    
                                    const colorPalette = ['#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6', '#14b8a6', '#f43f5e', '#84cc16', '#0ea5e9', '#d946ef'];

                                    // FUNGSI REKURSIF: Memaksa tiap cabang (branch) beda warna, anti "hijau semua"
                                    function applyColorToBranch(node, color) {
                                        node.itemStyle = node.itemStyle || {};
                                        node.itemStyle.color = color;
                                        if (node.children && node.children.length > 0) {
                                            node.children.forEach(child => applyColorToBranch(child, color));
                                        }
                                    }

                                    // FORMATTING PUSAT (ROOT) DAN ANAK-ANAKNYA
                                    if (sunburstData.length === 1) {
                                        sunburstData[0].itemStyle = { color: '#0f172a' }; // Pusat warna Navy Gelap
                                        sunburstData[0].label = { color: '#ffffff', fontSize: 13, fontWeight: 'bold' };
                                        
                                        if (sunburstData[0].children) {
                                            sunburstData[0].children.forEach((child, index) => { 
                                                let branchColor = colorPalette[index % colorPalette.length];
                                                applyColorToBranch(child, branchColor);
                                            });
                                        }
                                    }

                                    var chartDom = document.getElementById('sunburstLiveContainer');
                                    window.sunburstChartLive = echarts.init(chartDom);
                                    
                                    var option = {
                                        tooltip: { trigger: 'item', formatter: function(info) { return '<div style="max-width:300px; white-space:normal; font-size:13px;">' + info.name + '</div>'; } },
                                        series: {
                                            type: 'sunburst', 
                                            data: sunburstData, 
                                            radius: [0, '95%'], // '0' MEMBUAT ROOT MENGISI PENUH DI TENGAH (TIDAK BOLONG)
                                            sort: undefined, 
                                            emphasis: { focus: 'ancestor' },
                                            itemStyle: { borderRadius: 4, borderWidth: 1.5, borderColor: '#ffffff' },
                                            label: { 
                                                show: true, 
                                                formatter: '{b}', 
                                                minAngle: 4, 
                                                width: 80,  
                                                overflow: 'break', // Teks akan digulung ke bawah secara proporsional
                                                fontSize: 9.5, 
                                                fontWeight: 'bold', 
                                                fontFamily: 'sans-serif', 
                                                color: '#ffffff', 
                                                textBorderColor: 'rgba(0,0,0,0.6)', 
                                                textBorderWidth: 1.5 
                                            }
                                        }
                                    };
                                    window.sunburstChartLive.setOption(option);
                                    window.addEventListener('resize', function() { window.sunburstChartLive.resize(); });
                                }, 100);

                            } catch(err) { aiContent.innerHTML = '<div class="p-4 bg-red-50 text-red-600 rounded-xl mt-4">Gagal memproses data AI: ' + err.message + '</div>'; }
                            finally { aiBtn.innerHTML = '✨ Generate AI Summary'; aiBtn.disabled = false; isThinking = false; if (brainText) brainText.innerText = "NEURAL NETWORK IDLE"; }
                        };
                    })();
            </script>
        </body>
        </html>
        """
        components.html(html_code, height=2200, scrolling=True)
    # =====================================================================
    # TAB 2: FITUR OFFLINE TRANSCRIPTION
    # =====================================================================
    with tab2:
        st.markdown("### 📁 Transkripsi File Rekaman (Offline)")
        if not is_admin() and st.session_state.get("user_paket", "NON-AKTIF") == "NON-AKTIF":
            st.markdown("""<div style="background: rgba(239, 68, 68, 0.1); border: 2px solid #ef4444; border-radius: 12px; padding: 30px; text-align: center; margin-top: 20px;"><h1 style="font-size: 50px; margin: 0;">🔒</h1><h2 style="color: #ef4444; margin-top: 10px;">AKSES DITOLAK</h2><p style="color: #475569; font-size: 16px;">Fitur Upload Rekaman Audio ini hanya tersedia untuk pengguna aktif.</p></div>""", unsafe_allow_html=True)
        else:
            st.info("💡 Sistem ini menggunakan **LiteLLM Proxy** untuk proses Transkripsi (Whisper) sekaligus Summarization (Gemini). Menggunakan mesin FFmpeg native.")
            llm_key = st.text_input("🔑 API Key LiteLLM (All-in-One)", type="password", placeholder="sk-...", help="API Key untuk proxy LiteLLM Anda")
            uploaded_file = st.file_uploader("Upload File Rekaman Anda", type=["mp3", "wav", "m4a", "mp4"])

            if uploaded_file is not None:
                st.audio(uploaded_file)
                kuota_upload_sekarang = st.session_state.get("user_kuota_upload", 0)
                is_allowed_to_upload = is_admin() or kuota_upload_sekarang > 0

                if not is_allowed_to_upload: st.error("❌ Kuota Upload Anda telah habis. Silakan hubungi Admin atau top-up.")
                else:
                    if st.button("🎙️ Mulai Transkripsi (Smart Chunking FFmpeg)", use_container_width=True, type="primary"):
                        if not llm_key: st.warning("⚠️ Masukkan API Key LiteLLM terlebih dahulu!")
                        else:
                            if not is_admin():
                                st.session_state["user_kuota_upload"] -= 1
                                db.collection("users").document(st.session_state["user_uid"]).update({"kuota_upload": st.session_state["user_kuota_upload"]})
                                render_sidebar_profile()

                            with st.spinner("⏳ Memproses file secara efisien tanpa membebani RAM..."):
                                temp_dir = tempfile.mkdtemp()
                                input_path = None
                                try:
                                    import subprocess, glob, shutil
                                    file_extension = uploaded_file.name.split('.')[-1]
                                    input_path = os.path.join(temp_dir, f"input_audio.{file_extension}")
                                    with open(input_path, "wb") as f: f.write(uploaded_file.getvalue())
                                    
                                    status_text = st.empty()
                                    status_text.info("✂️ Memotong audio menjadi bagian-bagian kecil...")
                                    output_pattern = os.path.join(temp_dir, "chunk_%03d.mp3")
                                    subprocess.run(["ffmpeg", "-i", input_path, "-f", "segment", "-segment_time", "600", "-c:a", "libmp3lame", "-b:a", "64k", output_pattern], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                                    
                                    chunk_files = sorted(glob.glob(os.path.join(temp_dir, "chunk_*.mp3")))
                                    total_chunks = len(chunk_files)
                                    
                                    if total_chunks == 0: st.error("Gagal memotong audio.")
                                    else:
                                        full_transcript = ""
                                        progress_bar = st.progress(0)
                                        url = "https://litellm.koboi2026.biz.id/v1/audio/transcriptions"
                                        headers = {"Authorization": f"Bearer {llm_key}"}
                                        success_transcription = True
                                        all_segments = []
                                        for i, chunk_file in enumerate(chunk_files):
                                            status_text.markdown(f"**🔄 Mentranskripsi bagian {i+1} dari {total_chunks}...**")
                                            with open(chunk_file, "rb") as f: files = {"file": (os.path.basename(chunk_file), f.read(), "audio/mpeg")}
                                            # UBAH response_format ke verbose_json untuk memunculkan timestamp Time-Travel
                                            response = requests.post(url, headers=headers, files=files, data={"model": "whisper-1", "response_format": "verbose_json"})
                                            if response.status_code == 200: 
                                                res_data = response.json()
                                                full_transcript += res_data.get("text", "") + " "
                                                # Kalkulasi offset waktu tiap potongan (600 detik per chunk)
                                                offset = i * 600.0
                                                for seg in res_data.get("segments", []):
                                                    all_segments.append({
                                                        "start": seg.get("start", 0) + offset,
                                                        "end": seg.get("end", 0) + offset,
                                                        "text": seg.get("text", "")
                                                    })
                                            else: st.error(f"❌ Error API LiteLLM chunk {i+1}: {response.text}"); success_transcription = False; break
                                            progress_bar.progress((i + 1) / total_chunks)
                                        
                                        if success_transcription and full_transcript.strip():
                                            status_text.success("✅ Seluruh audio berhasil ditranskrip!")
                                            st.session_state["offline_transcript"] = full_transcript.strip()
                                            st.session_state["offline_segments"] = all_segments
                                            st.session_state["offline_audio_b64"] = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
                                            st.session_state["offline_audio_mime"] = uploaded_file.type
                                except subprocess.CalledProcessError as e: st.error(f"❌ Error FFmpeg: {e.stderr.decode('utf-8')}")
                                except Exception as e: st.error(f"Terjadi kesalahan: {str(e)}")
                                finally:
                                    if os.path.exists(temp_dir): shutil.rmtree(temp_dir, ignore_errors=True)

            if st.session_state["offline_transcript"]:
                st.markdown("#### 📝 Hasil Transkripsi")
                
                # --- TAMBAHAN UI TIME-TRAVEL MENGGANTIKAN TEXT-AREA BIASA ---
                if st.session_state.get("offline_segments") and st.session_state.get("offline_audio_b64"):
                    audio_uri = f"data:{st.session_state['offline_audio_mime']};base64,{st.session_state['offline_audio_b64']}"
                    
                    segments_html = ""
                    for seg in st.session_state["offline_segments"]:
                        total_sec = int(seg['start'])
                        m, s = divmod(total_sec, 60)
                        time_str = f"{m:02d}:{s:02d}"
                        segments_html += f'<div class="segment" data-start="{seg["start"]}" data-end="{seg["end"]}"><span class="timestamp">{time_str}</span> {seg["text"]}</div>'

                    time_travel_html = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <style>
                            body {{ font-family: -apple-system, BlinkMacSystemFont, sans-serif; background: transparent; margin: 0; padding: 5px; }}
                            .player-container {{ background: white; padding: 20px; border-radius: 16px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); border: 1px solid #e2e8f0; }}
                            audio {{ width: 100%; border-radius: 8px; margin-bottom: 20px; outline: none; }}
                            .transcript-box {{ height: 350px; overflow-y: auto; padding: 10px; border: 1px solid #cbd5e1; border-radius: 12px; background: #ffffff; scroll-behavior: smooth; }}
                            .segment {{ padding: 12px 15px; margin-bottom: 8px; border-radius: 8px; cursor: pointer; transition: all 0.2s ease; border-left: 4px solid transparent; font-size: 15px; color: #475569; }}
                            .segment:hover {{ background: #f1f5f9; }}
                            .segment.active {{ background: #eff6ff; border-left: 4px solid #3b82f6; color: #1e3a8a; font-weight: 600; box-shadow: 0 2px 4px rgba(59, 130, 246, 0.1); }}
                            .timestamp {{ font-size: 12px; color: #94a3b8; font-weight: bold; margin-right: 10px; background: #f1f5f9; padding: 2px 6px; border-radius: 4px; display: inline-block; }}
                        </style>
                    </head>
                    <body>
                    <div class="player-container">
                        <h3 style="margin-top:0; color:#1e293b; font-size:16px;">🎧 Interactive Timeline Audio (Time-Travel)</h3>
                        <audio id="audioPlayer" controls>
                            <source src="{audio_uri}" type="{st.session_state['offline_audio_mime']}">
                        </audio>
                        <div class="transcript-box" id="transcriptBox">
                            {segments_html}
                        </div>
                    </div>
                    <script>
                        const player = document.getElementById('audioPlayer');
                        const segments = document.querySelectorAll('.segment');
                        
                        segments.forEach(seg => {{
                            seg.addEventListener('click', () => {{
                                player.currentTime = parseFloat(seg.getAttribute('data-start'));
                                player.play();
                            }});
                        }});

                        player.addEventListener('timeupdate', () => {{
                            const currentTime = player.currentTime;
                            segments.forEach(seg => {{
                                const start = parseFloat(seg.getAttribute('data-start'));
                                const end = parseFloat(seg.getAttribute('data-end'));
                                if (currentTime >= start && currentTime <= end) {{
                                    if (!seg.classList.contains('active')) {{
                                        segments.forEach(s => s.classList.remove('active'));
                                        seg.classList.add('active');
                                        seg.scrollIntoView({{ behavior: 'smooth', block: 'center' }});
                                    }}
                                }}
                            }});
                        }});
                    </script>
                    </body>
                    </html>
                    """
                    components.html(time_travel_html, height=520, scrolling=True)
                # --- END UI TIME-TRAVEL ---

                st.session_state["offline_transcript"] = st.text_area("Edit Raw Text (Hanya jika diperlukan untuk Summary AI):", value=st.session_state["offline_transcript"], height=150)
                if st.button("✨ Generate AI Summary dari Teks Ini", use_container_width=True, type="secondary"):
                    if not llm_key: st.warning("⚠️ Masukkan API Key LiteLLM!")
                    elif not is_admin() and st.session_state.get("user_kuota_ai", 0) <= 0: st.error("❌ Kuota AI Summary habis! Silakan lakukan Top-Up.")
                    else:
                        if not is_admin():
                            st.session_state["user_kuota_ai"] -= 1
                            db.collection("users").document(st.session_state["user_uid"]).update({"kuota_ai": st.session_state["user_kuota_ai"]})
                            render_sidebar_profile()

                        with st.spinner("⏳ Tahap 1/2: AI sedang menyusun Ringkasan & Action Items..."):
                            prompt1 = f"""Anda adalah Ahli Pembuat Notulensi. Analisis transkrip rapat berikut dan WAJIB kembalikan output HANYA dalam format JSON. JANGAN menyalin ulang atau membuat transkrip dialog penuh.
                            STRUKTUR JSON: {{ "ringkasan_eksekutif": ["..."], "notulensi_rapat": {{ "agenda": "...", "peserta": ["..."], "jalannya_diskusi": ["..."], "keputusan": ["..."], "rencana_tindak_lanjut": [{{"tugas": "...", "pic": "...", "deadline": "...", "prioritas": "..."}}] }} }}
                            Transkrip: "{st.session_state['offline_transcript']}" """
                            res1 = None
                            try:
                                res1 = requests.post("https://litellm.koboi2026.biz.id/v1/chat/completions", headers={"Authorization": f"Bearer {llm_key}", "Content-Type": "application/json"}, json={"model":"gemini/gemini-2.5-flash", "messages": [{ "role": "user", "content": prompt1 }], "temperature": 0.2, "response_format": { "type": "json_object" }})
                                if res1.status_code != 200: st.error(f"Error AI (Tahap 1): {res1.status_code}"); res1 = None
                            except Exception as e: st.error(f"Koneksi LLM Gagal (Tahap 1): {str(e)}")

                        if res1:
                            with st.spinner("⏳ Tahap 2/2: AI sedang merancang Peta Konsep..."):
                                prompt2 = f"""Anda Ahli Visual Mapping. Buat rancangan JSON untuk visualisasi berdasarkan transkrip rapat. WAJIB HANYA JSON.
                                ATURAN MERMAID: 1. Hasilkan flowchart berstruktur pohon (graph LR). 2. ID Node HARUS 1 HURUF/ANGKA saja tanpa spasi. 3. Teks label WAJIB DIAPIT TANDA KUTIP GANDA.
                                STRUKTUR JSON: {{ "hubungan_topik": [{{"sumber": "...", "target": "...", "relasi": "..."}}], "visual_mindmap": "graph LR\\nA[\"Topik Utama\"] --> B[\"Sub Topik\"]", "markmap_code": "# Topik Utama\\n## Sub Topik" }}
                                Transkrip: "{st.session_state['offline_transcript']}" """
                                try:
                                    res2 = requests.post("https://litellm.koboi2026.biz.id/v1/chat/completions", headers={"Authorization": f"Bearer {llm_key}", "Content-Type": "application/json"}, json={"model":"gemini/gemini-2.5-flash", "messages": [{ "role": "user", "content": prompt2 }], "temperature": 0.2, "response_format": { "type": "json_object" }})
                                    if res2.status_code == 200:
                                        data_teks = json.loads(res1.json()["choices"][0]["message"]["content"])
                                        data_visual = json.loads(res2.json()["choices"][0]["message"]["content"])
                                        if "notulensi_rapat" not in data_teks: data_teks["notulensi_rapat"] = {}
                                        data_teks["visual_mindmap"] = data_visual.get("visual_mindmap", "")
                                        data_teks["markmap_code"] = data_visual.get("markmap_code", "")
                                        data_teks["notulensi_rapat"]["hubungan_topik"] = data_visual.get("hubungan_topik", [])
                                        st.session_state["offline_summary"] = data_teks
                                        st.success("✅ Analisis AI Lengkap & Selesai!")
                                    else: st.error(f"Error AI (Tahap 2): {res2.status_code}")
                                except Exception as e: st.error(f"Koneksi LLM Gagal (Tahap 2): {str(e)}")

            if st.session_state.get("offline_summary"):
                data = st.session_state["offline_summary"]
                st.markdown("---")
                col_t1, col_t2 = st.columns([3, 1])
                with col_t1: st.markdown("### 📋 Laporan Notulensi AI")
                
                txt_report = "NOTULENSI RAPAT SMARTDOSE ENTERPRISE\n========================================\n\n🌟 RINGKASAN EKSEKUTIF:\n"
                for r in data.get('ringkasan_eksekutif', []): txt_report += f"- {r}\n"
                notulensi = data.get('notulensi_rapat', {})
                txt_report += f"\n📌 AGENDA: {notulensi.get('agenda', '-')}\n👥 PESERTA: {', '.join(notulensi.get('peserta', [])) if isinstance(notulensi.get('peserta'), list) else notulensi.get('peserta', '-')}\n\n🗣️ JALANNYA DISKUSI:\n"
                for d in notulensi.get('jalannya_diskusi', []): txt_report += f"- {d}\n"
                txt_report += "\n✅ KEPUTUSAN:\n"
                for k in notulensi.get('keputusan', []): txt_report += f"- {k}\n"
                txt_report += "\n📅 ACTION ITEMS (Tugas | PIC | Deadline | Prioritas):\n"
                for t in notulensi.get('rencana_tindak_lanjut', []): txt_report += f"- {t.get('tugas', '-')} | {t.get('pic', '-')} | {t.get('deadline', '-')} | {t.get('prioritas', '-')}\n"

                with col_t2: 
                    st.download_button("📝 Download Laporan (TXT)", data=txt_report, file_name=f"Notulensi_{datetime.now().strftime('%Y%m%d')}.txt", mime="text/plain", use_container_width=True)
                    try: st.download_button("📄 Download Resmi (DOCX)", data=generate_notulensi_docx(data), file_name=f"Notulen_SmartDose_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True, type="primary")
                    except Exception as e: st.error(f"Gagal memuat DOCX: {e}")

                with st.container(border=True):
                    st.markdown("**🌟 RINGKASAN EKSEKUTIF:**")
                    st.markdown("<div style='background-color:#eff6ff; padding:15px; border-radius:10px; color:#1e3a8a; font-weight:bold; margin-bottom:15px;'><ul style='margin:0; padding-left:20px; line-height:1.6;'>" + "".join([f"<li>{r}</li>" for r in data.get('ringkasan_eksekutif', [])]) + "</ul></div>", unsafe_allow_html=True)
                    colA, colB = st.columns(2)
                    colA.markdown(f"**📌 AGENDA:** {data['notulensi_rapat'].get('agenda', '-')}")
                    colB.markdown(f"**👥 PESERTA:** {', '.join(data['notulensi_rapat'].get('peserta', []))}")
                    if 'transkrip_dialog' in data['notulensi_rapat']:
                        st.markdown("**💬 TRANSKRIP DIALOG:**")
                        dialog_html = "<div style='background-color:#f8fafc; padding:20px; border-radius:15px; border:1px solid #cbd5e1; max-height:350px; overflow-y:auto;'>"
                        for line in data['notulensi_rapat'].get('transkrip_dialog', []):
                            if ":" in line:
                                spk, txt = line.split(":", 1)
                                dialog_html += f"<div style='margin-bottom:12px;'><span style='font-size:12px; font-weight:bold; color:#475569;'>{spk.strip()}</span><div style='background-color:#e0f2fe; padding:10px 14px; border-radius:0 12px 12px 12px; font-size:14px; color:#1e293b; margin-top:2px;'>{txt.strip()}</div></div>"
                            else: dialog_html += f"<div style='margin-bottom:8px; font-style:italic; color:#64748b; font-size:13px;'>{line}</div>"
                        st.markdown(dialog_html + "</div>", unsafe_allow_html=True)
                    st.markdown("**🗣️ JALANNYA DISKUSI:**")
                    st.markdown("<div style='background-color:#fff; padding:15px; border-radius:10px; border:1px solid #e2e8f0; margin-bottom:15px;'><ul>" + "".join([f"<li style='margin-bottom:6px;'>- {d}</li>" for d in data['notulensi_rapat'].get('jalannya_diskusi', [])]) + "</ul></div>", unsafe_allow_html=True)
                    st.markdown("**📅 ACTION ITEMS:**")
                    st.table(pd.DataFrame(data['notulensi_rapat'].get('rencana_tindak_lanjut', [])))

               # Ambil raw string Mermaid
                clean_mer = data.get('visual_mindmap', '').replace("```mermaid", "").replace("```", "").strip()
                
                # Pastikan diawali dengan tipe grafiknya
                if not clean_mer.lower().startswith('graph') and not clean_mer.lower().startswith('flowchart') and not clean_mer.lower().startswith('mindmap'): 
                    clean_mer = "graph LR\n" + clean_mer
                
                # PENTING: JANGAN hapus double quote (") karena Mermaid butuh double quote untuk string berisikan spasi / tanda kurung
                clean_mer = clean_mer.replace('`', '').replace("'", "")
                
                # Tambahkan proteksi Regex seperti di Live Transkrip (opsional agar makin tangguh)
                clean_mer = re.sub(r'\[([A-Z0-9]+)\]', r'(\1)', clean_mer)

                mer_json_str = json.dumps(clean_mer)
                markmap_json_str = json.dumps(data.get('markmap_code', '').replace("```markdown", "").replace("```", "").strip())
                hubungan_json = json.dumps(data['notulensi_rapat'].get('hubungan_topik', []))
                
                st.markdown("### 🕸️ Visualisasi Terintegrasi")
                col_v1, col_v2 = st.columns(2)
                with col_v1:
                    st.markdown("**Cytoscape.js Network**")
                    components.html(f"""<!DOCTYPE html><html><head><script src="https://cdnjs.cloudflare.com/ajax/libs/cytoscape/3.26.0/cytoscape.min.js"></script></head><body style="margin:0; padding:10px; background:#f8fafc; position:relative;"><button onclick="dlCy()" style="position:absolute; top:20px; right:20px; z-index:100; background:#10b981; color:white; border:none; padding:5px 10px; border-radius:5px; cursor:pointer; font-weight:bold;">📸 PNG Full</button><div id="cy" style="width:100%; height:400px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px;"></div><script>const rawData = {hubungan_json}; const cyElements = []; const nodesSet = new Set(); rawData.forEach(rel => {{ if (!nodesSet.has(rel.sumber)) {{ nodesSet.add(rel.sumber); cyElements.push({{ data: {{ id: rel.sumber, label: rel.sumber }} }}); }} if (!nodesSet.has(rel.target)) {{ nodesSet.add(rel.target); cyElements.push({{ data: {{ id: rel.target, label: rel.target }} }}); }} cyElements.push({{ data: {{ source: rel.sumber, target: rel.target, label: rel.relasi }} }}); }}); var cy = cytoscape({{ container: document.getElementById('cy'), elements: cyElements, style: [ {{ selector: 'node', style: {{ 'background-color': '#f43f5e', 'label': 'data(label)', 'color': '#1e293b', 'font-size': '12px', 'text-valign': 'top' }} }}, {{ selector: 'edge', style: {{ 'width': 2, 'line-color': '#cbd5e1', 'target-arrow-shape': 'triangle', 'label': 'data(label)', 'font-size': '10px' }} }} ], layout: {{ name: 'cose' }} }}); function dlCy() {{ const a = document.createElement('a'); a.href = cy.png({{full: true, scale: 4, bg: 'white'}}); a.download = 'Cytoscape.png'; a.click(); }}</script></body></html>""", height=450)
                with col_v2:
                    st.markdown("**Mermaid (Mindmap)**")
                    components.html(f"""<!DOCTYPE html><html><head><script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script><script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js"></script></head><body style="margin:0; padding:10px; background:#f8fafc; position:relative;"><button id="dlBtn" onclick="downloadMermaidImage('merContainer', 'Offline')" style="position:absolute; top:20px; right:20px; z-index:100; background:#10b981; color:white; border:none; padding:5px 10px; border-radius:5px; cursor:pointer; font-weight:bold;">📸 PNG</button><div style="width:100%; height:400px; border:1px solid #e2e8f0; border-radius:8px; background:#ffffff; position:relative;"><div id="merContainerWrap" style="width:100%; height:100%; overflow:auto; border-radius:8px;"><pre class="mermaid" id="merContainer" style="background:transparent; margin:0; display:flex; justify-content:center; align-items:center;"></pre></div></div><script>document.getElementById('merContainer').textContent = {mer_json_str}; mermaid.initialize({{ startOnLoad: false, theme: 'default' }}); try {{ mermaid.run({{ querySelector: '.mermaid' }}).then(() => {{ const svgEl = document.querySelector('.mermaid svg'); if (svgEl) {{ svgEl.style.maxWidth = 'none'; svgEl.style.height = 'auto'; }} }}).catch(e => {{ document.getElementById('merContainerWrap').innerHTML = "<div style='padding:20px; font-size:12px; background:#fee2e2; color:#991b1b; border-radius:8px;'><b style='font-size:14px;'>⚠️ Format gambar dari AI kurang sempurna.</b><br>Berikut teks struktur aslinya agar data tidak hilang:<br><br><pre style='white-space:pre-wrap; font-family:monospace;'>" + {mer_json_str} + "</pre></div>"; }}); }} catch(e) {{}} window.downloadMermaidImage = function(wrapperId, title) {{ const mDiv = document.getElementById(wrapperId); const container = document.getElementById('merContainerWrap'); const svgEl = mDiv.querySelector('svg'); if (!svgEl) return; const btn = document.getElementById('dlBtn'); const originalText = btn.innerHTML; btn.innerHTML = "⏳..."; btn.disabled = true; setTimeout(() => {{ const bbox = svgEl.getBBox(); const padding = 40; const trueWidth = Math.max(bbox.width, svgEl.clientWidth) + padding*2; const trueHeight = Math.max(bbox.height, svgEl.clientHeight) + padding*2; const origSvgW = svgEl.style.width; const origSvgH = svgEl.style.height; const origSvgMaxW = svgEl.style.maxWidth; const origDivCssText = mDiv.style.cssText; const origContainerOverflow = container ? container.style.overflow : ''; mDiv.style.width = trueWidth + 'px'; mDiv.style.height = trueHeight + 'px'; mDiv.style.display = 'block'; mDiv.style.backgroundColor = '#ffffff'; mDiv.style.margin = '0 auto'; svgEl.style.width = trueWidth + 'px'; svgEl.style.height = trueHeight + 'px'; svgEl.style.maxWidth = 'none'; if (container) container.style.overflow = 'visible'; html2canvas(mDiv, {{ scale: 3, useCORS: true, backgroundColor: '#ffffff' }}).then(canvas => {{ svgEl.style.width = origSvgW; svgEl.style.height = origSvgH; svgEl.style.maxWidth = origSvgMaxW; mDiv.style.cssText = origDivCssText; if (container) container.style.overflow = origContainerOverflow; const link = document.createElement('a'); link.download = 'Mermaid_' + title + '.png'; link.href = canvas.toDataURL('image/png', 1.0); link.click(); btn.innerHTML = originalText; btn.disabled = false; }}).catch(() => {{ btn.innerHTML = originalText; btn.disabled = false; }}); }}, 500); }};</script></body></html>""", height=450)

                st.markdown("### 🌿 Visualisasi Markmap (Peta Konsep Rapat Horizontal)")
                components.html(f"""<!DOCTYPE html><html><head><script src="https://cdn.jsdelivr.net/npm/d3@7/dist/d3.min.js"></script><script src="https://cdn.jsdelivr.net/npm/markmap-lib@0.15.4/dist/browser/index.js"></script><script src="https://cdn.jsdelivr.net/npm/markmap-view@0.15.4/dist/browser/index.js"></script><script src="https://cdnjs.cloudflare.com/ajax/libs/html2canvas/1.4.1/html2canvas.min.js"></script></head><body style="margin:0; padding:10px; background:#f8fafc; position:relative;"><button id="dlBtnMM" onclick="downloadMarkmapImage('markmap-wrapper', 'Offline')" style="position:absolute; top:20px; right:20px; z-index:100; background:#10b981; color:white; border:none; padding:5px 10px; border-radius:5px; cursor:pointer; font-weight:bold;">📸 PNG HD</button><div id="markmap-wrapper" style="width:100%; height:550px; background:#ffffff; border:1px solid #e2e8f0; border-radius:8px; overflow:hidden;"><svg id="markmap" style="width:100%; height:100%;"></svg></div><script>const markdown = {markmap_json_str}; const {{ Transformer, Markmap }} = window.markmap; const {{ root }} = new Transformer().transform(markdown); Markmap.create('#markmap', null, root); window.downloadMarkmapImage = function(wrapperId, title) {{ const container = document.getElementById(wrapperId); const svgEl = container.querySelector('svg'); if (!svgEl) return; const btn = document.getElementById('dlBtnMM'); const originalText = btn.innerHTML; btn.innerHTML = "⏳..."; btn.disabled = true; const originalWidth = container.style.width; const originalHeight = container.style.height; const originalOverflow = container.style.overflow; const g = svgEl.querySelector('g'); const originalTransform = g ? g.getAttribute('transform') : null; if (g) g.setAttribute('transform', 'translate(50,50) scale(1)'); setTimeout(() => {{ const bbox = g ? g.getBBox() : svgEl.getBBox(); const padding = 50; const trueWidth = Math.max(bbox.width, 800) + (padding * 2); const trueHeight = Math.max(bbox.height, 600) + (padding * 2); container.style.width = trueWidth + 'px'; container.style.height = trueHeight + 'px'; container.style.overflow = 'visible'; svgEl.setAttribute('viewBox', `${{(bbox.x || 0) - padding}} ${{(bbox.y || 0) - padding}} ${{trueWidth}} ${{trueHeight}}`); html2canvas(container, {{ scale: 3, useCORS: true, backgroundColor: '#ffffff', width: trueWidth, height: trueHeight }}).then(canvas => {{ container.style.width = originalWidth; container.style.height = originalHeight; container.style.overflow = originalOverflow; if (g && originalTransform) g.setAttribute('transform', originalTransform); const link = document.createElement('a'); link.download = 'MindMap_' + title + '.png'; link.href = canvas.toDataURL('image/png', 1.0); link.click(); btn.innerHTML = originalText; btn.disabled = false; }}).catch(() => {{ btn.innerHTML = originalText; btn.disabled = false; }}); }}, 500); }};</script></body></html>""", height=600)
              # =====================================================================
                # FITUR BARU: SUNBURST HIERARCHY CHART (RODA ANATOMI RAPAT)
                # =====================================================================
                st.markdown("### ☀️ Sunburst Hierarchy Chart (Anatomi Rapat)")
                st.markdown("Klik pada salah satu potongan diagram untuk melakukan *Zoom-In*. Arahkan kursor (*hover*) untuk membaca teks lengkap.")
                
                # Kita oper data Markdown milik Markmap ke JS agar dikonversi jadi hierarki Sunburst yang sangat detail (Deep Layer)
                components.html(f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <script src="https://cdn.jsdelivr.net/npm/echarts@5.5.0/dist/echarts.min.js"></script>
                </head>
                <body style="margin:0; padding:10px; background:#f8fafc; position:relative;">
                    <div id="sunburst-chart" style="width:100%; height:600px; background:#ffffff; border:1px solid #e2e8f0; border-radius:12px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);"></div>
                    <script>
                        // 1. Ambil data markdown yang detail dari AI
                        const rawMarkdown = {markmap_json_str};
                        
                        // 2. Parser cerdas untuk mengubah Markdown menjadi struktur Tree (Cabang) ECharts
                        function parseMarkdownToSunburst(md) {{
                            const lines = md.split('\\n');
                            let root = {{ name: "Root", children: [] }};
                            let stack = [ {{level: 0, node: root}} ];

                            for (let i = 0; i < lines.length; i++) {{
                                let line = lines[i];
                                let trimmed = line.trimStart();
                                if (!trimmed) continue;

                                let level = 0;
                                let text = "";

                                // Cek apakah baris adalah Heading (#) atau Bullet (-)
                                let matchHeader = trimmed.match(/^(#+)\\s+(.*)/);
                                if (matchHeader) {{
                                    level = matchHeader[1].length;
                                    text = matchHeader[2];
                                }} else {{
                                    let matchList = trimmed.match(/^[-*]\\s+(.*)/);
                                    if (matchList) {{
                                        level = stack[stack.length - 1].level + 1;
                                        text = matchList[1];
                                    }}
                                }}

                                if (!text) continue;

                                // Bersihkan teks dari sisa-sisa karakter markdown seperti **
                                text = text.replace(/\\*\\*/g, '').trim();
                                let newNode = {{ name: text, children: [] }};

                                while (stack.length > 1 && stack[stack.length - 1].level >= level) {{
                                    stack.pop();
                                }}

                                let parent = stack[stack.length - 1].node;
                                parent.children.push(newNode);
                                stack.push({{ level: level, node: newNode }});
                            }}

                            // Beri value pada ujung rantai agar ukuran pie proporsional
                            function assignValues(node) {{
                                if (node.children.length === 0) {{
                                    node.value = 1;
                                }} else {{
                                    node.children.forEach(assignValues);
                                }}
                            }}
                            assignValues(root);

                            return root.children.length > 0 ? root.children : [{{name: "Data tidak tersedia", value: 1}}];
                        }}

                        const sunburstData = parseMarkdownToSunburst(rawMarkdown);

                        // === PERBAIKAN WARNA DITAMBAHKAN DI SINI ===
                        // Palet warna warni yang memanjakan mata
                        const colorPalette = ['#3b82f6', '#10b981', '#f59e0b', '#ef4444', '#8b5cf6', '#14b8a6', '#f43f5e', '#84cc16', '#0ea5e9', '#d946ef'];

                        // Jika root utama cuma 1 (Tema Rapat), kita warnai sub-topiknya agar berbeda
                        if (sunburstData.length === 1 && sunburstData[0].children) {{
                            sunburstData[0].itemStyle = {{ color: '#1e293b' }}; // Tema utama berwarna gelap/navy
                            sunburstData[0].children.forEach((child, index) => {{
                                child.itemStyle = {{ color: colorPalette[index % colorPalette.length] }};
                            }});
                        }} else {{
                            // Jika root lebih dari 1
                            sunburstData.forEach((child, index) => {{
                                child.itemStyle = {{ color: colorPalette[index % colorPalette.length] }};
                            }});
                        }}
                        // ===========================================

                        // 3. Render Grafik ECharts
                        var chartDom = document.getElementById('sunburst-chart');
                        var myChart = echarts.init(chartDom);
                        var option = {{
                            toolbox: {{
                                show: true,
                                feature: {{
                                    saveAsImage: {{
                                        show: true,
                                        title: 'Save PNG',
                                        name: 'Sunburst_Anatomi_Rapat',
                                        pixelRatio: 3, 
                                        iconStyle: {{ borderColor: '#10b981', borderWidth: 2 }}
                                    }}
                                }},
                                right: 20,
                                top: 20
                            }},
                            tooltip: {{ 
                                trigger: 'item',
                                formatter: function (info) {{ 
                                    return '<div style="max-width:300px; white-space:normal; font-family:sans-serif; font-size:13px; line-height:1.5;">' + info.name + '</div>'; 
                                }}
                            }},
                            series: {{
                                type: 'sunburst',
                                data: sunburstData,
                                radius: [0, '95%'],
                                sort: undefined,
                                emphasis: {{ focus: 'ancestor' }},
                                itemStyle: {{ 
                                    borderRadius: 5, 
                                    borderWidth: 1.5, 
                                    borderColor: '#ffffff' 
                                }},
                                label: {{ 
                                    show: true, 
                                    formatter: '{{b}}', 
                                    width: 85,
                                    overflow: 'break',
                                    minAngle: 12, 
                                    fontSize: 11, 
                                    fontWeight: 'bold',
                                    fontFamily: 'sans-serif',
                                    color: '#ffffff', 
                                    textBorderColor: 'rgba(0,0,0,0.6)', 
                                    textBorderWidth: 2 
                                }}
                            }}
                        }};
                        myChart.setOption(option);
                        window.addEventListener('resize', function() {{ myChart.resize(); }});
                    </script>
                </body>
                </html>
                """, height=620)
    # =====================================================================
    # TAB 3: SOP GENERATOR (KARS/JCI)
    # =====================================================================
    with tab_sop:
        st.markdown("### 🏥 AI Akreditasi SOP Generator")
        st.markdown("Generasi otomatis SOP Rumah Sakit berstandar Akreditasi KARS/JCI dengan output Flowchart Swimlane (PlantUML).")
        
        if not is_admin() and st.session_state.get("user_paket", "NON-AKTIF") not in ["EXECUTIVE", "MASTER"]:
            st.markdown("""<div style="background: rgba(239, 68, 68, 0.1); border: 2px solid #ef4444; border-radius: 12px; padding: 30px; text-align: center; margin-top: 20px;"><h1 style="font-size: 50px; margin: 0;">🔒</h1><h2 style="color: #ef4444; margin-top: 10px;">AKSES PREMIUM DIBUTUHKAN</h2><p style="color: #475569; font-size: 16px;">Fitur <b>AI Akreditasi SOP Generator</b> ini eksklusif hanya untuk pengguna paket <b>EXECUTIVE</b> dan <b>MASTER</b>. <br>Silakan upgrade paket Anda di tab Info Paket untuk membuka fitur ini.</p></div>""", unsafe_allow_html=True)
        else:
            with st.expander("🔑 Kredensial & Info Rumah Sakit", expanded=True):
                api_key_sop = st.text_input("LiteLLM / Gemini API Key", type="password", help="Masukkan API Key Anda", key="api_sop")
                logo_url = st.text_input("Link Logo Instansi (Copy Image Address)", value="https://upload.wikimedia.org/wikipedia/commons/thumb/e/eb/Coat_of_arms_of_Jakarta.svg/1280px-Coat_of_arms_of_Jakarta.svg.png", help="Paste link gambar logo (pastikan link berakhiran .png / .jpg)")
                col_i1, col_i2 = st.columns(2)
                with col_i1:
                    rs_name = st.text_input("Nama Instansi/RS", value="RSUD CONTOH")
                    title = st.text_input("Judul SOP", value="SOP Penerimaan Pasien Baru")
                    no_dok = st.text_input("Nomor Dokumen", value="001/SOP/2026")
                with col_i2:
                    dir_name = st.text_input("Nama Pimpinan", value="Dr. John Doe, MARS")
                    dir_nip = st.text_input("NIP Pimpinan", value="19800101 200501 1 001")

            st.markdown("#### 📂 Referensi (Max 5)")
            uploaded_files_sop = st.file_uploader("Upload Dokumen Referensi (PDF/Image)", accept_multiple_files=True, key="upload_sop")

           # --- BLOK TOMBOL GENERATE ---
            if st.button("🚀 Generate SOP Akreditasi", type="primary", use_container_width=True):
                if not api_key_sop: st.error("⚠️ API Key tidak ditemukan. Silakan input API Key!")
                elif not title: st.error("⚠️ Judul SOP wajib diisi!")
                elif len(uploaded_files_sop) > 5: st.error("⚠️ Maksimal 5 file referensi!")
                elif not is_admin() and st.session_state.get("user_kuota_ai", 0) <= 0: st.error("❌ Kuota AI Anda habis! Silakan lakukan Top-Up di tab Info Paket.")
                else:
                    if not is_admin():
                        st.session_state["user_kuota_ai"] -= 3
                        db.collection("users").document(st.session_state["user_uid"]).update({"kuota_ai": st.session_state["user_kuota_ai"]})
                        render_sidebar_profile()

                    try:
                        BASE_URL_KOBOI = "https://litellm.koboi2026.biz.id/v1/chat/completions"
                        headers = {'Authorization': f'Bearer {api_key_sop}', 'Content-Type': 'application/json'}
                        ref_instruction = "Buatlah SOP berdasarkan standar best practice akreditasi RS (KARS/JCI)."
                        file_data_array = []
                        if uploaded_files_sop:
                            ref_instruction = "SANGAT PENTING: Saya melampirkan referensi. WAJIB GUNAKAN ISI REFERENSI INI (dasar hukum, alat, langkah kerja)."
                            for f in uploaded_files_sop:
                                file_data_array.append({"type": "image_url", "image_url": { "url": f"data:{f.type};base64,{base64.b64encode(f.getvalue()).decode('utf-8')}" }})

                        prompt_cover = f"""Bertindaklah sebagai Ahli Hukum Kesehatan dan Auditor Akreditasi RS.
                        Tugas: Buat data teknis JSON untuk SOP: "{title}" di unit "{rs_name}".
                        {ref_instruction}
                        ATURAN PENTING: 1. DASAR HUKUM WAJIB TERBARU (UU No. 17 Tahun 2023). 2. NUMBERING (1., 2., dst). 3. KONSISTENSI.
                        OUTPUT JSON STRICT: {{ "dasar_hukum": ["1. ..."], "kualifikasi": ["1. ..."], "keterkaitan": ["1. ..."], "peralatan": ["1. ..."], "peringatan": "...", "pencatatan": ["1. ..."] }}"""

                        content_array_cover = [{"type": "text", "text": prompt_cover}]
                        if file_data_array: content_array_cover.extend(file_data_array)

                        with st.spinner("⏳ Robot Germic Membaca Referensi & Menyusun Struktur..."):
                            response_cover = requests.post(BASE_URL_KOBOI, headers=headers, json={"model": "gemini/gemini-2.5-flash", "messages": [{"role": "user", "content": content_array_cover}], "response_format": { "type": "json_object" }, "temperature": 0.2})
                            response_cover.raise_for_status()
                            data_cover = json.loads(clean_json_response(response_cover.json()['choices'][0]['message']['content']))

                        prompt_flow = f"""Tugas: Buat detail langkah kerja dan kode PlantUML untuk SOP "{title}".
                        ATURAN DECISION: Jika ada percabangan (disetujui/tidak), gunakan blok IF-ELSE di PlantUML.
                        ATURAN PLANTUML: 1. Swimlane SEBELUM 'start' (|Aktor|\\nstart). 2. Awali (:) akhiri (;). 3. JANGAN rubah Swimlane di dalam blok IF. 4. Tanpa 'repeat'/'while'.
                        KOLOM WAKTU: Angka dan Satuan (misal "5 Menit").
                        OUTPUT JSON: {{ "pelaksana_list": ["..."], "plantuml_code": "@startuml\\n|Petugas|\\nstart\\n:Kirim Draf;\\nstop\\n@enduml", "langkah": [ {{ "kegiatan": "...", "pelaksana": "...", "waktu": "...", "kelengkapan": "...", "output": "...", "keterangan": "-" }} ] }}"""

                        content_array_flow = [{"type": "text", "text": prompt_flow}]
                        if file_data_array: content_array_flow.extend(file_data_array)

                        with st.spinner("⏳ Robot Germic Meracik Alur Logika & Men-generate PlantUML..."):
                            response_flow = requests.post(BASE_URL_KOBOI, headers=headers, json={"model": "gemini/gemini-2.5-flash", "messages": [{"role": "user", "content": content_array_flow}], "response_format": { "type": "json_object" }, "temperature": 0.2})
                            response_flow.raise_for_status()
                            data_flow = json.loads(clean_json_response(response_flow.json()['choices'][0]['message']['content']))

                        st.toast("✅ Dokumen SOP Berhasil Dibuat!", icon="🎉")
                        preview_html, word_html = render_akreditasi_output({"cover": data_cover, "flow": data_flow}, rs_name, title, no_dok, dir_name, dir_nip, logo_url)

                        # SIMPAN HASILNYA KE MEMORI (SESSION STATE) AGAR TIDAK HILANG
                        st.session_state["sop_preview_html"] = preview_html
                        st.session_state["sop_word_html"] = word_html
                        st.session_state["sop_title"] = title

                    except requests.exceptions.RequestException as e: st.error(f"❌ Kesalahan API: {str(e)}")
                    except Exception as e: st.error(f"❌ Kesalahan pemrosesan: {str(e)}")

            # ====================================================================
            # MUNCULKAN PREVIEW & TOMBOL DOWNLOAD DI LUAR BLOK "IF BUTTON GENERATE"
            # PERHATIKAN: Posisi 'if' ini harus sejajar dengan 'if st.button(...)' di atas!
            # ====================================================================
            if st.session_state.get("sop_preview_html"):
                st.markdown("---")
                st.subheader("📄 Preview Dokumen SOP")
                st.components.v1.html(st.session_state["sop_preview_html"], height=1000, scrolling=True)
                
                st.download_button(
                    label="📄 DOWNLOAD WORD (SOP)", 
                    data=st.session_state["sop_word_html"].encode('utf-8-sig'), 
                    file_name=f"SOP_{st.session_state['sop_title'].replace(' ', '_')}.doc", 
                    mime="application/msword", 
                    use_container_width=True
                )
